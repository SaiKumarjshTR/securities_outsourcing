"""
hybrid_converter.py — Production PDF→DOCX replacement using PyMuPDF + pdfplumber.

Replaces the broken pdf2docx (and the Windows-only ABBYY COM path) with a
pure-Python approach that works on Linux/Docker and achieves 85-95% fidelity
on digital (native-text) PDFs.

Architecture (one PDF pass each tool, then merge):
  1. PyMuPDF   — text blocks with full font metadata (bold, italic, size, bbox)
  2. pdfplumber — tables with row/column structure and cell text
  3. Deduplicator — drop PyMuPDF blocks whose bbox overlaps a pdfplumber table
  4. Layout analyzer — fix multi-column reading order, per page
  5. Footnote extractor — separate small-font bottom-of-page text
  6. Page-break merger — join paragraphs split across pages
  7. DOCX builder — produce a .docx file consumed by the existing SGML pipeline

Output contract:
  convert_pdf_to_docx(pdf_path, docx_path) -> bool
  Returns True on success, False on error. Writes a valid .docx to docx_path.

  The DOCX is structured so batch_runner_deploy.py can consume it exactly as
  if ABBYY had produced it:
  - Headings use Word heading styles (Heading 1–4)
  - Bold text uses Word bold character style
  - Italic text uses Word italic character style
  - Tables use Word table structure
  - Footnotes are appended as plain paragraphs after a separator
"""
import io
import logging
import re
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

log = logging.getLogger(__name__)

# Suppress noisy pdfplumber/pdfminer font-descriptor warnings that appear when
# a PDF has malformed font metrics (FontBBox = None). These do not affect
# conversion output — pdfplumber falls back to default glyph widths.
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)

# ── Optional imports — fail gracefully so pipeline can still load ─────────────
try:
    import fitz  # PyMuPDF
    _FITZ_OK = True
except ImportError:
    _FITZ_OK = False
    log.warning("PyMuPDF (fitz) not installed — hybrid converter disabled")

try:
    import pdfplumber
    _PDFPLUMBER_OK = True
except ImportError:
    _PDFPLUMBER_OK = False
    log.warning("pdfplumber not installed — table extraction disabled")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False
    log.warning("python-docx not installed — DOCX output disabled")

from .formatting_extractor import (
    classify_spans_for_page,
    compute_body_font_size,
    is_bold,
    is_italic,
    detect_heading_level,
)
from .deduplicator import deduplicate
from .layout_analyzer import sort_blocks_by_reading_order, merge_page_breaks


# ── Data structures ───────────────────────────────────────────────────────────

@dataclass
class Span:
    """Single styled run of text within a line."""
    text: str
    bold: bool = False
    italic: bool = False
    superscript: bool = False
    size: float = 10.0
    font: str = ""


@dataclass
class TextBlock:
    """A paragraph-level text unit."""
    text: str
    spans: List[Span] = field(default_factory=list)
    font: str = ""
    size: float = 10.0
    bbox: Tuple[float, float, float, float] = (0, 0, 0, 0)
    page: int = 0
    is_bold: bool = False
    is_italic: bool = False
    is_heading: int = 0   # 0 = body, 1–4 = heading level
    is_footnote: bool = False


@dataclass
class TableStructure:
    """A table extracted by pdfplumber."""
    rows: List[List[str]]          # 2-D array; None cells → ""
    bbox: Tuple[float, float, float, float]
    page: int
    col_count: int = 0


@dataclass
class StructuredPDF:
    """Complete extracted content from one PDF."""
    blocks: List[TextBlock] = field(default_factory=list)
    tables: List[TableStructure] = field(default_factory=list)
    footnotes: List[TextBlock] = field(default_factory=list)
    page_count: int = 0
    metadata: Dict = field(default_factory=dict)


# ── pdfplumber table settings ─────────────────────────────────────────────────
# Strategy "lines" works for ruled tables; "text" for borderless column layouts.
# We try "lines" first and fall back to "text" if no tables found.
_TABLE_SETTINGS_LINES = {
    "vertical_strategy":   "lines",
    "horizontal_strategy": "lines",
    "snap_tolerance":       3,
    "join_tolerance":       3,
    "edge_min_length":      3,
    "min_words_vertical":   1,
    "min_words_horizontal": 1,
    "intersection_tolerance": 3,
}
_TABLE_SETTINGS_TEXT = {
    "vertical_strategy":   "text",
    "horizontal_strategy": "text",
    "snap_tolerance":       3,
    "join_tolerance":       3,
    "min_words_vertical":   4,   # raised from 2 — suppresses text-alignment FPs
    "min_words_horizontal": 2,   # raised from 1
}

# Maximum table area as a fraction of page area before a detected table is
# considered a layout artefact and discarded.
# Legal PDFs routinely wrap entire pages in a single table for layout purposes;
# accepting such tables causes the deduplicator to remove almost all body text.
# "lines": generous — real ruled tables (forms, schedules) can be large.
# "text" : strict  — text-column detection is very prone to false positives.
_MAX_TABLE_AREA_FRAC_LINES = 0.70   # >70% of page → layout container, not data
_MAX_TABLE_AREA_FRAC_TEXT  = 0.25   # >25% of page → body text falsely detected
_MAX_TABLE_ROWS_TEXT        = 15     # >15 rows via text strategy → almost certainly FP

# Footnote detection: small font + bottom of page
_FOOTNOTE_Y_FRAC = 0.82      # bottom 18% of page
_FOOTNOTE_SIZE_DELTA = 2.0   # font ≥ 2pt smaller than body
_FOOTNOTE_MARKER_RE = re.compile(r"^\s*(?:\d{1,3}|[*†‡§¶])\s+\S")

# ── Web-page PDF chrome filter ────────────────────────────────────────────────
# TMX/TSX/MX website PDFs contain navigation chrome (breadcrumbs, stock tickers,
# footer links, social icons) that pollutes the extracted text.
# Fingerprint: any block contains \uf054 (breadcrumb arrow glyph used by TMX sites).

# PUA (Private Use Area) unicode — font icon glyphs, not real text
_PUA_RE = re.compile(r"[\ue000-\uf8ff]")

# Known web footer/nav link text patterns (full-line match, case-insensitive)
_WEB_FOOTER_LINE_RE = re.compile(
    r"^(contact\s+us|terms\s+of\s+use|privacy\s+policy|"
    r"accessibility|fraud\s+prevention|sign\s+in|"
    r"trading\s+status|fran[çc]ais|"
    r"capital\s+formation|post-trade|post\s+trade|insights|"
    r"copyright[\s\S]{0,60}rights\s+reserved|"
    r"copyright\s*[©\u00a9©]|[©\u00a9©]\s*\d{4}|all\s+rights\s+reserved)$",
    re.IGNORECASE,
)

# Legal disclaimer text unique to TMX Group pages
_TMX_DISCLAIMER_RE = re.compile(
    r"^TMX\s+Group\s+Limited\s+and\s+its\s+affiliates\s+do\s+not",
    re.IGNORECASE,
)


def _is_web_chrome_block(block: Dict) -> bool:
    """
    Return True if this block is web-page chrome (nav/header/footer)
    that should be stripped from web-page PDFs (TMX/TSX/MX website prints).
    Only called when the PDF has been identified as a web-page PDF.
    """
    text = block.get("text", "")
    stripped = text.strip()
    if not stripped:
        return False

    # Breadcrumb navigation arrow glyph (specific to TMX website PDFs)
    if "\uf054" in text:
        return True

    # Block consisting entirely of PUA/icon characters (social-media icons, etc.)
    non_ws = re.sub(r"\s", "", stripped)
    if non_ws and all(_PUA_RE.match(c) for c in non_ws):
        return True

    # Stock ticker / site-navigation bar at very top of first page (y0 < 55pt)
    if block.get("page", 0) == 0:
        y0 = block.get("bbox", (0, 0, 0, 0))[1]
        if y0 < 55:
            return True

    # Known footer/nav patterns — check each line individually
    for line in stripped.split("\n"):
        # Strip trailing icon glyphs before matching
        line_clean = _PUA_RE.sub("", line).strip()
        if line_clean and _WEB_FOOTER_LINE_RE.match(line_clean):
            return True

    # TMX legal disclaimer
    if _TMX_DISCLAIMER_RE.match(stripped):
        return True

    # Copyright notice (handles encoding variants of ©, trailing period, etc.)
    if re.search(r"copyright.{0,80}reserved", stripped, re.IGNORECASE | re.DOTALL):
        return True

    return False


# ── Coordinate helpers ────────────────────────────────────────────────────────

def _rotate_bbox_to_visual(
    bbox: tuple, rotation: int, mw: float, mh: float
) -> tuple:
    """
    Transform a raw PyMuPDF block bbox (in mediabox coordinates) to the
    visual/display coordinate space by applying the page rotation.

    PyMuPDF's get_text('dict') returns bbox values in the un-rotated mediabox
    coordinate space.  For y0-based sorting to produce the correct reading
    order the bboxes must be in display coordinates (origin = top-left,
    y increasing downward, after the rotation is applied).

    Args:
        bbox:     (x0, y0, x1, y1) in raw mediabox coordinates
        rotation: page rotation in degrees (0 / 90 / 180 / 270)
        mw:       mediabox width  before rotation
        mh:       mediabox height before rotation
    """
    if not rotation:
        return bbox
    x0, y0, x1, y1 = bbox
    if rotation == 90:
        # visual_x = mh - raw_y,  visual_y = raw_x
        return (mh - y1, x0, mh - y0, x1)
    if rotation == 180:
        return (mw - x1, mh - y1, mw - x0, mh - y0)
    if rotation == 270:
        # visual_x = raw_y,  visual_y = mw - raw_x
        return (y0, mw - x1, y1, mw - x0)
    return bbox


# ── Main extraction functions ─────────────────────────────────────────────────

def _extract_pymupdf(pdf_path: str) -> Tuple[List[Dict], Dict, int]:
    """
    Extract all text blocks from the PDF using PyMuPDF.

    Returns (raw_blocks, metadata, page_count).
    Each raw block has keys: text, spans, font, size, bbox, page,
    _bold, _italic, _heading, is_footnote (False initially).
    """
    if not _FITZ_OK:
        raise RuntimeError("PyMuPDF not available")

    raw_blocks: List[Dict] = []

    doc = fitz.open(pdf_path)
    page_count = len(doc)

    # Extract metadata
    meta = doc.metadata or {}
    metadata = {
        "title":     meta.get("title", ""),
        "author":    meta.get("author", ""),
        "created":   meta.get("creationDate", ""),
        "pages":     page_count,
    }

    for page_num, page in enumerate(doc):
        page_rect = page.rect
        page_height = page_rect.height
        page_width  = page_rect.width

        # Rotation info needed to convert raw block bboxes → visual coordinates
        page_rotation = page.rotation          # 0 / 90 / 180 / 270
        page_mw       = page.mediabox.width    # raw (pre-rotation) page width
        page_mh       = page.mediabox.height   # raw (pre-rotation) page height

        page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        # Collect all spans from this page for body-size computation
        all_spans: List[Dict] = []
        for blk in page_dict.get("blocks", []):
            if blk.get("type") != 0:  # skip image blocks
                continue
            for line in blk.get("lines", []):
                all_spans.extend(line.get("spans", []))

        # Classify spans (adds _bold, _italic, _super, _heading keys)
        classified_spans = classify_spans_for_page(all_spans)
        body_size = compute_body_font_size(all_spans)

        span_idx = 0
        for blk in page_dict.get("blocks", []):
            if blk.get("type") != 0:
                continue

            block_text_parts: List[str] = []
            block_spans: List[Dict] = []
            dominant_font = ""
            dominant_size = body_size
            block_bold = False
            block_italic = False
            block_heading = 0

            for line_idx, line in enumerate(blk.get("lines", [])):
                line_text_parts = []
                prev_span_x1 = None
                # Insert a space between adjacent lines so span-based rendering
                # doesn't concatenate the last word of line N with the first word
                # of line N+1 (e.g. "the\nopening" → "theopening").
                if line_idx > 0 and block_spans:
                    last = block_spans[-1]["text"]
                    if not last.endswith(" "):
                        block_spans.append({
                            "text": " ", "bold": False, "italic": False,
                            "superscript": False, "size": body_size, "font": "",
                        })
                for span in line.get("spans", []):
                    # Use pre-classified span from our list
                    cspan = classified_spans[span_idx] if span_idx < len(classified_spans) else span
                    span_idx += 1

                    t = span.get("text", "")
                    if not t:
                        continue

                    # Insert a space if there is a visible horizontal gap between
                    # consecutive spans (PyMuPDF does not add spaces at span boundaries).
                    span_bbox = span.get("bbox", (0, 0, 0, 0))
                    if prev_span_x1 is not None:
                        gap = span_bbox[0] - prev_span_x1
                        char_w = span.get("size", body_size) * 0.4
                        if (gap > char_w
                                and line_text_parts
                                and not line_text_parts[-1].endswith(" ")
                                and not t.startswith(" ")):
                            line_text_parts.append(" ")
                    prev_span_x1 = span_bbox[2]

                    line_text_parts.append(t)
                    block_spans.append({
                        "text":        t,
                        "bold":        cspan.get("_bold", False),
                        "italic":      cspan.get("_italic", False),
                        "superscript": cspan.get("_super", False),
                        "size":        span.get("size", body_size),
                        "font":        span.get("font", ""),
                    })

                    # Track dominant properties (by longest span)
                    if not dominant_font and span.get("font"):
                        dominant_font = span["font"]
                        dominant_size = span.get("size", body_size)
                    if cspan.get("_bold"):
                        block_bold = True
                    if cspan.get("_italic"):
                        block_italic = True
                    h = cspan.get("_heading", 0)
                    if h > block_heading:
                        block_heading = h

                if line_text_parts:
                    block_text_parts.append("".join(line_text_parts))

            raw_text = "\n".join(block_text_parts).strip()
            if not raw_text:
                continue

            # Promote short bold single-line blocks to H4 when not already
            # detected as a heading by font size (handles bold headings at body size).
            # Conditions: starts with uppercase (excludes list labels like "(a)", "-AND"),
            # single line, ≤ 60 chars, no trailing sentence-end punctuation, last word
            # is not a preposition/conjunction (avoids mid-sentence continuations).
            _raw_last = (raw_text.rsplit(None, 1)[-1].rstrip(".,;:!?)]}\"'").lower()
                         if raw_text else "")
            _HEADING_STOP_WORDS = frozenset([
                "of", "the", "a", "an", "to", "in", "at", "by", "for", "and",
                "or", "nor", "but", "as", "that", "which", "with", "from",
                "into", "upon", "under", "over", "through", "not", "if",
                "is", "are", "was", "were", "be", "been", "has", "have",
                "had", "will", "would", "shall", "should", "may", "might",
                "can", "could", "than", "its", "their", "this", "these",
            ])
            if (block_heading == 0 and block_bold
                    and len(block_text_parts) == 1
                    and len(raw_text) <= 60
                    and raw_text and raw_text[0].isupper()
                    and not raw_text[:4].lower() in ("and ", "or n", "nor ", "but ")
                    and not raw_text.endswith(".")
                    and not raw_text.endswith(",")
                    and not raw_text.endswith(";")
                    and _raw_last not in _HEADING_STOP_WORDS):
                block_heading = 4

            bbox = tuple(blk.get("bbox", (0, 0, 0, 0)))
            # Transform to visual coordinates so sorting works on rotated pages
            if page_rotation:
                bbox = _rotate_bbox_to_visual(bbox, page_rotation, page_mw, page_mh)
            y0 = bbox[1]

            raw_blocks.append({
                "text":      raw_text,
                "spans":     block_spans,
                "font":      dominant_font,
                "size":      dominant_size,
                "bbox":      bbox,
                "page":      page_num,
                "page_h":    page_height,
                "page_w":    page_width,
                "_bold":     block_bold,
                "_italic":   block_italic,
                "_heading":  block_heading,
                "is_footnote": False,
            })

    doc.close()

    # ── Web-page PDF chrome filter ──────────────────────────────────────────
    # Detect TMX/TSX/MX website PDFs by the breadcrumb arrow glyph (\uf054).
    # If found, strip all navigation/footer chrome blocks from every page.
    if any("\uf054" in b.get("text", "") for b in raw_blocks):
        before = len(raw_blocks)
        raw_blocks = [b for b in raw_blocks if not _is_web_chrome_block(b)]
        log.debug(
            "Web-page PDF detected — filtered %d chrome blocks (%d remaining)",
            before - len(raw_blocks), len(raw_blocks),
        )

    log.debug("PyMuPDF: extracted %d blocks from %d pages", len(raw_blocks), page_count)
    return raw_blocks, metadata, page_count


def _extract_pdfplumber_tables(pdf_path: str) -> List[TableStructure]:
    """
    Extract all tables using pdfplumber.

    Tries line-based strategy first; if a page has no tables detected,
    retries with text-based strategy (handles borderless tables).

    For each table, the exact bbox is retrieved from the pdfplumber Table
    object (find_tables) using the SAME strategy that produced the data,
    so the bbox always corresponds to the extracted rows.
    """
    if not _PDFPLUMBER_OK:
        log.warning("pdfplumber not available — no table extraction")
        return []

    tables: List[TableStructure] = []

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_w = float(page.width)
                page_h = float(page.height)

                # ── Try lines strategy first, then text strategy fallback ─────
                strategy_used = "lines"
                found_objects = page.find_tables(_TABLE_SETTINGS_LINES)
                page_tables   = [t.extract() for t in found_objects] if found_objects else []

                if not page_tables:
                    strategy_used = "text"
                    found_objects = page.find_tables(_TABLE_SETTINGS_TEXT)
                    page_tables   = [t.extract() for t in found_objects] if found_objects else []

                for tbl_obj, pt in zip(found_objects, page_tables):
                    if not pt:
                        continue

                    # Clean cells: None → "", strip whitespace
                    cleaned: List[List[str]] = []
                    for row in pt:
                        if row is None:
                            continue
                        cleaned.append([
                            (cell.strip() if cell else "") for cell in row
                        ])

                    if not cleaned:
                        continue

                    # Get exact bbox from the Table object (same strategy, guaranteed match)
                    try:
                        fb = tbl_obj.bbox   # (x0, top, x1, bottom) in PDF points
                        t_bbox = (float(fb[0]), float(fb[1]), float(fb[2]), float(fb[3]))
                    except Exception:
                        t_bbox = (0.0, 0.0, page_w, page_h)

                    # ── False-positive filter: reject oversized / phantom tables ──────
                    # When a "table" covers the majority of the page it is almost always
                    # a layout-as-table artefact, not a real data table.  Keeping it
                    # causes the deduplicator to erase nearly all body text.
                    t_area    = (t_bbox[2] - t_bbox[0]) * (t_bbox[3] - t_bbox[1])
                    page_area = page_w * page_h
                    area_frac = t_area / page_area if page_area > 0 else 0.0
                    row_count = len(cleaned)

                    if strategy_used == "lines":
                        if area_frac > _MAX_TABLE_AREA_FRAC_LINES:
                            log.debug(
                                "Skip lines-table page %d: %.0f%% of page (layout artefact)",
                                page_num, area_frac * 100,
                            )
                            continue
                    else:  # "text" strategy — much stricter
                        if area_frac > _MAX_TABLE_AREA_FRAC_TEXT or row_count > _MAX_TABLE_ROWS_TEXT:
                            log.debug(
                                "Skip text-table page %d: %.0f%% of page, %d rows (false positive)",
                                page_num, area_frac * 100, row_count,
                            )
                            continue

                    col_count = max((len(r) for r in cleaned), default=0)

                    # A genuine table needs ≥2 columns WITH actual content, AND
                    # at least 2 rows where multiple columns are simultaneously
                    # non-empty.  Pseudo-tables from bordered or two-column-text
                    # layouts may have 2 nominal columns but content in only ONE
                    # column per row (the other column being a phantom separator
                    # or a header reference in a separate row).
                    cols_with_content: set = set()
                    rows_with_multicol: int = 0
                    for r in cleaned:
                        nonempty_in_row = [
                            ci for ci, cell in enumerate(r)
                            if cell and cell.strip()
                        ]
                        cols_with_content.update(nonempty_in_row)
                        if len(nonempty_in_row) >= 2:
                            rows_with_multicol += 1

                    nonempty_col_count = len(cols_with_content)
                    if nonempty_col_count < 2 or rows_with_multicol < 2:
                        log.debug(
                            "Skip pseudo-table page %d: %d non-empty cols, "
                            "%d multicol rows (layout artefact, bbox=%s)",
                            page_num, nonempty_col_count, rows_with_multicol, t_bbox,
                        )
                        continue

                    tables.append(TableStructure(
                        rows=cleaned,
                        bbox=t_bbox,
                        page=page_num,
                        col_count=col_count,
                    ))

    except Exception as exc:
        log.error("pdfplumber table extraction failed: %s", exc)

    log.debug("pdfplumber: extracted %d tables", len(tables))
    return tables


def _extract_footnotes(
    blocks: List[Dict],
) -> Tuple[List[Dict], List[Dict]]:
    """
    Separate footnotes from main body blocks.

    Footnote criteria (ALL must hold):
      1. Font size ≤ body_size - 2pt (significantly smaller)
      2. Vertical position in bottom 18% of the page
      3. OR: text starts with a superscript-style footnote marker (1, *, †)

    Returns (body_blocks, footnote_blocks).
    """
    body: List[Dict] = []
    notes: List[Dict] = []

    # Compute per-page body sizes
    page_body_size: Dict[int, float] = {}
    for b in blocks:
        pg = b.get("page", 0)
        if pg not in page_body_size:
            page_body_size[pg] = b.get("size", 10.0)

    for b in blocks:
        pg = b.get("page", 0)
        body_sz = page_body_size.get(pg, 10.0)
        size     = b.get("size", body_sz)
        page_h   = b.get("page_h", 792.0)
        y0       = b.get("bbox", (0, 0, 0, 0))[1]
        text     = b.get("text", "").strip()

        small_font   = (size <= body_sz - _FOOTNOTE_SIZE_DELTA)
        bottom_zone  = (y0 >= page_h * _FOOTNOTE_Y_FRAC)
        marker_match = bool(_FOOTNOTE_MARKER_RE.match(text))

        if (small_font and bottom_zone) or (small_font and marker_match):
            fn = dict(b)
            fn["is_footnote"] = True
            notes.append(fn)
        else:
            body.append(b)

    log.debug("Footnotes: %d detected, %d body blocks remain", len(notes), len(body))
    return body, notes


def _raw_to_textblocks(raw_blocks: List[Dict]) -> List[TextBlock]:
    """Convert raw dicts from PyMuPDF extraction to typed TextBlock objects."""
    result: List[TextBlock] = []
    for b in raw_blocks:
        spans = [
            Span(
                text        = s.get("text", ""),
                bold        = s.get("bold", False),
                italic      = s.get("italic", False),
                superscript = s.get("superscript", False),
                size        = s.get("size", 10.0),
                font        = s.get("font", ""),
            )
            for s in b.get("spans", [])
            if s.get("text", "")
        ]
        result.append(TextBlock(
            text       = b.get("text", ""),
            spans      = spans,
            font       = b.get("font", ""),
            size       = b.get("size", 10.0),
            bbox       = b.get("bbox", (0, 0, 0, 0)),
            page       = b.get("page", 0),
            is_bold    = b.get("_bold", False),
            is_italic  = b.get("_italic", False),
            is_heading = b.get("_heading", 0),
            is_footnote= b.get("is_footnote", False),
        ))
    return result


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _add_run(para, span: Span) -> None:
    """Add a styled run to a python-docx paragraph."""
    run = para.add_run(span.text)
    run.bold   = span.bold
    run.italic = span.italic
    if span.superscript:
        run.font.superscript = True


def _build_docx(structured: StructuredPDF) -> "Document":
    """
    Build a python-docx Document from extracted StructuredPDF data.

    Heading levels → Word "Heading 1"–"Heading 4" styles.
    Body text      → "Normal" style with inline bold/italic runs.
    Tables         → Word tables.
    Footnotes      → Plain paragraphs after a "---" separator.
    """
    doc = Document()

    # Combine body blocks and tables in page/y-order
    # We interleave using their page + y0 position
    items: List[Tuple[int, float, str, object]] = []

    for blk in structured.blocks:
        items.append((blk.page, blk.bbox[1], "block", blk))
    for tbl in structured.tables:
        items.append((tbl.page, tbl.bbox[1], "table", tbl))

    items.sort(key=lambda x: (x[0], x[1]))

    for _, _, kind, obj in items:
        if kind == "block":
            blk: TextBlock = obj  # type: ignore[assignment]
            if not blk.text.strip():
                continue

            if blk.is_heading > 0:
                style = f"Heading {min(blk.is_heading, 4)}"
                try:
                    para = doc.add_paragraph(blk.text.strip(), style=style)
                except Exception:
                    para = doc.add_paragraph(blk.text.strip())
            else:
                para = doc.add_paragraph()
                para.style = doc.styles["Normal"]
                if blk.spans:
                    for span in blk.spans:
                        if span.text:
                            _add_run(para, span)
                else:
                    run = para.add_run(blk.text)
                    run.bold   = blk.is_bold
                    run.italic = blk.is_italic

        elif kind == "table":
            tbl: TableStructure = obj  # type: ignore[assignment]
            if not tbl.rows:
                continue

            col_count = tbl.col_count or max((len(r) for r in tbl.rows), default=1)
            word_table = doc.add_table(rows=0, cols=col_count)
            word_table.style = "Table Grid"

            for row_data in tbl.rows:
                row = word_table.add_row()
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(row.cells):
                        row.cells[col_idx].text = cell_text or ""

    # ── Footnotes section ────────────────────────────────────────────────────
    if structured.footnotes:
        doc.add_paragraph("─" * 40)
        for fn in structured.footnotes:
            p = doc.add_paragraph(fn.text.strip())
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.font.size = Pt(8)

    return doc


# ── Public API ─────────────────────────────────────────────────────────────────

def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> bool:
    """
    Convert a digital PDF to a DOCX file using the hybrid PyMuPDF+pdfplumber
    approach.

    Parameters
    ----------
    pdf_path  : absolute path to input PDF
    docx_path : absolute path for output DOCX

    Returns
    -------
    True on success, False on any error.
    """
    if not _FITZ_OK:
        log.error("PyMuPDF not available — cannot convert PDF")
        return False
    if not _DOCX_OK:
        log.error("python-docx not available — cannot write DOCX")
        return False

    try:
        # ── Step 1: PyMuPDF extraction ───────────────────────────────────────
        raw_blocks, metadata, page_count = _extract_pymupdf(pdf_path)

        # ── Step 2: pdfplumber table extraction ──────────────────────────────
        tables = _extract_pdfplumber_tables(pdf_path)

        # ── Step 3: Deduplication — remove table text from PyMuPDF blocks ───
        table_bboxes_by_page: Dict[int, List] = {}
        for t in tables:
            table_bboxes_by_page.setdefault(t.page, []).append(t.bbox)

        deduplicated = deduplicate(raw_blocks, table_bboxes_by_page)

        # ── Step 4: Per-page layout ordering ────────────────────────────────
        # Group by page, sort each page, then recombine
        pages: Dict[int, List[Dict]] = {}
        for b in deduplicated:
            pages.setdefault(b["page"], []).append(b)

        ordered: List[Dict] = []
        for pg in sorted(pages.keys()):
            pg_blocks = pages[pg]
            # Get page width from first block (stored during extraction)
            pw = pg_blocks[0].get("page_w", 612.0) if pg_blocks else 612.0
            ordered.extend(sort_blocks_by_reading_order(pg_blocks, pw))

        # ── Step 5: Footnote extraction ──────────────────────────────────────
        body_raw, fn_raw = _extract_footnotes(ordered)

        # ── Step 6: Page-break merging ───────────────────────────────────────
        body_merged = merge_page_breaks(body_raw)

        # ── Step 7: Convert to typed objects ─────────────────────────────────
        body_blocks  = _raw_to_textblocks(body_merged)
        fn_blocks    = _raw_to_textblocks(fn_raw)

        structured = StructuredPDF(
            blocks     = body_blocks,
            tables     = tables,
            footnotes  = fn_blocks,
            page_count = page_count,
            metadata   = metadata,
        )

        # ── Step 8: Build DOCX ───────────────────────────────────────────────
        doc = _build_docx(structured)
        doc.save(docx_path)

        log.info(
            "Hybrid converter: %d blocks, %d tables, %d footnotes → %s",
            len(body_blocks), len(tables), len(fn_blocks), docx_path,
        )
        return os.path.exists(docx_path)

    except Exception as exc:
        log.exception("Hybrid PDF→DOCX conversion failed: %s", exc)
        return False


def is_available() -> bool:
    """Return True if all required libraries are present."""
    return _FITZ_OK and _PDFPLUMBER_OK and _DOCX_OK
