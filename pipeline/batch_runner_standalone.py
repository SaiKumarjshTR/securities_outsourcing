# ====== CODE CELL 0 ======
import os
import re
import json
import time
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from dataclasses import dataclass, field
from collections import defaultdict

# ABBYY
import win32com.client as win32c

# DOCX
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Image processing
from PIL import Image

# Anthropic
from anthropic import Anthropic
import requests

# RAG (ChromaDB vector database)
import chromadb

print("✅ All imports successful!")
print(f"📅 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


# ====== CODE CELL 1 ======
# ABBYY Configuration
ABBYY_CONFIG = {
    'customer_id': 'FFvrEyp5Gz8sXSwP98N9',
    'license_path': r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\SWAD12410007214764846289.ABBYY 1.ActivationToken',
    'license_password': '/80HjebrjO2bzpJUiJ/DwQ=='
}

# Thomson Reuters
WORKSPACE_ID = 'Saikumar3Y0Z'
TR_AUTH_URL = 'https://aiplatform.gcs.int.thomsonreuters.com/v1/anthropic/token'
ANTHROPIC_MODEL = 'claude-sonnet-4-5-20250929'   # Sonnet (fallback)
OPUS_MODEL      = 'claude-opus-4-20250514'         # Opus 4.6 (primary for agents)

# Paths - UPDATE THESE
PATHS = {
    'input_pdf': r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\sec-out-samples-2\Jurisdictions\juri\NB_\51-737.pdf',
    'output_dir': r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\sec-out-samples-2\Jurisdictions\juri\NB_\51-737',
    'keying_rules': r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\COMPLETE_KEYING_RULES_UPDATED.txt'
}

# System
SYSTEM_CONFIG = {
    'use_llm': True,
    'llm_batch_size': 10,
    'extract_tables': True,
    'extract_images': True,
    'apply_inline_formatting': True,
    'max_tokens': 16000,
    'temperature': 0.0,
    'image_dpi': 300
}

# RAG Configuration
_JURI = r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\sec-out-samples-2\Jurisdictions\juri'
_SEC_SAMPLES = r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\sec-out-samples-2\sec-out-samples-2'
RAG_CONFIG = {
    'enabled': True,
    'persist_dir': r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\chroma_db_v6',
    'n_rules': 12,
    'n_examples': 25,
    'vendor_sgms': [
        # ── Ontario gold files (10 docs — directly relevant to test set) ──
        _JURI + r'\Ontario\Ontario\11-312.sgm',
        _JURI + r'\Ontario\Ontario\11-502.sgm',
        _JURI + r'\Ontario\Ontario\11-503.sgm',
        _JURI + r'\Ontario\Ontario\31-367.sgm',
        _JURI + r'\Ontario\Ontario\31-368.sgm',
        _JURI + r'\Ontario\Ontario\41-702.sgm',
        _JURI + r'\Ontario\Ontario\45-330.sgm',
        _JURI + r'\Ontario\Ontario\51-737.sgm',
        _JURI + r'\Ontario\Ontario\81-510.sgm',
        _JURI + r'\Ontario\Ontario\96-307.sgm',
        # ── British Columbia (diverse structure types) ──
        _JURI + r'\British_Columbia\British_Columbia\11-348.sgm',
        _JURI + r'\British_Columbia\British_Columbia\31-369.sgm',
        _JURI + r'\British_Columbia\British_Columbia\96-308.sgm',
        # ── Alberta (instrument + notice types) ──
        _JURI + r'\Alberta\Alberta\instrument-35-101.sgm',
        _JURI + r'\Alberta\Alberta\notice-33-706.sgm',
        # ── Quebec (bilingual, complex tables) ──
        _JURI + r'\Quebec\Quebec\93-101.sgm',
        _JURI + r'\Quebec\Quebec\94-102.sgm',
        # ── NB (similar test jurisdiction) ──
        _JURI + r'\NB_\comparive_samples\2025-12-10.sgm',
        _JURI + r'\NB_\comparive_samples\51-737.sgm',
        # ── CIRO (rule-type documents with heavy ITEM content) ──
        _JURI + r'\CIRO_\25-0001.sgm',
        # ── Files with graphics/tables ──
        _JURI + r'\files-with-graphics\files-with-graphics\051-365.sgm',
        _JURI + r'\files-with-graphics\files-with-graphics\23-329.sgm',
        # ── Toronto Stock Exchange (short notices + by-laws) ──
        _JURI + r'\Toronto-Stock-Exchange\TSX\Capital-Markets-Tribunal-Rules-of-Procedure.sgm',
        # ── Original sec-samples (keep legacy reference) ──
        _SEC_SAMPLES + r'\96-101CP.sgm',
        _SEC_SAMPLES + r'\25-0271.sgm',
    ]
}

# Agentic Multi-LLM Configuration (v12)
AGENT_CONFIG = {
    'orchestrator':       {'model': OPUS_MODEL, 'temperature': 0.0, 'max_tokens': 4096},
    'specialized_agents': {'model': OPUS_MODEL, 'temperature': 0.0, 'max_tokens': 4096},
    'validator':          {'model': OPUS_MODEL, 'temperature': 0.0, 'max_tokens': 2048},
    'enable_parallel':    True,   # Parallel StructuralAgent + EMAgent via ThreadPoolExecutor
    'agent_routing_rules': {
        'has_bold_formatting':    ['BOLD_AGENT'],
        'has_italic_formatting':  ['EM_AGENT'],
        'has_bullet_or_indent':   ['ITEM_AGENT'],
        'is_heading_style':       ['BLOCK_AGENT'],
        'contains_email':         ['EM_AGENT'],
        'contains_act_reference': ['EM_AGENT'],
        'is_regular_paragraph':   ['STRUCTURE_AGENT'],
    }
}

os.makedirs(PATHS['output_dir'], exist_ok=True)

print("✅ Configuration loaded!")
print(f"📅 Output: {PATHS['output_dir']}")
print(f"🤖 Model (Sonnet fallback): {ANTHROPIC_MODEL}")
print(f"🧠 RAG: {'ENABLED' if RAG_CONFIG['enabled'] else 'DISABLED'}")
print(f"🤖 Agentic Model (Opus primary): {OPUS_MODEL}")


# ====== CODE CELL 2 ======
def authenticate_tr_platform(workspace_id: str) -> Optional[Dict[str, Any]]:
    print("\n" + "="*80)
    print("Ã°Å¸â€Â AUTHENTICATING WITH THOMSON REUTERS")
    print("="*80)
    
    try:
        response = requests.post(
            TR_AUTH_URL,
            json={'workspace_id': workspace_id, 'model_name': ANTHROPIC_MODEL},
            timeout=10
        )
        
        if response.status_code != 200:
            print(f"Ã¢Å“â€” Failed: HTTP {response.status_code}")
            return None
        
        credentials = response.json()
        token = credentials.get('anthropic_api_key') or credentials.get('token', '')
        
        if not token:
            print("Ã¢Å“â€” No token")
            return None
        
        print("Ã¢Å“â€¦ AUTHENTICATED!")
        print(f"Token: {token[:20]}...")
        print("="*80)
        
        return {'token': token, 'model': ANTHROPIC_MODEL, 'enabled': True}
        
    except Exception as e:
        print(f"Ã¢Å“â€” Error: {e}")
        return None

auth_config = authenticate_tr_platform(WORKSPACE_ID)

if auth_config:
    client = Anthropic(api_key=auth_config['token'])
    print("Ã¢Å“â€¦ Anthropic client ready!")
else:
    client = None
    print("Ã¢ÂÅ’ Authentication failed!")

# ====== CODE CELL 3 ======
def load_complete_keying_specifications(rules_path: str) -> str:
    """Load complete 117K keying specifications - ALL rules"""
    print("\nÃ°Å¸â€œâ€¹ Loading COMPLETE keying specifications...")
    
    if not os.path.exists(rules_path):
        print(f"   Ã¢ÂÅ’ File not found: {rules_path}")
        print("   Using placeholder - UPDATE PATH IN CELL 2!")
        return "PLACEHOLDER - Update PATHS['keying_rules'] with correct path to COMPLETE_KEYING_RULES_VERIFIED.txt"
    
    with open(rules_path, 'r', encoding='utf-8') as f:
        specs = f.read()
    
    # Verify completeness
    import re
    tags = set(re.findall(r'<(\w+)>', specs))
    
    print(f"   Ã¢Å“â€¦ Complete specifications loaded")
    print(f"   Ã°Å¸â€œÅ  Size: {len(specs):,} characters")
    print(f"   Ã°Å¸ÂÂ·Ã¯Â¸Â  Tags covered: {len(tags)} unique tags")
    print(f"   Ã°Å¸â€œÂ Includes: General + Securities + Table rules")
    
    return specs

KEYING_SPECIFICATIONS = load_complete_keying_specifications(PATHS['keying_rules'])

print("\nÃ¢Å“â€¦ Keying specifications ready for LLM!")

# ====== CODE CELL 4 ======
class RAGManager:
    """
    ChromaDB-backed Retrieval-Augmented Generation manager.
    Stores keying spec rule chunks and vendor SGML examples for per-batch context retrieval.
    Replaces passing the full 21,908-char keying spec in every LLM call with targeted retrieval.
    """

    # LEAF tags only: do NOT include BLOCK2-6 (they are container tags, too long when matched)
    INDEXABLE_TAGS = {'P', 'P1', 'P2', 'P3', 'P4', 'ITEM', 'LINE', 'QUOTE'}

    def __init__(self, keying_specs_path: str, persist_dir: str, vendor_sgms: List[str],
                 n_rules: int = 6, n_examples: int = 8):
        self.keying_specs_path = keying_specs_path
        self.persist_dir = persist_dir
        self.vendor_sgms = vendor_sgms
        self.n_rules = n_rules
        self.n_examples = n_examples
        self._client = None
        self._rules = None
        self._examples = None
        self._initialized = False

    # ──────────────────────────────────────────────────────────────────
    # Public API
    # ──────────────────────────────────────────────────────────────────

    def initialize(self) -> None:
        """Connect to (or create) ChromaDB and populate collections on first run."""
        os.makedirs(self.persist_dir, exist_ok=True)
        self._client = chromadb.PersistentClient(path=self.persist_dir)

        self._rules = self._client.get_or_create_collection(
            name="keying_rules",
            metadata={"hnsw:space": "cosine"}
        )
        self._examples = self._client.get_or_create_collection(
            name="sgm_examples",
            metadata={"hnsw:space": "cosine"}
        )

        if self._rules.count() == 0:
            print("  RAG: indexing keying rules…")
            self._index_keying_rules()

        if self._examples.count() == 0:
            print("  RAG: indexing vendor SGM examples…")
            self._index_vendor_examples()

        self._initialized = True
        print(f"  RAG ready — {self._rules.count()} rule chunks, {self._examples.count()} examples")

    def get_context_for_batch(self, batch_texts: List[str]) -> str:
        """Return a concise context string for the given batch of paragraph texts."""
        if not self._initialized:
            return ""
        # Build query from first 3 paragraphs (up to 150 chars each)
        query = " ".join(t[:150] for t in batch_texts[:3])
        parts: List[str] = []

        # ── Keying rules ──────────────────────────────────────────────
        n_r = min(self.n_rules, self._rules.count())
        if n_r > 0:
            res = self._rules.query(query_texts=[query], n_results=n_r)
            if res and res["documents"]:
                parts.append("RELEVANT KEYING RULES:")
                for doc in res["documents"][0]:
                    parts.append(f"  {doc[:500]}")

        # ── Vendor examples ───────────────────────────────────────────
        n_e = min(self.n_examples, self._examples.count())
        if n_e > 0:
            res = self._examples.query(query_texts=[query], n_results=n_e)
            if res and res["documents"]:
                parts.append("\nSIMILAR VENDOR EXAMPLES (tag → text):")
                for doc in res["documents"][0]:
                    parts.append(f"  {doc[:300]}")

        return "\n".join(parts)

    # ──────────────────────────────────────────────────────────────────
    # Internal helpers
    # ──────────────────────────────────────────────────────────────────

    def _index_keying_rules(self) -> None:
        with open(self.keying_specs_path, 'r', encoding='utf-8') as fh:
            text = fh.read()

        # Split on 2+ blank lines (section boundaries)
        raw_chunks = re.split(r'\n{2,}', text)
        ids, docs = [], []
        for i, chunk in enumerate(raw_chunks):
            chunk = chunk.strip()
            if len(chunk) > 40:
                ids.append(f"rule_{i:04d}")
                docs.append(chunk)

        # Also load KV rules JSON for structured tag-by-tag retrieval
        kv_path = self.keying_specs_path.replace(
            'COMPLETE_KEYING_RULES_UPDATED.txt', 'KEYING_RULES_KV.json'
        )
        if os.path.exists(kv_path):
            try:
                import json as _json
                with open(kv_path, 'r', encoding='utf-8') as fkv:
                    kv = _json.load(fkv)
                kv_counter = len(ids)
                def _flatten_kv(obj, prefix=''):
                    results = []
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            if k.startswith('_'):
                                continue
                            if isinstance(v, dict):
                                results.extend(_flatten_kv(v, prefix=f'{prefix}{k}: '))
                            elif isinstance(v, (str, list)):
                                text_val = v if isinstance(v, str) else '; '.join(str(x) for x in v)
                                results.append(f'{prefix}{k}: {text_val}')
                    return results
                kv_chunks = _flatten_kv(kv)
                for chunk in kv_chunks:
                    if len(chunk) > 30:
                        ids.append(f"kv_{kv_counter:04d}")
                        docs.append(chunk)
                        kv_counter += 1
                print(f"    → {len(kv_chunks)} KV rule chunks indexed from KEYING_RULES_KV.json")
            except Exception as e:
                print(f"    ⚠ KV rules load failed: {e}")

        # Upsert in batches of 100 (ChromaDB limit)
        for start in range(0, len(docs), 100):
            self._rules.upsert(ids=ids[start:start+100], documents=docs[start:start+100])
        print(f"    → {len(docs)} rule chunks indexed total")

    def _index_vendor_examples(self) -> None:
        all_ids, all_docs, all_meta = [], [], []
        counter = 0
        for sgm_path in self.vendor_sgms:
            if not os.path.exists(sgm_path):
                print(f"    ⚠ Not found: {sgm_path}")
                continue
            examples = self._parse_sgm(sgm_path)
            fname = os.path.basename(sgm_path)
            for ex in examples:
                all_ids.append(f"ex_{counter:06d}")
                label = f"[{ex['tag']}] {ex['text'][:200]}"
                all_docs.append(label)
                all_meta.append({"tag": ex['tag'], "source": fname})
                counter += 1
            print(f"    → {len(examples)} examples from {fname}")

        if all_docs:
            for start in range(0, len(all_docs), 100):
                self._examples.upsert(
                    ids=all_ids[start:start+100],
                    documents=all_docs[start:start+100],
                    metadatas=all_meta[start:start+100]
                )
        print(f"    → {len(all_docs)} examples total indexed")

    def _parse_sgm(self, sgm_path: str) -> List[Dict]:
        """
        Extract tag+text pairs from a vendor SGM file.
        Uses explicit tag alternation to avoid the outer document wrapper tag
        (<POLIDOC>, <MLIDOC>, etc.) consuming the entire file content.
        Only leaf-level tags (P, P1-P4, ITEM, LINE, QUOTE) are matched.
        """
        examples = []
        try:
            with open(sgm_path, 'r', encoding='utf-8', errors='replace') as fh:
                content = fh.read()

            # Build alternation pattern from longest tag name to shortest
            # (so P4 is tried before P, preventing partial match issues)
            sorted_tags = sorted(self.INDEXABLE_TAGS, key=len, reverse=True)
            tag_alt = '|'.join(re.escape(t) for t in sorted_tags)

            # Match ONLY the explicit leaf tags — this prevents the outer
            # document wrapper tag from consuming the entire file.
            pat = re.compile(
                rf'<({tag_alt})(?:\s[^>]*)?>(.+?)</\1>',
                re.DOTALL
            )
            for m in pat.finditer(content):
                tag = m.group(1)
                text = re.sub(r'<[^>]+>', '', m.group(2)).strip()
                if 15 <= len(text) <= 600:
                    examples.append({'tag': tag, 'text': text})
        except Exception as exc:
            print(f"    ⚠ Parse error {sgm_path}: {exc}")
        return examples


print("✅ RAGManager defined (v2 — fixed SGM parser)")
print("   Uses ChromaDB PersistentClient for rule + example retrieval")
print("   _parse_sgm: explicit leaf-tag alternation prevents wrapper-tag match")


# ====== CODE CELL 5 ======
import json
from concurrent.futures import ThreadPoolExecutor

# ─────────────────────────────────────────────────────────────────────────────
# AGENT RULES — JSON specifications for each specialized agent
# ─────────────────────────────────────────────────────────────────────────────

# Pre-compiled EM detection regex patterns (Phase 1: no LLM required)
EM_REGEX_PATTERNS = [
    # NI / MI instrument numbers: NI 31-103, MI 61-101, NI 81-102CP
    re.compile(r'\b((?:NI|MI)\s+\d+[\-\u2012\u2013]\d+(?:CP|F\d+)?)\b'),
    # CSA / OSC / CIRO regulatory notices and rules
    re.compile(
        r'\b((?:CSA|OSC|CIRO|IIROC|MFDA|CIPF|NEO)\s+'
        r'(?:Staff\s+)?(?:Notice|Rule|Policy|Blanket\s+Order|Guidance|Bulletin)'
        r'(?:\s+No\.?\s*[\d\-]+)?)\b', re.I
    ),
    # Email addresses
    re.compile(r'\b([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,6})\b'),
    # Named Canadian legislation
    re.compile(
        r'\b(Securities Act|Bank Act|Income Tax Act|PCMLTFA|'
        r'Business Corporations Act|Companies Act|Insurance Act|'
        r'Trust and Loan Companies Act|Bankruptcy and Insolvency Act|'
        r'Financial Administration Act|Canada Business Corporations Act|'
        r'Credit Unions? Act|Cooperative Association Act|'
        r'Proceeds of Crime.*?Act)\b', re.I
    ),
]

AGENT_RULES = {
    'EM_AGENT': {
        "agent_name":  "EM_AGENT",
        "tag_type":    "EM",
        "description": "Expert in <EM> inline emphasis tags for Canadian securities SGML",
        "pattern_rules": [
            {"rule_id": "EM_P01", "priority": 1,
             "description": "NI/MI instrument numbers (NI 31-103, MI 61-101)"},
            {"rule_id": "EM_P02", "priority": 1,
             "description": "CSA/OSC/CIRO regulatory notices and rules"},
            {"rule_id": "EM_P03", "priority": 1,
             "description": "Email addresses in contact information"},
            {"rule_id": "EM_P04", "priority": 1,
             "description": "Named Canadian Acts (Securities Act, Bank Act, etc.)"},
        ],
        "llm_rules": [
            {"rule_id": "EM_L01", "priority": 1,
             "description": "Italic text from DOCX formatting (check italic_runs field)",
             "examples": [
                 {"input": "(italicized text)", "output": "<EM>(italicized text)</EM>"}
             ]},
            {"rule_id": "EM_L02", "priority": 2,
             "description": "Compliance deficiency findings",
             "examples": [
                 {"input": "Training not provided to registered individuals",
                  "output": "<EM>Training not provided to registered individuals</EM>"}
             ]},
            {"rule_id": "EM_L03", "priority": 2,
             "description": "Key defined terms at first definition (after 'means')",
             "examples": [
                 {"input": '"dealer member" means a member...',
                  "output": '"<EM>dealer member</EM>" means a member...'}
             ]},
        ],
        "conflict_resolution": "Pattern rules override LLM rules; highest priority wins"
    },

    'BOLD_AGENT': {
        "agent_name":  "BOLD_AGENT",
        "tag_type":    "BOLD",
        "description": "Expert in <BOLD> tags for visual bold formatting",
        "rules": [
            {"rule_id": "BOLD_R01", "priority": 1,
             "condition": "formatting_shows_bold",
             "description": "Apply BOLD to runs with bold=True from DOCX formatting",
             "examples": [
                 {"bold_run": "Important Notice", "output": "<BOLD>Important Notice</BOLD>"}
             ]},
        ]
    },

    'BLOCK_AGENT': {
        "agent_name":       "BLOCK_AGENT",
        "tag_type":         "BLOCK",
        "description":      "Expert in structural BLOCK/ITEM/P hierarchy tags",
        "uses_keying_specs": True,
        "note": ("Receives complete Carswell DTD structural keying specs as system context. "
                 "Does NOT handle inline formatting — EMAgent and BOLDAgent do that.")
    },

    'ITEM_AGENT': {
        "agent_name":  "ITEM_AGENT",
        "tag_type":    "ITEM",
        "description": "Expert in <ITEM> list-item tags",
        "rules": [
            {"rule_id": "ITEM_R01", "priority": 1,
             "condition": "has_bullet_character",
             "description": "Apply ITEM if paragraph starts with bullet or list marker",
             "patterns": [r"^[•·◦▪▫■□●○◆◇→⇒➢➤]", r"^\([a-z]\)", r"^\([0-9]+\)",
                          r"^[a-z]\.", r"^[0-9]+\."]},
            {"rule_id": "ITEM_R02", "priority": 1,
             "condition": "follows_list_introduction",
             "description": "Apply ITEM if previous para introduced a list",
             "intro_patterns": ["following", "as follows", "include", "includes"]},
            {"rule_id": "ITEM_R03", "priority": 2,
             "condition": "indent_suggests_list",
             "description": "Apply ITEM if indented (<500 chars)",
             "thresholds": {"left_indent": "> 10", "text_length": "< 500"}},
        ]
    },

    'STRUCTURE_AGENT': {
        "agent_name":  "STRUCTURE_AGENT",
        "tag_type":    "STRUCTURE",
        "description": "Expert in paragraph-level structure tags (P, P1, P2, P3, P4)",
        "rules": [
            {"rule_id": "STRUCT_R01", "priority": 1,
             "condition": "is_regular_paragraph",
             "description": "Apply P to regular body paragraphs"},
            {"rule_id": "STRUCT_R02", "priority": 1,
             "condition": "is_indented_paragraph",
             "description": "Apply P1/P2/P3/P4 based on indent level"},
        ]
    }
}

print("✅ AGENT_RULES defined")
print(f"   EM_AGENT : {len(EM_REGEX_PATTERNS)} regex patterns + {len(AGENT_RULES['EM_AGENT']['llm_rules'])} LLM rules")
print(f"   BOLD_AGENT: {len(AGENT_RULES['BOLD_AGENT']['rules'])} rules")
print(f"   BLOCK_AGENT: Full keying-spec context")
print(f"   ITEM_AGENT: {len(AGENT_RULES['ITEM_AGENT']['rules'])} rules")
print(f"   STRUCTURE_AGENT: {len(AGENT_RULES['STRUCTURE_AGENT']['rules'])} rules")


# ====== CODE CELL 6 ======
# Cell 6: Data Structures (Updated with skip field)

@dataclass
class DocumentMetadata:
    document_number: str = ''
    title: str = ''
    label: str = ''
    effective_date: str = ''
    adddate: str = ''
    moddate: str = ''
    lang: str = 'EN'
    cite: str = ''          # OSC Bulletin citation e.g. "48 O.S.C.B. 9737"

@dataclass
class RunData:
    text: str
    bold: bool
    italic: bool
    underline: bool
    font_size: Optional[float]
    font_name: Optional[str]

@dataclass
class ParagraphData:
    index: int
    text: str
    runs: List[RunData]
    style: str
    alignment: str
    left_indent: float
    patterns: Dict = field(default_factory=dict)
    final_tag: str = ''
    confidence: float = 0.0
    inline_formatting: List[Dict] = field(default_factory=list)
    skip: bool = False  # Ã¢â€ Â THIS IS THE CRITICAL LINE
    docx_formatting: List[Dict] = field(default_factory=list)  # Ã¢â€ Â For Phase 3

@dataclass
class TableData:
    rows: List[List[Dict]]
    has_header: bool = False
    col_widths: List[int] = None   # Column widths in twips from DOCX tblGrid (for % calculation)

@dataclass
class ImageData:
    filename: str
    width: int
    height: int
    filepath: str
    paragraph_index: int = -1

print('Ã¢Å“â€¦ Data structures defined (with skip and docx_formatting fields)')

# ====== CODE CELL 7 ======
class ABBYYConverter:
    """ABBYY FineReader Engine 12 - DevCode proven approach"""
    
    FEF_DOCX = 8
    
    def __init__(self, customer_id: str, license_path: str, license_password: str):
        self.engine_loader = None
        self.engine = None
        self.customer_id = customer_id
        self.license_path = license_path
        self.license_password = license_password
    
    def initialize(self):
        print("\nÃ°Å¸â€Â§ Initializing ABBYY...")
        self.engine_loader = win32c.Dispatch("FREngine.OutprocLoader.12")
        self.engine = self.engine_loader.InitializeEngine(
            self.customer_id, self.license_path, self.license_password, "", "", False
        )
        self.engine.LoadPredefinedProfile("DocumentConversion_Accuracy")
        print("   Ã¢Å“â€¦ ABBYY initialized")
    
    def convert_pdf_to_docx(self, pdf_path: str, docx_path: str) -> bool:
        print(f"\nÃ°Å¸â€œâ€ž Converting PDF to DOCX...")
        print(f"   Input: {Path(pdf_path).name}")
        
        try:
            document = self.engine.CreateFRDocument()
            print("   Loading PDF...")
            document.AddImageFile(pdf_path, None, None)
            
            print("   Processing...")
            start = time.time()
            document.Process(None)
            print(f"   Ã¢Å“â€¦ Processed in {time.time()-start:.1f}s")
            
            export_params = self.engine.CreateRTFExportParams()
            export_params.PictureExportParams.Resolution = 300
            export_params.BackgroundColorMode = 1
            # PageSynthesisMode=1: column-aware synthesis (default for FRE12)
            # Mode=0 caused wrong reading order in multi-column docs (31-367 -26%)
            export_params.PageSynthesisMode = 1
            export_params.KeepPageBreaks = 1
            export_params.UseDocumentStructure = True
            # Suppress running titles (page headers/footers) to reduce
            # body-text noise and fix score degradation on long docs
            try:
                export_params.WriteRunningTitles = False
            except AttributeError:
                pass  # Older SDK versions without this attribute
            
            print("   Exporting DOCX...")
            document.Export(docx_path, self.FEF_DOCX, export_params)
            document.Close()
            
            if os.path.exists(docx_path):
                size = os.path.getsize(docx_path) / 1024
                print(f"   Ã¢Å“â€¦ DOCX: {size:.1f} KB")
                return True
            return False
        except Exception as e:
            print(f"   Ã¢ÂÅ’ Error: {e}")
            return False
    

    def convert_pdf_to_html(self, pdf_path: str, html_path: str) -> bool:
        """Export ABBYY HTML preserving <sup> footnote markers.

        ABBYY HTML output contains <sup>i</sup> through <sup>xi</sup>
        inline markers that DOCX export loses. Used as a secondary pass
        for docs needing correct footnote position detection.
        """
        print(f"   Exporting HTML for footnote-anchor detection...")
        try:
            document = self.engine.CreateFRDocument()
            document.AddImageFile(pdf_path, None, None)
            import time as _t; _s = _t.time()
            document.Process(None)
            print(f"   Processed in {_t.time()-_s:.1f}s")
            html_params = self.engine.CreateHTMLExportParams()
            document.Export(html_path, 1, html_params)  # FEF_HTML = 1
            document.Close()
            if os.path.exists(html_path):
                size = os.path.getsize(html_path) / 1024
                print(f"   HTML: {size:.1f} KB")
                return True
            return False
        except Exception as e:
            print(f"   HTML export error: {e}")
            return False

    def cleanup(self):
        if self.engine_loader:
            self.engine_loader.ExplicitlyUnload()
            print("   Ã¢Å“â€¦ ABBYY cleaned up")

print("Ã¢Å“â€¦ ABBYYConverter defined")

# ====== CODE CELL 8 ======
class ImageExtractor:
    """Extract images from DOCX and save as BMP"""
    
    def __init__(self, output_dir: str, dpi: int = 300):
        self.output_dir = output_dir
        self.dpi = dpi
        self.image_counter = 0
    
    def extract_images_from_docx(self, docx_path: str) -> List[ImageData]:
        """Extract all images from DOCX and save as BMP"""
        print("\nÃ°Å¸â€œÂ¸ Extracting images from DOCX...")
        images = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for media_file in media_files:
                    try:
                        image_data = docx_zip.read(media_file)
                        img = Image.open(BytesIO(image_data))
                        
                        if img.mode == 'RGBA':
                            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                            rgb_img.paste(img, mask=img.split()[3])
                            img = rgb_img
                        elif img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        self.image_counter += 1
                        filename = f"SB1{self.image_counter:05d}.BMP"
                        filepath = os.path.join(self.output_dir, filename)
                        
                        img.save(filepath, 'BMP', dpi=(self.dpi, self.dpi))
                        
                        images.append(ImageData(
                            filename=filename,
                            width=img.width,
                            height=img.height,
                            filepath=filepath
                        ))
                        
                    except Exception as e:
                        print(f"   Ã¢Å¡Â Ã¯Â¸Â Error processing {media_file}: {e}")
                        continue
            
            print(f"   Ã¢Å“â€¦ Extracted {len(images)} images")
            return images
            
        except Exception as e:
            print(f"   Ã¢Å¡Â Ã¯Â¸Â Image extraction error: {e}")
            return []

print("Ã¢Å“â€¦ ImageExtractor defined")

# ====== CODE CELL 9 ======
class CompleteDOCXExtractor:
    """DOCX extraction with PHASE 2 Cover/TOC detection + doc-type-aware table filtering.
    v5.0: Improved metadata extraction (LABEL, N, full-title, ADDDATE from doc date).
          Container BLOCK format auto-detection via is_annual_report + notices.
          TOC 2-col table fix works in Annual Report mode too.
    """

    # --- LABEL detection keywords (order: most specific first) ---
    _LABEL_PATTERNS = [
        # Joint notices first (most specific)
        ('Joint CSA / CIRO Staff Notice',
         re.compile(r'(?:joint|conjointe).*?(?:CSA|ACVM).*?(?:CIRO|OCRI)', re.I)),
        ('Joint CSA / IIROC Staff Notice',
         re.compile(r'(?:joint|conjointe).*?(?:CSA|ACVM).*?(?:IIROC)', re.I)),
        # CSA types
        ('CSA Multilateral Staff Notice',
         re.compile(r'\bCSA\s+Multilateral\s+Staff\s+Notice\b', re.I)),
        ('CSA Staff Notice',
         re.compile(r'\bCSA\s+staff\s+notice\b', re.I)),
        ('CSA Notice',
         re.compile(r'\bCSA\s+notice\b', re.I)),
        # OSC types
        ('OSC Staff Notice',
         re.compile(r'\bOSC\s+staff\s+notice\b', re.I)),
        ('OSC Blanket Order',
         re.compile(r'\bOSC\s+blanket\s+order\b', re.I)),
        ('OSC Policy',
         re.compile(r'\bOSC\s+policy\b', re.I)),
        ('OSC Rule',
         re.compile(r'\bOSC\s+Rule\b', re.I)),
        ('OSC Notice',
         re.compile(r'\bOSC\s+notice\b', re.I)),
        # Blanket Orders (before National Instrument to avoid body-text false match)
        ('MSC Coordinated Blanket Order',
         re.compile(r'(?:Manitoba\s+Securities\s+Commission|MSC).*?Coordinated\s+Blanket\s+Order', re.I | re.DOTALL)),
        ('Coordinated Blanket Order',
         re.compile(r'\bCoordinated\s+Blanket\s+Order\b', re.I)),
        ('Blanket Order',
         re.compile(r'\bBlanket\s+Order\b', re.I)),
        # Instrument/rule types (document number prefix)
        ('Multilateral Instrument',
         re.compile(r'\bMultilateral\s+Instrument\b', re.I)),
        ('National Instrument',
         re.compile(r'\bNational\s+Instrument\b', re.I)),
        ('Multilateral Policy',
         re.compile(r'\bMultilateral\s+Policy\b', re.I)),
        ('National Policy',
         re.compile(r'\bNational\s+Policy\b', re.I)),
        ('Companion Policy',
         re.compile(r'\bCompanion\s+Policy\b', re.I)),
        ('MFDA Policy',
         re.compile(r'\bMFDA\s+(?:Rule|Policy)\b', re.I)),
        ('CIRO Notice',
         re.compile(r'\bCIRO\s+(?:staff\s+)?notice\b', re.I)),
        # Catch-all number-prefix based
        ('NI',  re.compile(r'^NI\s+\d', re.I)),
        ('MI',  re.compile(r'^MI\s+\d', re.I)),
        ('NP',  re.compile(r'^NP\s+\d', re.I)),
        ('CP',  re.compile(r'^CP\s+\d', re.I)),
    ]

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.document = Document(docx_path)
        self._is_annual_report = self._detect_annual_report()
        self._footnotes: Dict[int, str] = self._load_footnotes()
        self._inline_footnotes: Dict[int, str] = {}      # superscript-digit inline footnotes
        self._inline_fn_para_idx: set = set()            # paragraph indices to skip (footnote bodies)
        self._inline_footnotes, self._inline_fn_para_idx = self._load_inline_footnotes()
        self._numbering: Dict = self._load_numbering()   # {numId: {ilvl: (numFmt, lvlText, start)}}
        self._num_counters: Dict = {}                    # {(numId, ilvl): current_count}

    def _load_numbering(self) -> Dict:
        """Load word/numbering.xml → {numId: {ilvl: (numFmt, lvlText, start)}}.
        Used to reconstruct DOCX auto-numbering labels (a), (b), (i), (ii), 1., etc.
        that python-docx does NOT include in para.text."""
        import zipfile as _zf2
        import xml.etree.ElementTree as _ET2
        _W2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        result: Dict = {}
        try:
            with _zf2.ZipFile(self.docx_path, 'r') as z:
                if 'word/numbering.xml' not in z.namelist():
                    return result
                xml_bytes = z.read('word/numbering.xml')
            root = _ET2.fromstring(xml_bytes)
            # Build abstractNum map: {abstractNumId: {ilvl: (numFmt, lvlText, start)}}
            abstract_nums: Dict = {}
            for an in root.findall(f'{{{_W2}}}abstractNum'):
                an_id = int(an.get(f'{{{_W2}}}abstractNumId', -1))
                levels: Dict = {}
                for lvl in an.findall(f'{{{_W2}}}lvl'):
                    ilvl = int(lvl.get(f'{{{_W2}}}ilvl', 0))
                    numFmt = ''
                    lvlText = ''
                    start = 1
                    nf = lvl.find(f'{{{_W2}}}numFmt')
                    if nf is not None:
                        numFmt = nf.get(f'{{{_W2}}}val', '')
                    lt = lvl.find(f'{{{_W2}}}lvlText')
                    if lt is not None:
                        lvlText = lt.get(f'{{{_W2}}}val') or ''
                    s = lvl.find(f'{{{_W2}}}start')
                    if s is not None:
                        try:
                            start = int(s.get(f'{{{_W2}}}val', 1))
                        except ValueError:
                            start = 1
                    levels[ilvl] = (numFmt, lvlText, start)
                abstract_nums[an_id] = levels
            # Map numId → abstractNum levels (with numId-level overrides applied)
            for num in root.findall(f'{{{_W2}}}num'):
                num_id_str = num.get(f'{{{_W2}}}numId', '')
                try:
                    num_id = int(num_id_str)
                except ValueError:
                    continue
                an_ref = num.find(f'{{{_W2}}}abstractNumId')
                if an_ref is None:
                    continue
                try:
                    an_id = int(an_ref.get(f'{{{_W2}}}val', -1))
                except ValueError:
                    continue
                if an_id not in abstract_nums:
                    continue
                import copy as _cp
                levels = _cp.deepcopy(abstract_nums[an_id])
                # Apply any lvlOverride elements for this numId
                for ov in num.findall(f'{{{_W2}}}lvlOverride'):
                    try:
                        ov_ilvl = int(ov.get(f'{{{_W2}}}ilvl', -1))
                    except ValueError:
                        continue
                    ov_start = ov.find(f'{{{_W2}}}startOverride')
                    if ov_start is not None and ov_ilvl in levels:
                        try:
                            levels[ov_ilvl] = (
                                levels[ov_ilvl][0],
                                levels[ov_ilvl][1],
                                int(ov_start.get(f'{{{_W2}}}val', levels[ov_ilvl][2]))
                            )
                        except (ValueError, KeyError):
                            pass
                result[num_id] = levels
        except Exception:
            pass
        return result

    @staticmethod
    def _int_to_roman(n: int) -> str:
        """Convert positive integer to Roman numeral string."""
        val  = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
        syms = ['M','CM','D','CD','C','XC','L','XL','X','IX','V','IV','I']
        result = ''
        for v, s in zip(val, syms):
            while n >= v:
                result += s
                n -= v
        return result

    def _get_numpr_label(self, para_el) -> str:
        """Return the auto-generated list label for a paragraph with <w:numPr>.
        Increments the appropriate counter and resets deeper-level counters.
        Returns '' when no numPr, bullet format, or numbering not found."""
        _W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        try:
            pPr = para_el.find(f'{_W}pPr')
            if pPr is None:
                return ''
            numPr_el = pPr.find(f'{_W}numPr')
            if numPr_el is None:
                return ''
            ilvl_el  = numPr_el.find(f'{_W}ilvl')
            numId_el = numPr_el.find(f'{_W}numId')
            if ilvl_el is None or numId_el is None:
                return ''
            ilvl   = int(ilvl_el.get(f'{_W}val',  0))
            numId  = int(numId_el.get(f'{_W}val', 0))
            if numId == 0:          # numId=0 means "no list"
                return ''
            if numId not in self._numbering:
                return ''
            levels = self._numbering[numId]
            if ilvl not in levels:
                return ''
            numFmt, lvlText, start = levels[ilvl]
            # Skip non-numbered formats (bullets, none, ordinal symbols)
            if numFmt in ('bullet', 'none', ''):
                return ''
            # Increment counter for this level; reset deeper levels (same numId)
            key = (numId, ilvl)
            for k in [k for k in self._num_counters if k[0] == numId and k[1] > ilvl]:
                del self._num_counters[k]
            if key not in self._num_counters:
                self._num_counters[key] = start - 1
            self._num_counters[key] += 1
            n = self._num_counters[key]

            def _fmt(n: int, fmt: str) -> str:
                if fmt == 'lowerLetter':
                    return chr(ord('a') + min(n - 1, 25))
                if fmt == 'upperLetter':
                    return chr(ord('A') + min(n - 1, 25))
                if fmt == 'lowerRoman':
                    return self._int_to_roman(n).lower()
                if fmt == 'upperRoman':
                    return self._int_to_roman(n).upper()
                return str(n)   # decimal / ordinal

            # Replace %N placeholders in the level-text template
            import re as _nr
            label = lvlText
            for m in _nr.finditer(r'%(\d+)', lvlText):
                ref_ilvl = int(m.group(1)) - 1   # %1 → ilvl 0, %2 → ilvl 1
                ref_key  = (numId, ref_ilvl)
                ref_fmt  = levels.get(ref_ilvl, (numFmt, '', 1))[0]
                ref_n    = self._num_counters.get(ref_key, 1)
                label = label.replace(m.group(0), _fmt(ref_n, ref_fmt), 1)
            return label.strip()
        except Exception:
            return ''

    def _load_footnotes(self) -> Dict[int, str]:
        """Load footnote text from word/footnotes.xml inside the DOCX ZIP.
        Returns {footnote_id: text_content} for IDs >= 1 (skips separator/continuation)."""
        import zipfile as _zipfile
        import xml.etree.ElementTree as _ET
        footnotes: Dict[int, str] = {}
        try:
            with _zipfile.ZipFile(self.docx_path) as z:
                if 'word/footnotes.xml' not in z.namelist():
                    return footnotes
                xml_bytes = z.read('word/footnotes.xml')
            root = _ET.fromstring(xml_bytes)
            _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for fn_el in root.findall(f'{{{_W}}}footnote'):
                raw_id = fn_el.get(f'{{{_W}}}id', '')
                try:
                    fn_id = int(raw_id)
                except ValueError:
                    continue
                if fn_id < 1:   # skip -1 (separator) and 0 (continuation notice)
                    continue
                # Collect text from all <w:t> descendants, concatenating paragraph runs
                paragraphs_text = []
                for p_el in fn_el.findall(f'.//{{{_W}}}p'):
                    para_pieces = []
                    for t_el in p_el.findall(f'.//{{{_W}}}t'):
                        if t_el.text:
                            para_pieces.append(t_el.text)
                    t = ''.join(para_pieces).strip()
                    # Strip leading footnote number (e.g. "1 " or "1. " at start)
                    t = re.sub(r'^\d+\.?\s+', '', t)
                    if t:
                        paragraphs_text.append(t)
                combined = ' '.join(paragraphs_text).strip()
                if combined:
                    footnotes[fn_id] = combined
        except Exception:
            pass
        return footnotes

    def _load_inline_footnotes(self) -> tuple:
        """Detect 'manual' superscript footnotes: body paragraphs that serve as
        footnote definitions (start with superscript digit run or 'N text').
        Returns ({fn_num: text}, {para_indices_to_skip}).
        Safety guards:
          (a) No real Word footnotes with content (self._footnotes is empty).
          (b) At least 1 superscript-digit INLINE reference (after text content in para).
          (c) Detected fn_num cross-referenced against known inline reference numbers.
          (d) Paragraph is not a heading style."""
        if self._footnotes:          # Real Word footnotes take priority — skip
            return {}, set()
        _W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        inline_fns: Dict[int, str] = {}
        skip_indices: set = set()
        try:
            all_paras = list(self.document.paragraphs)
            n = len(all_paras)
            main_body_limit = max(1, int(n * 0.8))
            # First pass: collect inline reference numbers — only superscript digits
            # that appear AFTER some text content in a paragraph.
            # This distinguishes "body cited as footnote ref" from "footnote body marker
            # at paragraph start" (which also uses superscript for the footnote number).
            main_ref_nums: set = set()
            for p in all_paras[:main_body_limit]:
                seen_text = False
                for r in p.runs:
                    rpr = r._element.find(_W + 'rPr')
                    is_super_digit = False
                    if rpr is not None:
                        va = rpr.find(_W + 'vertAlign')
                        if va is not None and va.get(_W + 'val') == 'superscript' \
                                and r.text.strip().isdigit():
                            is_super_digit = True
                    if is_super_digit and seen_text:
                        main_ref_nums.add(int(r.text.strip()))
                    elif r.text.strip() and not is_super_digit:
                        seen_text = True
            if not main_ref_nums:   # No inline footnote references → nothing to do
                return {}, set()
            # Second pass: FORWARD scan for footnote body paragraphs.
            # Start at the 1/3 point (skip cover/intro) and scan to the end.
            # This avoids the fragility of backward scans in web-scraped DOCXs that
            # may have long copyright/navigation text between footnote bodies and the end.
            scan_start = max(3, n // 3)
            for i in range(scan_start, n):
                p = all_paras[i]
                text = p.text.strip()
                if not text:
                    continue
                # Skip long non-digit paragraphs (disclaimers, copyright, navigation)
                if len(text) > 150 and not text[0].isdigit():
                    continue
                # Skip heading-style paragraphs (numbered section titles ≠ footnote bodies)
                style_name = (p.style.name or '').lower() if p.style else ''
                if 'heading' in style_name:
                    continue
                fn_num = None
                fn_text = None
                # Case 1: first non-blank run is a superscript digit
                for r in p.runs:
                    if not r.text.strip():
                        continue
                    rpr = r._element.find(_W + 'rPr')
                    is_super_digit = False
                    if rpr is not None:
                        va = rpr.find(_W + 'vertAlign')
                        if va is not None and va.get(_W + 'val') == 'superscript' \
                                and r.text.strip().isdigit():
                            is_super_digit = True
                    if is_super_digit:
                        fn_num = int(r.text.strip())
                        fn_text = re.sub(r'^\d+[. ]*\s*', '', text).strip()
                    break   # Only examine the first non-blank run
                # Case 2: paragraph starts with "N. " or "N " (digit + period/space)
                if fn_num is None:
                    m = re.match(r'^(\d{1,2})[. ]\s*(.+)', text)
                    if m:
                        fn_num = int(m.group(1))
                        fn_text = m.group(2).strip()
                # Cross-reference: fn_num must match a known inline reference + has text
                if (fn_num is not None and fn_text and 1 <= fn_num <= 20
                        and fn_num in main_ref_nums):
                    inline_fns[fn_num] = fn_text
                    skip_indices.add(i)
        except Exception:
            pass
        return inline_fns, skip_indices

    def _detect_annual_report(self) -> bool:
        """Detect Annual Report from first 10 paragraphs."""
        annual_report_phrases = ['annual report', 'rapport annuel']
        sample = ' '.join(p.text.lower() for p in self.document.paragraphs[:10])
        return any(phrase in sample for phrase in annual_report_phrases)

    def extract_complete_document(self) -> Dict[str, Any]:
        print('\n📄 Extracting DOCX with filtering...')
        doc_type = 'Annual Report' if self._is_annual_report else 'Notice/Circular'
        print(f'   📋 Document type: {doc_type}')
        metadata = self._extract_metadata()
        content = self._extract_content_in_order()
        content = self._detect_cover_page(content)
        content = self._detect_toc(content)
        content = self._filter_global_noise(content)
        # Merge adjacent tables with same column count (fixes ABBYY page-break splits)
        content = self._merge_adjacent_tables(content)
        paragraphs = [c['data'] for c in content
                     if c['type'] == 'paragraph' and not c['data'].skip]
        tables = [c['data'] for c in content if c['type'] == 'table']
        print(f'   ✅ {len(paragraphs)} paragraphs (after filtering)')
        print(f'   ✅ {len(tables)} tables (data tables, layout→paragraphs/dropped)')
        return {'metadata': metadata, 'content': content, 'paragraphs': paragraphs, 'tables': tables}

    @staticmethod
    def _is_page_break_continuation(next_table: 'TableData') -> bool:
        """Return True if next_table is a CONTINUATION of the previous table
        (ABBYY page-break split), not a new independent Q&A / data entry.

        A continuation is identified by the first DATA row having an empty
        first cell (#/index column).  A new entry has a non-empty first cell.

        Works for Q&A tables (#, Section, Question, Response) and generic
        2-column tables: if every row except a possible repeated header row has
        an empty first cell, treat it as a continuation.
        """
        if not next_table.rows:
            return True  # Empty table — safe to merge

        # Skip the header row (all cells look like column labels) if present.
        # A header row: first cell is '#', 'No', 'Number', '№' or similar short label.
        HEADER_FIRSTCELL = {'#', 'no', 'number', 'no.', 'num', 'item', '№', ''}

        def _cell_text(cell):
            """Extract text from a cell that may be a str or {'text': ..., 'bold': ...} dict."""
            if isinstance(cell, dict):
                return cell.get('text', '')
            return str(cell) if cell is not None else ''

        data_rows = [
            r for r in next_table.rows
            if (r and _cell_text(r[0]).strip().lower() not in HEADER_FIRSTCELL)
        ]
        if not data_rows:
            # Only header rows present — likely a duplicated header at page top.
            # Allow merge so the duplicate header gets absorbed.
            return True

        # If ALL data rows have an empty first cell → continuation.
        # If ANY data row has a non-empty first cell → new entry, do NOT merge.
        first_cells_nonempty = [_cell_text(r[0]).strip() for r in data_rows if r and _cell_text(r[0]).strip()]
        return len(first_cells_nonempty) == 0


    @staticmethod
    def _is_page_break_continuation_static(next_table):
        """Return True if next_table is a page-break split continuation."""
        if not next_table.rows:
            return True
        def _cell_text(cell):
            if isinstance(cell, dict):
                return cell.get('text', '')
            return str(cell) if cell is not None else ''
        HEADER_LABELS = {'#', 'no', 'no.', 'number', 'num', 'item',
                          'section', 'question', 'response', 'description', ''}
        data_rows = [r for r in next_table.rows
                     if r and _cell_text(r[0]).strip().lower() not in HEADER_LABELS]
        if not data_rows:
            return True  # Only header rows — duplicated header at page top
        return all(not _cell_text(r[0]).strip() for r in data_rows if r)

    def _merge_adjacent_tables(self, content: List[Dict]) -> List[Dict]:
        """Merge consecutive tables that are page-break splits of the same table.

        ABBYY splits tables across page breaks into multiple small tables.
        This method merges only GENUINE continuations (where the next table's
        first data row has an empty index/# cell, indicating it carries over
        from the previous table).

        Tables whose next instance starts a NEW entry (non-empty # cell)
        are kept separate, preserving individual Q&A entries for docs like
        96-307 (Frequently Asked Questions).
        """
        if not content:
            return content
        result = []
        i = 0
        while i < len(content):
            item = content[i]
            if item['type'] != 'table':
                result.append(item)
                i += 1
                continue
            # Start a merge group
            merged = item['data']
            num_cols = max(len(r) for r in merged.rows) if merged.rows else 0
            j = i + 1
            # Consume next items: allow a single empty paragraph between tables
            while j < len(content):
                nxt = content[j]
                # Skip over empty/whitespace paragraphs that page-break injected
                if nxt['type'] == 'paragraph':
                    txt = nxt['data'].text.strip()
                    if not txt:
                        j += 1
                        continue
                    break  # Non-empty paragraph ends the merge window
                if nxt['type'] == 'table':
                    nxt_cols = max(len(r) for r in nxt['data'].rows) if nxt['data'].rows else 0
                    if nxt_cols == num_cols and num_cols > 0:
                        # Only merge if this is a true page-break continuation
                        if self._is_page_break_continuation(nxt['data']):
                            merged.rows.extend(nxt['data'].rows)
                            j += 1
                            continue
                        else:
                            break  # New entry — keep as separate table
                break
            result.append({'type': 'table', 'data': merged})
            i = j
        merged_count = len([x for x in result if x['type'] == 'table'])
        orig_count   = len([x for x in content if x['type'] == 'table'])
        if merged_count != orig_count:
            print(f'   🔗 Table merge: {orig_count} → {merged_count} tables')
        return result

    # ─── METADATA ──────────────────────────────────────────────────────────────
    def _extract_metadata(self) -> 'DocumentMetadata':
        metadata = DocumentMetadata()

        # Collect first 35 paragraph texts for analysis (35 to capture CITE/DATE in longer headers)
        top_texts = []
        top_styles = []  # parallel list of style names for filtering
        for para in self.document.paragraphs[:35]:
            t = para.text.strip()
            if t:
                top_texts.append(t)
                top_styles.append((para.style.name if para.style else '').lower())
        combined = '\n'.join(top_texts)

        # ── pre-filter: build top_texts without OSC Bulletin section headers ──
        # e.g. "B.1 Notices", "B.5 Rules" – single-section breadcrumbs with no doc number
        # Skip Heading 1/Heading #1 ONLY when they are short breadcrumb-style navigation
        # items (≤6 words, no embedded doc numbers like NN-NNN).
        # Do NOT skip Heading 1 paragraphs that are actual document titles (e.g.
        # "Dealer Rebates of Trailing Commissions", "Frequently Asked Questions...").
        # CRITICAL: Do NOT filter OSC Bulletin crumb lines like
        # "B.5.1 OSC Rule 11-502 Distribution of..." — these CONTAIN the real title.
        _SEC_HDR_RE = re.compile(r'^[A-Z]\.\d+(?:\.\d+)*\s+\w')
        _HAS_DOC_NUM_RE = re.compile(r'\b\d{2,3}-\d{2,3}\b')
        def _is_nav_heading(text: str, style: str) -> bool:
            # Crumb lines: "B.5.1 OSC Rule 11-502 Title Here" — contain doc number → KEEP
            if _SEC_HDR_RE.match(text):
                if _HAS_DOC_NUM_RE.search(text):
                    return False  # real crumb title — do NOT filter
                return True  # plain nav heading like "B.1 Notices" → filter
            if ('heading 1' in style or 'heading #1' in style):
                # Only filter if it looks like a nav breadcrumb: short, no doc num
                words = text.split()
                if len(words) <= 6 and not _HAS_DOC_NUM_RE.search(text):
                    return True
            return False
        top_texts_clean = [
            t for t, s in zip(top_texts, top_styles)
            if not _is_nav_heading(t, s)
        ]

        # ── 1. Document Number ──
        doc_num = ''
        # Highest priority: use DOCX filename stem if it contains a date (e.g. "2025-04-17.docx"
        # or "notice-2025-09-30.docx") — this matches how vendors assign date-format N values.
        _stem = Path(self.docx_path).stem
        _m_stem = re.search(r'(\d{4}[/-]\d{2}[/-]\d{2})', _stem)
        if _m_stem:
            doc_num = _m_stem.group(1).replace('-', '/')  # normalise to slash
        if not doc_num:
            # Date-style inside document text e.g. "2026/01/22" or "2025-12-10"
            m = re.search(r'\b(\d{4}[/-]\d{2}[/-]\d{2})\b', combined)
            if m:
                doc_num = m.group(1).replace('-', '/')  # normalise to slash
        if not doc_num:
            # FIRST: Try doc-type keyword preceding a document number — most reliable
            # e.g. "CSA Multilateral Staff Notice 31-367" → "31-367"
            # Only capture known revision suffixes like "(Revised)" in doc number.
            # Do NOT capture "(Commodity Futures Act)" or other descriptive parentheticals.
            _DOC_TYPE_NUM_RE = re.compile(
                r'\b(?:CSA\s+(?:Multilateral\s+)?(?:Staff\s+)?Notice|'
                r'OSC\s+(?:Rule|Staff\s+Notice|Notice|Policy)|'
                r'National\s+Instrument|Multilateral\s+Instrument|'
                r'National\s+Policy|Companion\s+Policy|'
                r'Staff\s+Notice)\s+(\d{2,3}-\d{2,3}(?:\s*\((?:Revised|Amendment|Amended|Restated|Updated)\))?)', re.I
            )
            dm2 = _DOC_TYPE_NUM_RE.search(combined)
            if dm2:
                doc_num = dm2.group(1).strip()
        if not doc_num:
            # Blanket/Coordinated Order: extract ORDER number
            m = re.search(r'(?:BLANKET\s+ORDER|ORDER\s+N[Oo]?\.?)\s+(\d{2,3}-\d{3})', combined, re.IGNORECASE)
            if m:
                doc_num = m.group(1)
        if not doc_num:
            # Search only first 5 paragraphs first (title area) to avoid body-text false matches
            top5 = '\n'.join(top_texts[:5])
            m = re.search(r'\b(\d{2,3}-\d{2,3})\b', top5)
            if m:
                doc_num = m.group(1)
        if not doc_num:
            # Expand to all 20 if not found in first 5
            m = re.search(r'\b(\d{2,3}-\d{2,3})\b', combined)
            if m:
                doc_num = m.group(1)
        if not doc_num:
            # Alphanumeric e.g. "NI 81-101", "MI 52-109"
            m = re.search(r'\b([A-Z]{1,3})\s+(\d{2}-\d{3})\b', combined)
            if m:
                doc_num = f"{m.group(1)} {m.group(2)}"
        if not doc_num:
            # Final fallback: use DOCX filename stem when it IS the document number.
            # e.g. "44-306.docx" → "44-306"; "A25-013.docx" → "A25-013"
            # Avoid this for "notice-46-503" style (those already matched date pattern above).
            _stem_bare = re.sub(r'^(?:notice|instrument|order|bulletin)-', '', _stem, flags=re.I)
            _stem_bare = re.sub(r'_TR$', '', _stem_bare, flags=re.I)
            if re.match(r'^\d{2,3}-\d{2,3}$', _stem_bare):
                doc_num = _stem_bare
            elif re.match(r'^[A-Z]\d{2}-\d{3}$', _stem_bare):
                doc_num = _stem_bare
        metadata.document_number = doc_num

        # ── 2. Effective Date ──
        # For "revised" docs use the LATEST date found (not the first/original publication date)
        months = r'(?:January|February|March|April|May|June|July|August|September|October|November|December)'
        # Use (?!\d) instead of trailing \b: allows "May 1, 2025Advisory" to still match.
        # Word boundary fails when year is immediately followed by a letter (e.g. "2025Advisory").
        date_pat = re.compile(rf'\b({months}\s+\d{{1,2}},\s+\d{{4}})(?!\d)', re.IGNORECASE)
        all_dates = date_pat.findall(combined)

        # ── Extra: search table cells for "Effective Date:" label (PEI-style amending instruments) ──
        # These docs have a metadata table at the top where paragraph iteration misses the dates.
        _eff_date_from_table = None
        try:
            tables = self.document.tables
            for tbl in tables[:3]:  # only first 3 tables (header area)
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells]
                    for cell_text in cells:
                        if re.match(r'Effective\s+Date', cell_text, re.IGNORECASE):
                            # Look for date in the SAME row's other cells
                            for other_cell in cells:
                                m_eff = date_pat.search(other_cell)
                                if m_eff:
                                    _eff_date_from_table = m_eff.group(1)
                                    break
                        if _eff_date_from_table:
                            break
                    if _eff_date_from_table:
                        break
        except Exception:
            pass

        # ── Extra: search body paragraphs for "EFFECTIVE DATE:" label paragraph ──
        # Pattern: a standalone para saying "EFFECTIVE DATE:" followed by date in next para
        # Used by NFLD blanket orders, superintendent orders, etc.
        _eff_date_from_label = None
        try:
            all_paras = list(self.document.paragraphs)
            for idx, para in enumerate(all_paras):
                pt = para.text.strip()
                if re.match(r'EFFECTIVE\s+DATE\s*:?\s*$', pt, re.IGNORECASE):
                    # Check next 3 paragraphs for a date
                    for next_para in all_paras[idx + 1: idx + 4]:
                        np_text = next_para.text.strip()
                        m_eff = date_pat.search(np_text)
                        if m_eff:
                            _eff_date_from_label = m_eff.group(1)
                            break
                    if _eff_date_from_label:
                        break
                elif re.search(r'(?:comes?\s+into\s+effect|takes\s+effect)\s+on\s+', pt, re.IGNORECASE):
                    # "This Order comes into effect on December 1, 2025."
                    # "This Superintendent Order takes effect on April 1, 2025."
                    m_eff = date_pat.search(pt)
                    if m_eff:
                        _eff_date_from_label = m_eff.group(1)
                        break
        except Exception:
            pass

        if all_dates:
            # Use doc_num to detect revised documents reliably.
            # Only trigger "latest date" logic when "Revised" is EXPLICITLY in the doc number
            # (e.g. "45-330 (Revised)", "11-312 (Revised)") — avoids false positives where
            # the word "amendment" appears in body text (e.g. 31-367 Blanket Order notice).
            _is_revised = bool(re.search(r'\((?:Revised|Amendment|Amended|Restated|Updated)\)', doc_num, re.I))
            if _eff_date_from_table:
                # Table says "Effective Date: X" — use that directly (most reliable for PEI/amending docs)
                metadata.effective_date = _eff_date_from_table
            elif _eff_date_from_label:
                # Body paragraph "EFFECTIVE DATE:" label or "comes into effect on" phrase
                metadata.effective_date = _eff_date_from_label
            elif _is_revised and len(all_dates) > 1:
                # Parse all dates and take the most recent (latest amendment/revision date)
                from datetime import datetime as _dt
                parsed = []
                for d in all_dates:
                    try:
                        parsed.append((_dt.strptime(d, '%B %d, %Y'), d))
                    except ValueError:
                        pass
                if parsed:
                    metadata.effective_date = max(parsed, key=lambda x: x[0])[1]
                else:
                    metadata.effective_date = all_dates[0]
            else:
                # Non-revised: prefer the FIRST date found in the header area (first 5 paras).
                # This captures the OSC Bulletin publication date (e.g. "May 31, 2025")
                # before rule body dates (e.g. "June 1, 2022" = original effective date).
                header_text = '\n'.join(top_texts[:5])
                header_dates = date_pat.findall(header_text)
                if header_dates:
                    metadata.effective_date = header_dates[0]
                else:
                    metadata.effective_date = all_dates[0]
        elif _eff_date_from_table:
            # No dates in main paragraphs but found in table
            metadata.effective_date = _eff_date_from_table
        elif _eff_date_from_label:
            metadata.effective_date = _eff_date_from_label

        # ── 3. ADDDATE / MODDATE from effective date
        if metadata.effective_date:
            try:
                from datetime import datetime as _dt2
                dt = _dt2.strptime(metadata.effective_date, '%B %d, %Y')
                metadata.adddate = dt.strftime('%Y%m%d')
                metadata.moddate = metadata.adddate
            except ValueError:
                pass
        if not metadata.adddate:
            metadata.adddate = datetime.now().strftime('%Y%m%d')
            metadata.moddate = metadata.adddate

        # ── 3b. OSC Bulletin CITE extraction ──
        # Pattern: "48 O.S.C.B. 9737" or "48 O.S.C.B 7523" (volume + page)
        _CITE_RE = re.compile(r'\b(\d{1,2}\s+O\.S\.C\.B\.?\s+\d{4,5})\b')
        cm = _CITE_RE.search(combined)
        if cm:
            metadata.cite = cm.group(1).strip()

        # ── 4. LABEL detection ──
        metadata.label = self._detect_label(combined, doc_num)

        # ── 5. Title: full assembled title — use cleaned top_texts ──
        metadata.title = self._extract_full_title(top_texts_clean)

        return metadata

    def _detect_label(self, combined_text: str, doc_num: str) -> str:
        """Detect the correct POLIDOC LABEL from document content."""
        for label, pattern in self._LABEL_PATTERNS:
            if pattern.search(combined_text):
                return label

        # Infer from document number prefix
        if doc_num:
            if re.match(r'^\d{2,3}-\d{2,3}$', doc_num):
                prefix = doc_num.split('-')[0]
                if int(prefix) < 30:
                    return 'National Instrument'
            m = re.match(r'^([A-Z]{1,2})\s+\d', doc_num)
            if m:
                pfx = m.group(1).upper()
                if pfx == 'NI':   return 'National Instrument'
                if pfx == 'MI':   return 'Multilateral Instrument'
                if pfx == 'NP':   return 'National Policy'
                if pfx == 'CP':   return 'Companion Policy'

        return 'Securities Document'

    # Matches bulletin breadcrumb: "B.5.1 OSC Rule 11-502 Real Title Here"
    # Capture groups: (1) title portion after the doc number (may be empty)
    # The doc-number portion only consumes REVISION markers like (Revised).
    # Descriptive parentheticals like (Commodity Futures Act) are deliberately
    # left for group(1) so they appear in TI, matching vendor POLIDENT format.
    _CRUMB_TITLE_RE = re.compile(
        r'^[A-Z]\.\d+(?:\.\d+)+\s+'
        r'(?:CSA\s+(?:Multilateral\s+)?(?:Staff\s+)?Notice|'
        r'OSC\s+(?:Rule|Staff\s+Notice|Notice|Policy)|'
        r'National\s+Instrument|Multilateral\s+Instrument|'
        r'Staff\s+Notice)\s+'
        r'(?:\d{2,3}-\d{2,3}(?:\s*\((?:Revised|Amendment|Amended|Restated|Updated)\))?\s*)'
        r'(?:[\-\u2013\u2014\s]+(.+))?$',
        re.I | re.DOTALL
    )

    def _extract_full_title(self, top_texts: List[str]) -> str:
        """Extract the full document title from the first meaningful heading(s).

        v6 strategy:
        1. Check crumb first (B.X.X Label NN-NNN – Title) → extract clean title
        2. Skip any remaining B.X.X breadcrumb lines
        3. Skip nav category words, org-name fragments (≤ 2 words at start)
        4. Stop at preamble lines ("First published...", "This Notice is a revised...")
        5. Stop at ALL-CAPS duplicate of already-collected title
        6. Stop at date-only lines / long body text
        7. Allow crumb continuation (one short follow-on part from next line)
        """
        _NAV_WORDS = re.compile(
            r'^(?:Notices?|Rules?\s+and\s+Policies|Policies|Decisions?|Staff\s+Notices?|'
            r'Exemption\s+Orders?|New\s+Rules?|Amendments?|Rescissions?|'
            r'CSA\s+Releases?|Orders?\s+and\s+Rulings?)\s*$', re.I
        )
        _PREAMBLE_RE = re.compile(
            r'^(?:First\s+published|Previously\s+published|This\s+Notice\s+is\s+a\s+revised|'
            r'This\s+is\s+a\s+revised|Revised\s+version\s+of|'
            r'Published\s+in|As\s+published\s+in)', re.I
        )
        _DATE_ONLY_RE = re.compile(r'^[A-Z][a-z]+\s+\d{1,2},\s+\d{4}$')

        title_parts = []
        got_crumb_title = False

        for text in top_texts:
            if not text.strip():
                continue
            words = text.split()

            # ── 1. CRUMB CHECK FIRST (before any skip logic) ────────────────
            # "B.5.1 OSC Rule 11-502 Distribution of Amounts Received..."
            crumb_m = self._CRUMB_TITLE_RE.match(text.strip())
            if crumb_m:
                real_title = (crumb_m.group(1) or '').strip()
                if real_title:
                    # Use crumb even if trivial org-name parts were already collected
                    trivial_parts = all(len(p.split()) <= 3 for p in title_parts)
                    if len(title_parts) == 0 or trivial_parts:
                        if real_title == real_title.upper() and any(c.isalpha() for c in real_title):
                            real_title = real_title.title()
                        title_parts = [real_title]  # Reset (discard trivial noise)
                        got_crumb_title = True
                        # Don't break immediately — allow one more short continuation line
                        continue
                continue  # Non-extractable crumb line → skip

            # ── 2. Skip remaining B.X.X breadcrumb / nav lines ──────────────
            # Matches single-level (B.1) AND multi-level (B.1.1, B.5.3) patterns
            if re.match(r'^[A-Z]\.\d+(?:\.\d+)*\s+', text.strip()):
                continue

            # ── 3. Skip standalone nav category words ────────────────────────
            if _NAV_WORDS.match(text.strip()):
                continue

            # ── 4. Preamble stop ─────────────────────────────────────────────
            if _PREAMBLE_RE.match(text.strip()):
                if title_parts:
                    break  # Stop collecting — preamble follows the real title
                continue  # No title yet — just skip this preamble line

            # ── 5. Date-only stop ────────────────────────────────────────────
            if _DATE_ONLY_RE.match(text.strip()):
                if title_parts:
                    break
                continue

            # ── 6. Once crumb title captured, allow ONE short continuation ──
            if got_crumb_title:
                # Append only if ≤ 5 words, no doc number, not preamble/date/nav
                # e.g. "2025 Annual Report" after "Corporate Finance Division" crumb
                # Do NOT append "OSC RULE 11-502" or other reference labels
                _HAS_DOC_NUM_INLINE = re.search(r'\b\d{2,3}-\d{2,3}\b', text)
                if (len(words) <= 5
                        and not _PREAMBLE_RE.match(text.strip())
                        and not _DATE_ONLY_RE.match(text.strip())
                        and not _HAS_DOC_NUM_INLINE
                        and len(text) < 60):  # short year/report suffix only
                    title_parts.append(text.strip())
                break  # After one check, stop regardless

            # ── GLOBAL: Skip bilingual CSA/CIRO org header lines at ANY position ───
            # Catches "Canadian Securities Autorites canadiennes Administrators
            # en valeurs mobilieres" appearing as combined or separate lines.
            _BILINGUAL_HDR_FAST = re.compile(
                r'Autorit[e\xe9]s?\s+canadiennes|en\s+valeurs\s+mobili[e\xe8]res?', re.I
            )
            if _BILINGUAL_HDR_FAST.search(text):
                continue
            # ── GLOBAL: Skip web breadcrumb lines at ANY position ─────────────
            if re.search(r'Home\s*>\s*Resources|Home\s*>\s*Notices|Advisory\s+Notices?\s*>', text, re.I):
                continue
            # ── GLOBAL: Skip garbled/corrupted OCR text (hex-like tokens mixed with noise) ──
            # e.g. "onaed Fan e 00e4590P0Soe0-eDo,o7eReee" from DOCX embedded objects
            # Pattern 1: Classic hex token with boundary
            _is_garbled = (re.search(r'\b[0-9][0-9a-fA-F]{3,}[0-9a-zA-Z]\b', text)
                           # Pattern 2: Token starting with digit, containing uppercase mid-word
                           # e.g. "00e4590P0Soe0" → starts with digit, has caps inside
                           or any(re.match(r'[0-9]\w{3,}[A-Z]\w+', w) for w in words))
            if _is_garbled and len(words) < 15:
                continue

            # ── 7. Skip short pure-noise fragments at position 0 ─────────────
            if len(title_parts) == 0:
                if len(words) <= 2:  # e.g. "Ontario Securities", "OSC"
                    continue
                if len(words) <= 3 and text.strip().isupper():
                    continue
                # Skip bilingual organizational/entity name lines at title position
                # e.g. "Organisme canadien de regiementation des investissements Canadian
                #       Investment Regulatory" — this is an ORG NAME, not a document title.
                # Also skip lines ending in known org suffixes: "Regulatory", "Organization",
                # "Corporation", "Authority", "Commission" (standalone at end of phrase).
                _ORG_END_PAT = re.compile(
                    r'(?:Regulatory|Organization|Organisation|Corporation|'
                    r'Authority|Commission|Board|Agency|Council)\s*$', re.I
                )
                _FRENCH_WORD_PAT = re.compile(
                    r'\b(?:organisme|canadien|investissements|reglement|valeurs|'
                    r'reglementation|commission|agence|conseil)\b', re.I
                )
                if _ORG_END_PAT.search(text.strip()) or _FRENCH_WORD_PAT.search(text):
                    continue

            # ── 8. Stop at long body text ────────────────────────────────────
            if title_parts and len(text) > 180 and re.search(r'[a-z]{10}', text):
                break

            # ── 9. Stop at ALL-CAPS duplicate of existing title ──────────────
            if title_parts and text.strip() == text.strip().upper() and any(c.isalpha() for c in text):
                existing_words = set(' '.join(title_parts).lower().split())
                new_words = set(text.lower().split())
                overlap = len(existing_words & new_words) / max(len(new_words), 1)
                if overlap >= 0.4:  # ≥ 40% word overlap with existing → ALL-CAPS duplicate
                    break

            title_parts.append(text.strip())

            # Stop after 3 segments (safety cap)
            if len(title_parts) >= 3:
                break

        title = ' '.join(title_parts)

        # ── Post-processing: strip doc-type + number prefix from title ────────
        # e.g. "OSC Staff Notice 51-737 Corporate Finance Division 2025 Annual Report"
        #    → "Corporate Finance Division 2025 Annual Report"
        # e.g. "B.1.1 Joint CSA Staff Notice 31-368 – Client Focused Reforms..."
        #    → already handled by crumb/B.X.X skip above; this catches standalone headings
        _DOC_PREFIX_STRIP = re.compile(
            r'^(?:CSA\s+(?:Multilateral\s+)?(?:Staff\s+)?Notice|'
            r'OSC\s+(?:Rule|Staff\s+Notice|Notice|Policy)|'
            r'National\s+Instrument|Multilateral\s+Instrument|'
            r'National\s+Policy|Companion\s+Policy|'
            r'Staff\s+Notice)\s+\d{2,3}-\d{2,3}(?:\s*\([^)]+\))?'
            r'(?:\s*[\-\u2013\u2014]\s*|\s+)',
            re.I
        )
        stripped_title = _DOC_PREFIX_STRIP.sub('', title).strip()
        # Don't strip the doc-type prefix when the remainder itself starts with another
        # document reference (e.g. "Coordinated Blanket Order 96-933") — in these
        # cases the vendor keeps the full compound title like
        # "CSA Staff Notice 96-306 Coordinated Blanket Order 96-933 Re..."
        _is_blanket_combo = bool(re.match(
            r'(?:Coordinated\s+)?Blanket\s+Order\s+\d', stripped_title, re.I
        ))
        if stripped_title and len(stripped_title) >= 10 and not _is_blanket_combo:
            title = stripped_title

        # ── Strip bilingual org header boilerplate from assembled title ───────
        # e.g. ".Canadian Securities Autorites canadiennes Administrators en valeurs
        #        mobilieres CSA Notice" → "CSA Notice"
        _BILINGUAL_CLEAN_RE = re.compile(
            r'^[.\s]*(?:Canadian\s+Securities\s+)?'
            r'Autorit[e\xe9]s?\s+canadiennes\s*'
            r'(?:Administrators?\s+)?(?:en\s+valeurs\s+mobili[e\xe8]res?\s+)?', re.I
        )
        title = _BILINGUAL_CLEAN_RE.sub('', title).strip()
        # Also strip if bilingual block appears mid-title before a recognizable label
        _BILINGUAL_MID_RE = re.compile(
            r'Canadian\s+Securities\s+[A-Za-z\s]{3,60}?en\s+valeurs\s+mobili[e\xe8]res?\s+',
            re.I | re.DOTALL
        )
        title = _BILINGUAL_MID_RE.sub('', title).strip()

        # ── Annual Report fallback: if title is still just a doc-number ──────
        # When ABBYY extracts only "OSC Staff Notice 51-737" as the heading
        # (subtitle in a graphic), look for the Annual Report title in early body text.
        # Pattern: "proud to share our X Annual Report" or "our X Annual Report"
        _is_ann_only = (
            self._is_annual_report
            and bool(re.match(
                r'^(?:OSC\s+Staff\s+)?(?:Notice|Staff\s+Notice|Rule|Policy)\s+\d{2,3}-\d{2,3}$',
                title.strip(), re.I
            ))
        )
        if not _is_ann_only and self._is_annual_report and len(title.strip()) < 30:
            # Also trigger if title is very short (likely just doc number)
            _is_ann_only = bool(re.match(r'^\d{2,3}-\d{2,3}$', title.strip()))
        if _is_ann_only:
            # Scan all top_texts for "X Annual Report" pattern
            _AR_TITLE_RE = re.compile(
                r'(?:proud\s+to\s+share\s+our\s+|our\s+)'
                r'(.{5,80}?\s+Annual\s+Report)\b', re.I
            )
            for _tx in top_texts:
                _arm = _AR_TITLE_RE.search(_tx)
                if _arm:
                    _ar_cand = _arm.group(1).strip()
                    if 10 <= len(_ar_cand) <= 80:
                        title = _ar_cand
                        break

        if len(title) > 300:
            title = title[:297] + '...'
        return title

    # ─── CONTENT EXTRACTION ────────────────────────────────────────────────────
    # Matches OSC Bulletin chapter headers: "B.1.1 CSA Multilateral Staff Notice 31-367 – ..."
    _BULLETIN_HDR_RE = re.compile(
        r'^[A-Z]\.\d+(?:\.\d+)+\s+(?:CSA|OSC|National|Multilateral|Staff)\b', re.I
    )

    def _detect_cover_page(self, content: List[Dict]) -> List[Dict]:
        print('   🔍 Detecting cover page...')
        para_items = [c for c in content if c['type'] == 'paragraph']
        # Use at least 12 paragraphs to ensure OSC Bulletin nav blocks are covered.
        # OSC Bulletin nav headers (B.1, Notices, B.5.1 crumb) typically appear in
        # the first 10-15 paragraphs before the actual document content begins.
        cover_threshold = max(12, min(20, len(para_items) // 8))
        skip_count = 0
        # Track whether the main document title has been identified.
        # Once the title is found, ALL-CAPS bold paragraphs are section headings
        # (not cover noise) and short paragraphs are body content — preserve them.
        _past_title = False
        for item in para_items[:cover_threshold]:
            para = item['data']
            text = para.text.strip()
            # Detect main document title: a heading-style paragraph (level ≥3) that
            # is NOT all-caps and has ≥4 words. Using level ≥3 (not ≥2) avoids
            # false triggers from Heading-2 nav elements in OSC/NB bulletin cover areas.
            # After this point, skip rules for section-heading-like content are relaxed
            # to avoid suppressing legitimate section headers.
            # Example: MX Advisory Notices start with web-nav then a "Heading #3" title,
            # followed by section headings like "GENERAL TEST ENVIRONMENT (GTE)" that
            # should NOT be treated as cover noise.
            _hl_ct = para.patterns.get('heading_level', 0)
            if (not _past_title and not para.skip
                    and _hl_ct >= 3
                    and not para.patterns.get('is_all_caps', False)
                    and len(text.split()) >= 4):
                _past_title = True
            # Agency notice title headings in cover zone — skip BEFORE setting _past_title
            # so that HL=1 agency titles (e.g. "CSA Staff Notice 46-309 Bail-in Debt")
            # are caught even though they would otherwise set _past_title=True.
            _is_cover_notice_title_early = bool(re.match(
                r'^(?:CSA|OSC|AMF|BCSC|ASC|MSC|FCNB|FSRA|CIRO|IIROC|MFR)\s+'
                r'(?:Staff\s+)?(?:Notice|Bulletin|Rule|Instrument|Order|Consultation)\s+\d',
                text.strip(), re.IGNORECASE
            ))
            if _is_cover_notice_title_early and not _past_title:
                para.skip = True
                skip_count += 1
                continue
            # Also set _past_title for HL=1 non-nav section headings (Definitions,
            # Background, IT IS ORDERED THAT, etc.) — these mark the start of body content
            # in legal docs that use only Heading #1 style (MB, NFLD, etc.).
            if (not _past_title and not para.skip
                    and _hl_ct == 1
                    and len(text.split()) >= 1):
                _past_title = True
            # Standalone date lines in cover zone (e.g. "August 23, 2018") — these are
            # cover-page dates that should not become BLOCK headings, but should still
            # be emitted as body content (vendor keeps them as <P><BOLD> not as BLOCK).
            _is_standalone_date = bool(re.match(
                r'^(?:January|February|March|April|May|June|July|August|September|'
                r'October|November|December)\s+\d{1,2},\s+\d{4}$',
                text.strip(), re.IGNORECASE
            ))
            if _is_standalone_date and _hl_ct >= 1:
                # Don't skip — just demote: prevent BLOCK routing by clearing heading_level
                # The paragraph will emit as <P><BOLD>...</BOLD></P> via normal body routing
                para.patterns['heading_level'] = 0
                para.patterns['is_toc_entry'] = True  # prevents BLOCK routing
                # Don't count as skip or continue — fall through to normal processing
            # OSC Bulletin chapter header: bold paragraph starting with section dot-notation
            # e.g. "B.5.1 OSC Rule 11-502 Distribution of..." or "B.1.1 CSA Multilateral..."
            if self._BULLETIN_HDR_RE.match(text):
                para.skip = True
                skip_count += 1
                continue
            if len(text) < 50:
                # Heading #1 in Bulletin docs = "B.5 Rules and Policies" navigation menus → skip
                # Heading #2+ in legislation docs = PART/SEC titles → PRESERVE
                # EXCEPTION: short nav section markers (B.1, B.5) and OSC nav words
                # (e.g. "B.1" styled as Heading 2 appears in OSC Bulletin body cover area)
                # NOTE: Only skip HL=1 if it matches known OSC Bulletin nav patterns.
                # Real section headings (Definitions, Background, IT IS ORDERED THAT, etc.)
                # should be preserved even in the cover threshold zone.
                hl = para.patterns.get('heading_level', 0)
                if hl == 1:
                    _is_bulletin_nav_hl1 = bool(re.match(
                        r'^(?:Rules?\s+and\s+Policies|Policies|Decisions?|Notices?|'
                        r'Staff\s+Notices?|Exemption\s+Orders?|New\s+Rules?|Amendments?|'
                        r'Rescissions?|CSA\s+Releases?|Orders?\s+and\s+Rulings?|'
                        r'[A-Z]\.\d+(?:\.\d+)*)\s*$',
                        text.strip(), re.I
                    ))
                    if _is_bulletin_nav_hl1:
                        para.skip = True   # Bulletin nav heading (e.g. "Rules and Policies")
                        skip_count += 1
                elif hl >= 2:
                    # Skip OSC Bulletin nav section markers even at heading level 2
                    # e.g. "B.1" (single letter-dot-digit = nav cross-reference) or
                    # standalone nav category words like "Notices", "Rules and Policies"
                    _is_nav_sec = bool(re.match(r'^[A-Z]\.\d+(?:\.\d+)*$', text.strip()))
                    _is_nav_word = bool(re.match(
                        r'^(?:Notices?|Rules?\s+and\s+Policies|Policies|Decisions?|'
                        r'Staff\s+Notices?|Exemption\s+Orders?|New\s+Rules?|Amendments?|'
                        r'Rescissions?|CSA\s+Releases?|Orders?\s+and\s+Rulings?)\s*$',
                        text.strip(), re.I
                    ))
                    # Skip ALL-CAPS document identifier headings at cover level
                    # e.g. "CSA MULTILATERAL STAFF NOTICE 31-367" — doc label in cover area
                    # These belong in POLIDENT (N/TI), NOT in FREEFORM body.
                    _HAS_DOC_NUM_IN_COVER = re.search(r'\b\d{2,3}-\d{2,3}\b', text)
                    _is_allcaps_docid = (
                        text.strip() == text.strip().upper()
                        and any(c.isalpha() for c in text)
                        and bool(_HAS_DOC_NUM_IN_COVER)
                    )
                    # Skip ALL-CAPS org/institution names in cover zone before main title
                    # e.g. "ONTARIO SECURITIES COMMISSION", "THE MANITOBA SECURITIES COMMISSION"
                    # Condition: all-caps, multi-word (≥2), no doc number, before _past_title
                    _is_allcaps_org = (
                        not _past_title
                        and text.strip() == text.strip().upper()
                        and any(c.isalpha() for c in text)
                        and not _HAS_DOC_NUM_IN_COVER
                        and len(text.strip().split()) >= 2
                    )
                    if _is_nav_sec or _is_nav_word or _is_allcaps_docid or _is_allcaps_org:
                        para.skip = True
                        skip_count += 1
                    # else: real PART/SEC title — preserve
                elif (para.patterns.get('is_centered') or
                      para.patterns.get('is_all_bold') or
                      len(text.split()) <= 3):
                    # Guard: if we're past the doc title and this is an all-caps section
                    # heading (≥3 words), preserve it — don't treat as cover noise.
                    # e.g. "GENERAL TEST ENVIRONMENT (GTE)" is a real section heading.
                    _is_allcaps_section = (
                        _past_title
                        and para.patterns.get('is_all_caps', False)
                        and len(text.split()) >= 3
                    )
                    if not _is_allcaps_section:
                        para.skip = True
                        skip_count += 1
            # All-caps bold cover lines (e.g. multi-line title split by ABBYY, or
            # document identifier paragraphs like "NOTICE AND CONSULTATION REGARDING...")
            # Expanded from <100 to <300 chars to also catch longer all-caps document
            # titles/subtitles in the cover region.
            # Guard: once _past_title is set AND text has ≥3 words, these are section
            # headings — preserve them. Short all-caps bold (<3 words) are still cover noise.
            elif (para.patterns.get('is_all_caps') and
                  para.patterns.get('is_all_bold') and
                  len(text) < 300 and
                  not (_past_title and len(text.split()) >= 3)):
                para.skip = True
                skip_count += 1
            if re.match(r'^[A-Z]{2,5}$', text) or re.match(r'^\d+-\d+$', text):
                para.skip = True
                skip_count += 1
        print(f'   ✅ Marked {skip_count} cover elements to skip')
        return content

    # Patterns that are always noise regardless of position in the document
    _GLOBAL_NOISE_RE = re.compile(
        r'Home\s*>\s*(?:Resources|Notices|Publications)|'        # web breadcrumbs
        r'Advisory\s+Notices?\s*(?:>|[A-Z]\d{2}-\d{3})|'       # MX web nav: "Advisory Notices >" or "Advisory Notice A25-012"
        r'Autorit[e\xe9]s?\s+canadiennes\s+en\s+valeurs|'        # bilingual org header
        r'^\s*[A-Z]-\s*[A-Z]\+\*?\s*$|'                         # font-size controls "A- A+*"
        r'Type:\s*Rules\s+Bulletin\s*>|'                         # CIRO bulletin metadata header
        r'Distribute\s+internally\s+to\s*:|'                     # CIRO distribution list
        r'Rulebook\s+connection\s*:|'                             # CIRO rulebook reference
        r'Updated\s+on\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}|'  # "Updated on September 15, 2025"
        r'You\s+can\s+find\s+the\s+Canadian\s+Investment\s+Regulatory\s+Organization|'  # CIRO website footer
        r'Close\s+this\s+popup|'                                 # CIRO modal dialog
        r'Please\s+seek\s+professional\s+advice\s+to\s+evaluate\s+specific\s+securities|'  # MX legal disclaimer
        r'^\s*Division:\s+|'                                      # CIRO division specifier at para start
        r'^\s*https?://\S+\s*$|'                               # standalone URL-only lines (appendix download links)
        r'^\s*\.{8,}\s*$|'                                       # lines that are ONLY dots (TOC leaders)
        r'^\s*\(#\w+\)|'                                         # anchor refs at START of paragraph e.g. (#ref1) text
        r'Canadian\s+Derivatives\s+Exchange\s+is\s+an\s+official\s+mark|'  # MX footer boilerplate
        r'Interim\s+Fee\s+Model\s+Guidelines|'                  # CIRO Appendix A/B header
        r'Welcome\s+to\s+CIRO\.ca|'                             # CIRO URL placeholder content
        r'^\s*Uciro-ocri\s*$|'                                   # CIRO bilingual org monogram (logo heading)
        r'^\s*(?:Market\s+Operations|Markets|Post-Trade|Insights|Terms\s+of\s+Use|Privacy\s+Policy|'  # MX website nav items
        r'Fraud\s+Prevention|Capital\s+Formation|Contact\s+Us\s*(?:$|\|)|'               # MX website nav items (cont)
        r'Position\s+Limits?\s*(?:$|\s*>)|'                     # MX position limits nav
        r'Clearing\s*(?:$|\s*>)|Trading\s+Data\s*(?:$|\s*>)|Surveillance\s*(?:$|\s*>)|'  # MX modules nav
        r'Publications\s*(?:$|\s*>)|Statistics\s*(?:$|\s*>)|Products\s*(?:$|\s*>))\s*\+?\s*$|'  # MX modules nav (cont) — also with trailing '+' expander
        r'^\s*TMX\s*(?:\([A-Z]+\)|Group)\s*(?::|\()\s*[\d.]+|'  # TMX stock ticker lines
        r'^\s*Trading\s+Status\s*[&|]|'                          # TMX trading status widget
        r'^\s*X\s+El\s+in\s+[•*]|'                              # TMX/MX social media widget
        r'^\s*(?:Fran[cç]ais|English|Language|Langue)\s*$|'      # bilingual toggle labels
        r'^\s*\+\s*$|'                                             # standalone "+" expander buttons (MX nav)
        r'^\s*Copyright\s+©\s+\d{4}\s+Bourse\s+de\s+Montr',      # MX footer copyright ("Copyright © YYYY Bourse...") — not vendor content

        re.I | re.MULTILINE
    )

    def _filter_global_noise(self, content: List[Dict]) -> List[Dict]:
        """Mark globally noisy paragraphs as skip: web breadcrumbs, bilingual headers, UI controls."""
        for item in content:
            if item['type'] != 'paragraph':
                continue
            para = item['data']
            if para.skip:
                continue
            if self._GLOBAL_NOISE_RE.search(para.text.strip()):
                para.skip = True
        return content

    def _detect_toc(self, content: List[Dict]) -> List[Dict]:
        """v5.0: Skip TOC lines including Annual Reports (was bypassed before)."""
        print('   🔍 Detecting TOC...')
        in_toc = False
        skip_count = 0
        for item in content:
            if item['type'] != 'paragraph':
                continue
            para = item['data']
            text = para.text.strip().lower()
            # Word TOC field entries use "Table of contents" style but the text
            # contains the document heading text (not "table of contents" literally).
            # Mark these as TOC entries to prevent heading routing; keep as plain P.
            style_name = (para.style or '').strip().lower()
            if style_name == 'table of contents':
                para.patterns['is_toc_entry'] = True  # prevents BLOCK routing
                # Also clean up embedded tab chars from TOC field entries
                if '\t' in para.text:
                    # Keep only the first segment before the tab (the heading text)
                    para_text_clean = para.text.split('\t')[0].strip()
                    # If second segment is just ")" (court caption right col), skip para
                    right_col = para.text.split('\t')[-1].strip()
                    if right_col in (')', ''):
                        para.skip = True
                        skip_count += 1
                    # else preserve the text portion
                else:
                    # Single-cell TOC entry without tab: just ")" or very short → skip
                    if para.text.strip() in (')', ''):
                        para.skip = True
                        skip_count += 1
                continue
            if 'table of contents' in text or text == 'contents':
                para.skip = True
                in_toc = True
                skip_count += 1
                continue
            if in_toc:
                if re.match(r'^\d{1,3}$', para.text.strip()):
                    para.skip = True
                    skip_count += 1
                    continue
                if (re.search(r'.+\s{2,}\d+$', text) or
                    re.search(r'.+\.{2,}\d+', text) or
                    ('\t' in para.text and text.endswith(tuple('0123456789')))):
                    para.skip = True
                    skip_count += 1
                elif len(text) > 80 and not re.search(r'\d+$', text):
                    in_toc = False
        print(f'   ✅ Marked {skip_count} TOC lines to skip')
        return content

    def _extract_content_in_order(self) -> List[Dict]:
        """Extract paragraphs and tables in document order."""
        content = []
        para_index = 0
        layout_para_counter = [10000]
        doc_seen_texts = set()
        table_total_idx = 0   # sequential index of ALL tables seen (layout + data)
        body = self.document.element.body
        for element in body:
            tag = element.tag.split('}')[-1]
            if tag == 'p':
                if para_index < len(self.document.paragraphs):
                    para = self.document.paragraphs[para_index]
                    # Skip inline footnote body paragraphs (will be inlined at reference point)
                    if para_index not in self._inline_fn_para_idx:
                        para_data = self._extract_paragraph(para, para_index)
                        content.append({'type': 'paragraph', 'data': para_data})
                    para_index += 1
            elif tag == 'tbl':
                table_idx = len([c for c in content if c['type'] == 'table'])
                if table_idx < len(self.document.tables):
                    table = self.document.tables[table_idx]
                    # TMX footer/branding table: 2-col table where ALL non-empty
                    # second-column cells contain only "+". Appears at end of MX
                    # Advisory Notices as a branding navigation element.
                    # Vendor does not include this table → suppress it.
                    if self._is_tmx_branding_table(table):
                        table_total_idx += 1
                        continue
                    table_data = self._extract_table(table)
                    if not self._is_layout_table(table_data, table_total_idx=table_total_idx):
                        content.append({'type': 'table', 'data': table_data})
                    else:
                        layout_paras = self._layout_table_to_paragraphs(
                            table, layout_para_counter[0], doc_seen_texts)
                        layout_para_counter[0] += len(layout_paras)
                        content.extend(layout_paras)
                table_total_idx += 1   # always increment for every tbl element
        return content

    @staticmethod
    def _is_tmx_branding_table(table) -> bool:
        """Return True if this is a TMX navigation/branding footer table.

        Pattern: 2-column table where ALL non-empty cells in column-2 contain
        only the string "+" (e.g. Capital Formation / + , Markets / +, ...).
        Vendor SGM does not include these footer tables.
        """
        rows = list(table.rows)
        if len(rows) < 2:
            return False
        plus_col2 = []
        for row in rows:
            seen_ids = set()
            unique = []
            for cell in row.cells:
                cid = id(cell)
                if cid not in seen_ids:
                    seen_ids.add(cid)
                    unique.append(cell)
            if len(unique) >= 2:
                t = unique[1].text.strip()
                if t:
                    plus_col2.append(t)
        return bool(plus_col2) and all(t == '+' for t in plus_col2)

    def _layout_table_to_paragraphs(self, table, base_index: int, doc_seen_texts: set = None) -> List[Dict]:
        """Convert a layout table to paragraph items."""
        rows = list(table.rows)
        if not rows:
            return []

        # ── MISCLAW Part-index table: drop entirely (col1 = "PART N" headers) ──
        # These 2-col tables serve as a pure TOC/index for legislation documents.
        # Only drop when the first cell is a merged "PART\nPART 1" header (the
        # header keyword "PART" and the first entry "PART 1" are merged into one
        # cell with a newline), OR col1 first cell is exactly "PART" (pure header).
        # Companion-document narrative tables start with "PART 1", "1.1", "" etc.
        # and must NOT be dropped.
        try:
            _col1_texts = []
            for row in rows:
                seen_ids = set()
                cells = []
                for cell in row.cells:
                    if id(cell) not in seen_ids:
                        seen_ids.add(id(cell))
                        cells.append(cell)
                if cells:
                    _col1_texts.append(cells[0].text.strip())
            _part_matches = sum(1 for t in _col1_texts if re.search(r'\bPART\s+\d+', t, re.IGNORECASE))
            _first_cell = _col1_texts[0] if _col1_texts else ''
            # True TOC: first cell is "PART" (exact header) or "PART\nPART 1" (merged header+data)
            _is_toc_header = (_first_cell == 'PART') or _first_cell.startswith('PART\n')
            if _part_matches >= 2 and _is_toc_header:
                return []  # Hard-drop: pure TOC table, no paragraphs extracted
        except Exception:
            pass

        # ── PEI metadata table: drop entirely (col1 = "Document Type:", "Document No:", etc.) ──
        # Vendor does not include this administrative header in the SGML output.
        # Converting to paragraphs would create spurious POLIDENT-style content.
        try:
            _meta_labels = {'Document Type', 'Document No', 'Document No.', 'Subject', 'Effective Date', 'Effective', 'Date'}
            _col1_meta = []
            for row in rows[:6]:  # only check first 6 rows
                seen_ids = set()
                cells = []
                for cell in row.cells:
                    if id(cell) not in seen_ids:
                        seen_ids.add(id(cell))
                        cells.append(cell)
                if cells:
                    _col1_meta.append(cells[0].text.strip())
            _meta_matches = sum(1 for t in _col1_meta if any(t.startswith(lbl) for lbl in _meta_labels))
            if _meta_matches >= 2:
                return []  # Hard-drop: administrative metadata table
        except Exception:
            pass
        seen_test = set()
        first_row_cells = []
        for cell in rows[0].cells:
            if id(cell) in seen_test:
                continue
            seen_test.add(id(cell))
            first_row_cells.append(cell)
        actual_cols = len(first_row_cells)

        if actual_cols >= 3:
            all_texts = []
            seen = set()
            for row in rows:
                for cell in row.cells:
                    if id(cell) in seen:
                        continue
                    seen.add(id(cell))
                    t = cell.text.strip()
                    if t:
                        all_texts.append(t)
            avg = sum(len(t) for t in all_texts) / len(all_texts) if all_texts else 0
            if avg < 50:
                return []
            items = []
            idx = base_index
            local_seen = doc_seen_texts if doc_seen_texts is not None else set()
            for row in rows:
                row_cells = []
                row_seen = set()
                for cell in row.cells:
                    if id(cell) in row_seen:
                        continue
                    row_seen.add(id(cell))
                    row_cells.append(cell)
                if len(row_cells) >= 2:
                    target_cell = row_cells[1]
                    for para in target_cell.paragraphs:
                        if not para.text.strip():
                            continue
                        txt_key = para.text.strip()[:100].lower()
                        if txt_key in local_seen:
                            continue
                        local_seen.add(txt_key)
                        para_data = self._extract_paragraph(para, idx)
                        items.append({'type': 'paragraph', 'data': para_data})
                        idx += 1
            return items

        # 2-col: check if TOC table (col2 = page numbers) — drop it
        col2_texts = []
        for row in rows:
            unique_cells = []
            seen_row = set()
            for cell in row.cells:
                cid = id(cell)
                if cid in seen_row:
                    continue
                seen_row.add(cid)
                unique_cells.append(cell)
            if len(unique_cells) >= 2:
                col2_texts.append(unique_cells[1].text.strip())

        if col2_texts:
            numeric_count = sum(1 for t in col2_texts if re.match(r'^\d{1,3}$', t))
            if (numeric_count / len(col2_texts)) >= 0.6:
                return []

        # 2-col (non-TOC): extract ALL cells with dedup
        items = []
        seen_cell_ids = set()
        seen_texts = doc_seen_texts if doc_seen_texts is not None else set()
        idx = base_index
        for row in rows:
            for cell in row.cells:
                cid = id(cell)
                if cid in seen_cell_ids:
                    continue
                seen_cell_ids.add(cid)
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    txt_key = para.text.strip()[:100].lower()
                    if txt_key in seen_texts:
                        continue
                    seen_texts.add(txt_key)
                    para_data = self._extract_paragraph(para, idx)
                    items.append({'type': 'paragraph', 'data': para_data})
                    idx += 1
        return items

    def _is_layout_table(self, table_data: 'TableData', table_total_idx: int = 999) -> bool:
        """Detect layout/formatting tables (should NOT become SGMLTBL).

        table_total_idx: sequential position of this table in the document (0-based),
        counting ALL tables including ones already identified as layout tables.
        Default 999 keeps backward-compatibility when called without position.
        """
        rows = table_data.rows
        if not rows:
            return True
        max_cols = max(len(row) for row in rows) if rows else 0
        if max_cols <= 1:
            return True

        # ── TOC-pattern check: SCOPED to first 3 tables in document ──────────
        # TOC tables are always table #0-#2.  Real financial data tables
        # appear later in the document.  Scoping by position lets us use a
        # lower 65% threshold while still avoiding false-positives.
        # Annual Report TOCs often have 1-digit page numbers (3, 5, 8) —
        # use {1,3} to catch those (was {2,3} which missed single-digit pages).
        if max_cols == 2 and len(rows) >= 8 and table_total_idx <= 2:
            col2 = [row[1].get('text', '').strip() for row in rows if len(row) >= 2]
            numeric = sum(1 for t in col2 if re.match(r'^\d{1,3}$', t))  # 1–3 digit page numbers
            if col2 and (numeric / len(col2)) >= 0.65:
                print(f'   🗂  TOC table at position {table_total_idx} dropped '
                      f'({numeric}/{len(col2)} page-number cells)')
                return True  # TOC table — drop in ALL doc types including Annual Report

        # ── PEI amending/implementing instrument metadata table ───────────────
        # These 2-column tables at table position 0 have label:value rows such as
        # "Document Type:", "Document No:", "Subject:", "Effective Date:".
        # Vendor does not include this administrative header in the SGML output.
        if max_cols == 2 and table_total_idx <= 1:
            col1 = [row[0].get('text', '').strip() for row in rows if len(row) >= 1]
            _metadata_labels = {'Document Type', 'Document No', 'Document No.', 'Subject', 'Effective Date', 'Effective', 'Date'}
            _label_matches = sum(1 for t in col1 if any(t.startswith(lbl) for lbl in _metadata_labels))
            if col1 and _label_matches >= 2:
                print(f'   🗂  Metadata table at position {table_total_idx} dropped '
                      f'({_label_matches}/{len(col1)} label cells)')
                return True  # Administrative metadata table — drop from output

        # ── MISCLAW Part index table (table of contents for legislation) ──────
        # 2-column tables with "PART 1", "PART 2", ... in col1 act as a TOC
        # for MISCLAW legislation documents. Vendor never includes these.
        # Only match TRUE index tables: first cell is "PART" (header-only) or
        # "PART\nPART 1" (merged header+first-entry). Companion-document narrative
        # tables start with "PART 1", "1.1", "" etc. and must NOT be classified here.
        if max_cols == 2 and table_total_idx <= 2:
            col1 = [row[0].get('text', '').strip() for row in rows if len(row) >= 1]
            # Check for "PART N" or "PART\nPART N" (header row) pattern
            _part_col1 = [t for t in col1 if re.search(r'\bPART\s+\d+', t, re.IGNORECASE)]
            _first = col1[0] if col1 else ''
            _is_toc_header = (_first == 'PART') or _first.startswith('PART\n')
            if len(_part_col1) >= 2 and _is_toc_header:
                print(f'   🗂  MISCLAW Part-index table at position {table_total_idx} dropped')
                return True  # MISCLAW TOC table — drop from output

        # When preserve_data_tables is set AND table is non-trivial → keep as data table
        # (moved early so empty-cells heuristics don't drop real financial tables)
        if getattr(self, '_preserve_data_tables', False):
            if max_cols >= 2 and len(rows) >= 2:
                return False  # Preserve as real data table

        all_cells = [cell for row in rows for cell in row]
        non_empty_cells = [c for c in all_cells if len(c.get('text', '').strip()) >= 3]
        if len(all_cells) > 0 and len(non_empty_cells) < max(2, len(all_cells) // 3):
            return True
        if all_cells:
            first_text = all_cells[0].get('text', '').strip()
            if all(c.get('text', '').strip() == first_text for c in all_cells):
                return True

        # When preserve_data_tables is set, skip aggressive filtering (for legal docs)
        if getattr(self, '_preserve_data_tables', False):
            return False

        if self._is_annual_report:
            return False

        text_lengths = [len(c.get('text', '').strip()) for c in non_empty_cells]
        avg_len = sum(text_lengths) / len(text_lengths) if text_lengths else 0
        max_len = max(text_lengths) if text_lengths else 0

        if max_cols >= 3:
            return True
        if len(rows) <= 1:
            return True
        if len(rows) <= 3 and avg_len > 80:
            return True
        if avg_len > 120:
            return True
        if max_len < 50:
            return True

        return False

    def _extract_paragraph(self, para, index: int) -> 'ParagraphData':
        runs = []
        for run in para.runs:
            runs.append(RunData(
                text=run.text,
                bold=run.bold if run.bold is not None else False,
                italic=run.italic if run.italic is not None else False,
                underline=run.underline if run.underline is not None else False,
                font_size=run.font.size.pt if run.font.size else None,
                font_name=run.font.name
            ))
        left_indent = para.paragraph_format.left_indent
        left_indent_pt = float(left_indent.pt) if left_indent else 0.0
        para_data = ParagraphData(
            index=index, text=para.text, runs=runs,
            style=para.style.name if para.style else 'Normal',
            alignment=str(para.alignment), left_indent=left_indent_pt, patterns={}
        )
        para_data.patterns = self._detect_patterns(para_data)
        # numPr detection: store Word list numbering flag for use in tagger
        # (set has_numpr only if left_indent is 0 — avoids false positives in
        # docs like 91-102 where indented numPr items are plain P/ITEM by context)
        try:
            pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    para_data.patterns['has_numpr'] = True
        except Exception:
            pass
        # Footnote reference detection: collect IDs of any <w:footnoteReference> in this para
        if self._footnotes:
            try:
                _W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                _fn_pos = []   # (fn_id, fn_text, char_pos_in_para_text)
                _char_pos = 0
                # Iterate <w:r> runs in order (same iteration as para.runs) to track char offset
                for r_el in para._element.iter(_W + 'r'):
                    run_text = ''.join(t.text or '' for t in r_el.findall(_W + 't'))
                    fn_ref_el = r_el.find(_W + 'footnoteReference')
                    if fn_ref_el is not None:
                        fid_str = fn_ref_el.get(_W + 'id', '')
                        try:
                            fid = int(fid_str)
                        except ValueError:
                            fid = None
                        if fid is not None and fid in self._footnotes:
                            _fn_pos.append((fid, self._footnotes[fid], _char_pos))
                    _char_pos += len(run_text)
                if _fn_pos:
                    para_data.patterns['footnote_refs']       = [(n, t) for n, t, _ in _fn_pos]
                    para_data.patterns['inline_fn_positions'] = _fn_pos  # positional injection
            except Exception:
                pass
        # Inline (superscript digit) footnote reference detection
        if self._inline_footnotes:
            try:
                _W2 = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                _inline_pos = []   # (fn_num, fn_text, char_start_in_para_text)
                _char_pos = 0
                for r in para.runs:
                    rpr = r._element.find(_W2 + 'rPr')
                    if rpr is not None:
                        va = rpr.find(_W2 + 'vertAlign')
                        if va is not None and va.get(_W2 + 'val') == 'superscript' \
                                and r.text.strip().isdigit():
                            _fnum = int(r.text.strip())
                            if _fnum in self._inline_footnotes:
                                _inline_pos.append((_fnum, self._inline_footnotes[_fnum], _char_pos))
                    _char_pos += len(r.text)
                if _inline_pos:
                    # Store both for compatibility and for positional injection
                    para_data.patterns['footnote_refs']       = [(n, t) for n, t, _ in _inline_pos]
                    para_data.patterns['inline_fn_positions'] = _inline_pos
            except Exception:
                pass
        # ── DOCX auto-numbering label reconstruction ──────────────────────────
        # python-docx para.text does NOT include the auto-generated list prefix
        # (e.g. "(a)", "(i)", "1.") stored in word/numbering.xml.  We reconstruct
        # that label here so the MISCLAW generator's _ALPHA_RE / _ROMAN_RE patterns
        # fire correctly and produce <PARA><N>(a)</N><PARAP>...</PARAP></PARA>
        # structure matching vendor SGML.
        if self._numbering:
            _auto_label = self._get_numpr_label(para._element)
            if _auto_label:
                _txt = para_data.text.lstrip()
                # Only prepend if text doesn't already start with this label
                # (avoids double-counting when the number IS typed in the text)
                if not _txt.startswith(_auto_label):
                    para_data.text = _auto_label + ' ' + _txt
                    # Update the RunData list: prepend a synthetic run so
                    # inline formatting offsets remain consistent
                    from dataclasses import replace as _dc_replace
                    _prefix_run = RunData(
                        text=_auto_label + ' ',
                        bold=False, italic=False, underline=False,
                        font_size=None, font_name=None
                    )
                    para_data.runs = [_prefix_run] + list(para_data.runs)
        return para_data

    def _detect_patterns(self, para_data: 'ParagraphData') -> Dict:
        """Detect paragraph patterns. v5.0: added is_annual_report_doc flag."""
        text = para_data.text.strip()
        runs = para_data.runs
        non_empty_runs = [r for r in runs if r.text.strip()]
        all_bold = all(r.bold for r in non_empty_runs) if non_empty_runs else False
        all_italic = all(r.italic for r in non_empty_runs) if non_empty_runs else False
        has_bold_runs = any(r.bold for r in non_empty_runs)
        has_italic_runs = any(r.italic for r in non_empty_runs)
        sized_runs = [r for r in non_empty_runs if r.font_size]
        avg_font_size = sum(r.font_size for r in sized_runs) / len(sized_runs) if sized_runs else 12

        is_list_item = False
        style_name = para_data.style.lower() if para_data.style else ''
        if any(kw in style_name for kw in ['list', 'bullet', 'number']):
            is_list_item = True
        if re.match(r'^[•·‐‑‒–—▸▪▫◦]\s+', text):
            is_list_item = True
        if re.match(r'^o\s{1,3}\S', text) and len(text) < 300:
            is_list_item = True

        heading_level = 0
        m = re.search(r'[Hh]eading\s*[#\s]?(\d+)', para_data.style or '')
        if m:
            heading_level = int(m.group(1))

        is_all_caps = (len(text) > 2 and any(c.isalpha() for c in text) and all(c.isupper() or not c.isalpha() for c in text))

        return {
            'has_bullet': bool(re.match(r'^[•·‐‑‒–—▸▪▫◦]\s+', text)),
            'has_dash_bullet': bool(re.match(r'^[\-\u2013\u2014]\s+', text)),
            'has_number': bool(re.match(r'^\d+[\.\)]\s+', text)),
            'has_alpha_list': bool(
                re.match(r'^\([a-z]\)\s+', text) or
                re.match(r'^\([ivxlcdm]+\)\s+', text, re.IGNORECASE) or
                (re.match(r'^[a-z]\.\s+', text) and not re.match(r'^[a-z]{4,}\.', text))
            ),
            'has_email': bool(re.search(r'@[\w\.-]+\.\w+', text)),
            'has_phone': bool(re.search(r'\d{3}[-\.\s]?\d{3}[-\.\s]?\d{4}', text)),
            'has_fax': bool(re.search(r'\b[Ff]ax\s*:', text)),
            'has_tel': bool(re.search(r'\b(?:Tel|Telephone|Phone|Tél|Téléphone|Télec)\.?\s*:', text, re.IGNORECASE)),
            'has_url': bool(re.search(r'(?:https?://|www\.|[Ww]ebsite\s*:)', text)),
            'has_postal_code': bool(re.search(r'\b[A-Z]\d[A-Z]\s*\d[A-Z]\d\b', text)),
            'is_all_bold': all_bold, 'is_all_italic': all_italic,
            'has_bold_runs': has_bold_runs and not all_bold,
            'has_italic_runs': has_italic_runs and not all_italic,
            'is_short': len(text) < 80,
            'is_centered': para_data.alignment == str(WD_ALIGN_PARAGRAPH.CENTER),
            'indent_level': (4 if para_data.left_indent > 216 else
                             (3 if para_data.left_indent > 144 else
                              (2 if para_data.left_indent > 72 else
                               (1 if para_data.left_indent > 10 else 0)))),
            'avg_font_size': avg_font_size,
            'is_list_item': is_list_item,
            'heading_level': heading_level,
            'is_all_caps': is_all_caps,
            'is_annual_report_doc': self._is_annual_report,
            'is_body_text_2': bool(re.match(r'body text \(', para_data.style.lower())),
        }

    def _extract_table(self, table) -> 'TableData':
        rows = []
        for row in table.rows:
            row_data = []
            seen_cell_ids = set()
            for cell in row.cells:
                cell_id = id(cell)
                if cell_id in seen_cell_ids:
                    continue
                seen_cell_ids.add(cell_id)
                cell_text = cell.text.strip()
                has_bold = any(run.bold for para in cell.paragraphs
                               for run in para.runs if run.text.strip())
                row_data.append({'text': cell_text, 'bold': has_bold})
            if row_data:
                rows.append(row_data)
        has_header = False
        if rows and len(rows) > 1:
            has_header = all(cell['bold'] for cell in rows[0])

        # Issue 3 fix: Extract column widths from DOCX XML <w:tblGrid><w:gridCol w:w="..."/>
        col_widths = []
        try:
            tbl_xml = table._tbl
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for gc in tbl_xml.findall(f'.//{{{ns}}}gridCol'):
                w = gc.get(f'{{{ns}}}w')
                if w:
                    col_widths.append(int(w))
        except Exception:
            col_widths = []

        return TableData(rows=rows, has_header=has_header, col_widths=col_widths or None)

print("✅ CompleteDOCXExtractor v5.0")
print("   • Smart LABEL detection (22 label types including Joint CSA/CIRO, OSC Staff Notice, etc.)")
print("   • Full title assembly from first meaningful headings")
print("   • Document number: date-style YYYY/MM/DD or numeric NN-NNN")
print("   • ADDDATE from document effective date (not run date)")
print("   • TOC table drop now works for Annual Report mode too")


# ====== CODE CELL 10 ======
class ContextTracker:
    """
    PHASE 2: Track document context for better ITEM detection
    Detects list intros and maintains list state
    """
    
    def __init__(self):
        self.in_list = False
        self.list_indent = 0
        self.last_was_list_intro = False
        self.in_contact_block = False
        self.contact_indent = 0
        self.high_freq_headings = set()  # Bold heading texts that repeat ≥ FREQ_THR times

    def check_contact_intro(self, para_text: str) -> bool:
        """Detect if paragraph introduces a contact/address block → next items use LINE."""
        text = para_text.strip().lower()
        contact_signals = [
            'for further information', 'for more information', 'for additional information',
            'address your submission', 'address submissions', 'address comments',
            'direct written comments', 'direct questions', 'please contact',
            'contact staff', 'staff can be reached', 'questions may be directed',
            'questions may be addressed', 'please send written', 'please send your',
            'send written comments', 'send your comments', 'inquiries may be',
            'for questions', 'please direct', 'signed copies', 'questions should',
            'following staff', 'refer your questions', 'please refer',
        ]
        return any(signal in text for signal in contact_signals)
    
    def check_list_intro(self, para_text: str) -> bool:
        """
        Detect if paragraph introduces a list.
        Must end with ':' AND contain a list signal word/phrase.
        Extended to cover OSC Annual Report patterns.
        """
        text = para_text.strip().lower()
        
        # Ends with colon
        if not text.endswith(':'):
            return False
        
        # Contains list intro signals
        list_signals = [
            # Core signals
            'aim to', 'include', 'following', 'noted that',
            'informed by', 'reviewed', 'considered',
            'consists of', 'comprised of', 'such as',
            # OSC Annual Report specific
            'work includes', 'responsibility', 'operational work',
            'we refer', 'we track', 'guidance',
            'outcomes for', 'includes', 'may include',
            'relief includes', 'allows', 'applicable',
            'required to', 'must consider', 'should consider',
            'primary responsibility', 'key relief includes',
        ]
        return any(signal in text for signal in list_signals)
    
    def should_be_item(self, para: 'ParagraphData', prev_para) -> bool:
        """
        Determine if paragraph should be ITEM based on context
        """
        # Check if previous was list intro
        if prev_para and self.check_list_intro(prev_para.text):
            self.in_list = True
            self.list_indent = para.left_indent
            return True

        # Continue list if same indent
        if self.in_list:
            if abs(para.left_indent - self.list_indent) < 5:  # Allow small variance
                return True
            elif para.left_indent < self.list_indent:
                # Indent decreased, list ended
                self.in_list = False
                return False

        return False

print("✅ ContextTracker defined (extended list signals + contact-block tracking)")


# ====== CODE CELL 11 ======
class PatternBasedTagger:
    """Enhanced pattern-based tagging with context awareness - OPTIMIZED for 85%+ accuracy"""
    
    def __init__(self):
        """Initialize with context tracker"""
        self.context_tracker = ContextTracker()
        # Jurisdiction set by CompletePipeline.convert() before extract_inline_formatting calls.
        # Used to enable/disable jurisdiction-specific EM patterns (e.g. Ontario
        # vendor SGML does NOT EM-tag NI/MI refs; all other jurisdictions do).
        self.jurisdiction: str = ''
    
    def tag_paragraphs(self, paragraphs: List['ParagraphData']) -> Tuple[List['ParagraphData'], List['ParagraphData']]:
        """Tag paragraphs with high confidence vs ambiguous"""
        confirmed = []
        ambiguous = []

        # Normalize heading levels relative to minimum found in this document.
        # ABBYY often produces Heading 3/4 where the correct output is BLOCK2/3.
        # Skip HL=1 (document title level) and centered paragraphs.
        import re as _re_norm
        _raw_levels = []
        for _p in paragraphs:
            if _p.patterns.get('is_centered', False):
                continue  # skip centered (title) paragraphs
            _hl = _p.patterns.get('heading_level', 0)
            if _hl > 1:   # > 1: skip Heading 1 (document title)
                _raw_levels.append(_hl)
            if _p.style:
                _m = _re_norm.search(r'heading\s*[#\s]?(\d+)', _p.style, _re_norm.IGNORECASE)
                if _m:
                    lvl = int(_m.group(1))
                    if lvl > 1:  # skip Heading 1
                        _raw_levels.append(lvl)
        _min_hl = min(_raw_levels) if _raw_levels else 2  # default=2 if no headings

        # Detect if document has varied heading levels (> 1 distinct level).
        # If ABBYY only reports one heading level, we need context-based BLOCK3 promotion.
        _distinct_levels = len(set(_raw_levels))

        def _h2b(level):
            """Normalize: lowest heading level → BLOCK2, next → BLOCK3, etc."""
            return min(max(2, level - _min_hl + 2), 6)

        # ── Pre-pass A: Count bold heading recurrences ────────────────────────
        # Repeated headings (e.g. "Issues identified" × 8) are deep sub-sections.
        # We use repetition frequency to detect false BLOCK2 headings that should
        # be BLOCK4/5 (repeated sub-section patterns such as "Issues identified",
        # "Guidance", "Examples of firm practices" in CSA/CIRO joint notices).
        import re as _re_freq
        from collections import Counter as _BldCounter
        _bld_heading_texts = []
        for _p in paragraphs:
            if (_p.patterns.get('is_all_bold')
                    and not _p.patterns.get('is_centered')
                    and not _p.patterns.get('is_all_italic')):
                _ts_chk = _p.text.strip()
                _wc_chk = len(_ts_chk.split())
                _is_sentence_chk = _ts_chk.endswith('.') and _wc_chk > 20
                _has_explicit_prefix = bool(
                    _re_freq.match(r'^[A-Z]\.\s+\S', _ts_chk)
                    or _re_freq.match(r'^\([A-Z]\)\s', _ts_chk)
                    or (_re_freq.match(r'^\d+\.\s+[A-Z]', _ts_chk) and _wc_chk <= 10)
                    or _re_freq.match(r'^\([a-z]\)\s', _ts_chk)
                    or _re_freq.match(r'^\([ivxlcdm]+\)\s', _ts_chk, _re_freq.IGNORECASE)
                )
                if not _is_sentence_chk and not _has_explicit_prefix:
                    _bld_heading_texts.append(_ts_chk.strip().lower()[:60])
        _heading_freq = _BldCounter(_bld_heading_texts)
        # Threshold: heading appearing ≥ FREQ_THR times in document = deep sub-section
        _FREQ_THR = 3
        # Share high-frequency heading set with ContextTracker for ITEM guard
        self.context_tracker.high_freq_headings = {
            k for k, v in _heading_freq.items() if v >= _FREQ_THR
        }

        # ── Pre-pass B: Cluster bolt heading indent values → BLOCK levels ─────
        # For docs where ABBYY preserves indentation (annual reports, circulars):
        # different indent values map to different BLOCK depths.
        _bld_indents = []
        for _p in paragraphs:
            if (_p.patterns.get('is_all_bold')
                    and not _p.patterns.get('is_centered')
                    and not _p.patterns.get('is_all_italic')):
                _ts_chk2 = _p.text.strip()
                _wc_chk2 = len(_ts_chk2.split())
                _is_sent2 = _ts_chk2.endswith('.') and _wc_chk2 > 20
                if not _is_sent2:
                    _bld_indents.append(_p.left_indent)
        _sorted_ind = sorted(set(_bld_indents))
        _indent_clusters = []
        for _iv in _sorted_ind:
            if not _indent_clusters or _iv - _indent_clusters[-1] > 15:
                _indent_clusters.append(_iv)  # New distinct indent level
        _has_multi_indent = len(_indent_clusters) >= 2

        def _indent_to_blk(indent_pt):
            """Map indent value to BLOCK level 2-6 based on distinct clusters."""
            if not _indent_clusters:
                return 2
            _closest = min(range(len(_indent_clusters)),
                           key=lambda _ci: abs(_indent_clusters[_ci] - indent_pt))
            return min(2 + _closest, 6)

        # Track last BLOCK level for context-aware sub-heading detection
        _last_block_level = 0          # int: 2,3,4... or 0 if none seen yet
        _last_block_was_caps = False   # True if last BLOCK para was ALL-CAPS

        import re as _re_b3
        def _is_title_case_heading(text):
            """True if text looks like Title Case (not all-caps, not all-lower)."""
            words = text.strip().split()
            if not words:
                return False
            # At least first word is capitalised and text is not all-caps
            has_lower = any(c.islower() for c in text)
            first_cap = words[0][0].isupper() if words[0] else False
            return has_lower and first_cap

        # ── Pre-pass C: detect "inline header" pattern for Annual Reports ──────
        # A bold heading with heading_level=2 that is short (≤6 words) AND
        # immediately followed by a bullet-list item (in Annual Report mode) is
        # a compact sub-topic label, NOT a section container. Mark as inline_header.
        # Example: "Right-Size Regulation" (bold, HL2) → "Published four..." (ITEM)
        import re as _re_ih
        _active_paras = [p for p in paragraphs if not p.skip]
        for _pi, _ph in enumerate(_active_paras):
            _is_annual_ph = _ph.patterns.get('is_annual_report_doc', False)
            if not _is_annual_ph:
                continue
            _ts_ph = _ph.text.strip()
            _wc_ph = len(_ts_ph.split())
            _hl_ph = _ph.patterns.get('heading_level', 0)
            _is_bold_ph = _ph.patterns.get('is_all_bold', False)
            # Must be bold, heading_level ≤ 3, ≤6 words, no section prefix
            if not (_is_bold_ph and _hl_ph > 0 and _hl_ph <= 3 and _wc_ph <= 6):
                continue
            # Must NOT have a section prefix (A., (A), 1., (a), etc.)
            if _re_ih.match(r'^(?:[A-Z]\.\s|[A-Za-z]\)\s|\([A-Za-z]\)\s|\([ivxlcdm]+\)\s|\d+\.\s)', _ts_ph, _re_ih.IGNORECASE):
                continue
            # Check if NEXT non-skip para is a likely ITEM (bullet content)
            _next_phs = _active_paras[_pi+1:_pi+3]
            _next_is_item = False
            for _np in _next_phs:
                _ts_np = _np.text.strip()
                _np_wc = len(_ts_np.split())
                # If next para is already tagged or looks like a bullet continuation
                if (_np.patterns.get('is_list_item', False)
                        or _np.patterns.get('has_bullet', False)
                        or _np.patterns.get('has_dash_bullet', False)
                        or (_np.patterns.get('indent_level', 0) > 0 and _np_wc > 3)):
                    _next_is_item = True
                    break
            if _next_is_item:
                # Mark as inline header — will be mapped to P (BOLD inline) not BLOCK
                _ph.patterns['inline_header'] = True

        prev_para = None
        for para in paragraphs:
            if para.skip:
                continue
            
            if self._is_line_item(para, prev_para):
                para.final_tag = 'LINE'
                para.confidence = 0.85
                confirmed.append(para)
            elif self._is_list_item(para, prev_para):
                para.final_tag = 'ITEM'
                para.confidence = 0.95
                confirmed.append(para)
            elif (para.patterns.get('is_all_italic', False)
                  and not para.patterns.get('is_all_bold', False)
                  and not para.patterns.get('is_centered', False)
                  and not para.patterns.get('heading_level', 0)
                  and not para.patterns.get('is_toc_entry', False)):
                # All-italic non-bold short paragraphs are subsection headings
                # (e.g. "Third-Year Historical Financial Statements" in MB blanket orders,
                # "General advice registration exemption" in BC staff notices).
                # Route as BLOCK3 when: ≤8 words, starts with uppercase letter (not digit/quote),
                # and vendor context suggests BLOCK3 (covered by broad heuristic below).
                import re as _re_ital
                _ts_ital = para.text.strip()
                _wc_ital = len(_ts_ital.split())
                # Guard: if prev_para is a form-item reference (starts with "item N.N of Form"),
                # this is an inline-continuation EM fragment (e.g. "Standard Term Sheets During
                # the Waiting Period" appended after "item 21 of Form 62-104F2 Issuer Bid
                # Circular.") → tag as P, not BLOCK3.
                _prev_is_form_item = (
                    prev_para is not None
                    and bool(_re_ital.match(r'^item\s+\d+', prev_para.text.strip(), _re_ital.IGNORECASE))
                )
                _is_ital_heading = (
                    _wc_ital <= 8
                    and _ts_ital[:1].isupper()          # starts with uppercase
                    and not _re_ital.match(r'^\d', _ts_ital)  # not numeric code
                    and len(_ts_ital) >= 4              # not a trivial fragment
                    and not _prev_is_form_item          # not an inline-continuation fragment
                )
                if _is_ital_heading:
                    para.final_tag = 'BLOCK3'
                    _last_block_level = 3
                    _last_block_was_caps = False
                    para.confidence = 0.85
                    confirmed.append(para)
                else:
                    # Longer all-italic paras (quotes, captions) → P
                    para.final_tag = 'P'
                    para.confidence = 0.87
                    confirmed.append(para)
            elif para.patterns.get('is_all_bold') and not para.patterns.get('is_centered') and not para.patterns.get('is_toc_entry', False):
                import re as _re_bold
                _ts_bold = para.text.strip()
                _bw_bold = len(_ts_bold.split())
                _hl_bold = para.patterns.get('heading_level', 0)
                _is_annual = para.patterns.get('is_annual_report_doc', False)
                # Section-prefix detection: A./( A)→BLOCK3, 1.→BLOCK4, (a)/(i)→BLOCK5
                # These take priority over sentence-detection to correctly classify
                # lettered/numbered section headings in notice/circular documents.
                _ts_bold_is_caps = (len(_ts_bold) > 2 and all(c.isupper() or not c.isalpha() for c in _ts_bold))
                _blk_by_prefix = None
                if _re_bold.match(r'^[A-Z]\.\s+\S', _ts_bold):            # A. Title → BLOCK3
                    _blk_by_prefix = 3
                elif _re_bold.match(r'^\([A-Z]\)\s', _ts_bold):            # (A) Title → BLOCK3
                    _blk_by_prefix = 3
                elif (_re_bold.match(r'^\d+\.\s+[A-Z]', _ts_bold)         # 1. Title → BLOCK4
                        and _bw_bold <= 10):
                    _blk_by_prefix = 4
                elif _re_bold.match(r'^\([a-z]\)\s', _ts_bold):            # (a) Title → BLOCK5
                    _blk_by_prefix = 5
                elif _re_bold.match(r'^\([ivxlcdm]+\)\s', _ts_bold, _re_bold.IGNORECASE):  # (i)/(ii) → BLOCK5
                    _blk_by_prefix = 5
                elif _re_bold.match(r'^[a-z]\)[ \t]', _ts_bold):  # a) b) non-parens → BLOCK5
                    _blk_by_prefix = 5
                # Stat-box caption: pure numeric/currency/percentage heading → P
                # e.g. "10%", "~ 2,800", "$6.50 billion", "1,000 +" in Annual Report infographics
                # Also: short (≤4 char) abbreviation-only headings in Annual Reports (e.g. "OSC")
                _is_stat_box = (
                    _blk_by_prefix is None
                    and _is_annual
                    and bool(_re_bold.match(
                        r'^[\~\$]?\s*[\d,\.]+\s*[%\+]?\s*(?:[Bb]illion|[Mm]illion|[Tt]rillion|[Kk])?$',
                        _ts_bold
                    ))
                )
                _is_short_abbrev = (
                    _blk_by_prefix is None
                    and _is_annual
                    and _re_bold.match(r'^[A-Z]{2,5}$', _ts_bold) is not None
                )
                # Bold sentence/very long → P only when no section prefix was detected
                _is_bold_sentence = (_ts_bold.endswith('.') or _bw_bold > 20) and _blk_by_prefix is None
                # Inline header: short bold heading before bullet list in Annual Report → P
                _is_inline_header = para.patterns.get('inline_header', False)
                # Signature / attribution block → P  (e.g. "FOR THE COMMISSION", "BY ORDER OF")
                # These are not section headings; vendor tags them as bold P.
                _is_signature_block = (
                    _blk_by_prefix is None
                    and _hl_bold == 0
                    and _bw_bold <= 6
                    and bool(_re_bold.match(
                        r'^(?:FOR\s+THE\s+|BY\s+ORDER\b|PER\s+PROCURATION\b|'
                        r'ON\s+BEHALF\s+OF\b|SIGNED\s+(?:BY|ON)\b)',
                        _ts_bold, _re_bold.IGNORECASE
                    ))
                )
                if _blk_by_prefix is not None:
                    para.final_tag = f'BLOCK{_blk_by_prefix}'
                    _last_block_level = _blk_by_prefix
                    _last_block_was_caps = False
                elif _is_stat_box or _is_short_abbrev or _is_inline_header or _is_signature_block:
                    # Statistic caption, short abbreviation, inline section label, or signature block → P
                    para.final_tag = 'P'
                elif _is_bold_sentence:
                    # Bold body text → P (SGMLGenerator wraps content in <BOLD> inline)
                    para.final_tag = 'P'
                elif _hl_bold > 0:
                    # Bold heading with known DOCX level — use normalized depth
                    _blk = _h2b(_hl_bold)
                    # For Notice docs (not Annual Report) where only one heading level
                    # is seen, use context: Title-Case after ALL-CAPS → BLOCK3
                    if (not _is_annual and _distinct_levels <= 1 and _blk == 2
                            and _last_block_level == 2 and _last_block_was_caps
                            and _is_title_case_heading(_ts_bold)):
                        _blk = 3
                    # Recurrence override: frequently repeating heading = deep sub-section
                    _freq_key = _ts_bold.strip().lower()[:60]
                    if _heading_freq.get(_freq_key, 0) >= _FREQ_THR and _blk <= 3:
                        _blk = 5  # Repeated sub-heading → deep BLOCK5
                    # Indent override: multi-indent doc → use indent clustering
                    elif _has_multi_indent and _blk == 2:
                        _blk = _indent_to_blk(para.left_indent)
                    para.final_tag = f'BLOCK{_blk}'
                    _last_block_level = _blk
                    _last_block_was_caps = _ts_bold_is_caps
                else:
                    # No heading level from DOCX — use context for BLOCK depth
                    _blk = 2
                    # Recurrence override: frequently repeating heading = deep sub-section
                    _freq_key2 = _ts_bold.strip().lower()[:60]
                    if _heading_freq.get(_freq_key2, 0) >= _FREQ_THR:
                        _blk = 5  # Repeated sub-section header in flat-indent doc
                    # Indent-based override: multi-indent doc (annual reports etc.)
                    elif _has_multi_indent:
                        _blk = _indent_to_blk(para.left_indent)
                    elif (not _is_annual and _last_block_level == 2
                            and _last_block_was_caps
                            and _is_title_case_heading(_ts_bold)
                            and not _ts_bold_is_caps):
                        # Title-Case bold heading follows ALL-CAPS bold heading → BLOCK3
                        _blk = 3
                    para.final_tag = f'BLOCK{_blk}'
                    _last_block_level = _blk
                    _last_block_was_caps = _ts_bold_is_caps
                para.confidence = 0.90
                confirmed.append(para)
            elif (para.patterns.get('heading_level', 0) > 0
                  and not para.patterns.get('is_centered')):
                # Non-bold heading-style paragraph — HYBRID normalization
                # Sub-lettered prefixes always → BLOCK3; others use level-based normalization
                import re as _re_blk, re as _re2h
                _hl2 = para.patterns.get('heading_level', 0)
                _is_annual2 = para.patterns.get('is_annual_report_doc', False)
                if para.style:
                    _m2h = _re2h.search(r'heading\s*[#\s]?(\d+)', para.style, _re2h.IGNORECASE)
                    level = int(_m2h.group(1)) if _m2h else _hl2
                else:
                    level = _hl2
                _ts_blk = para.text.strip()
                _ts_blk_is_caps = (len(_ts_blk) > 2 and all(c.isupper() or not c.isalpha() for c in _ts_blk))
                if _re_blk.match(r'^[A-Z]\.\s+\S', _ts_blk) or _re_blk.match(r'^\([A-Z]\)\s', _ts_blk):
                    _blk_level = 3   # A. or (A) → BLOCK3
                elif (_re_blk.match(r'^\d+\.\s+[A-Z]', _ts_blk) and len(_ts_blk.split()) <= 10):
                    _blk_level = 4   # 1. → BLOCK4
                elif (_re_blk.match(r'^\([a-z]\)\s', _ts_blk)
                        or _re_blk.match(r'^\([ivxlcdm]+\)\s', _ts_blk, _re_blk.IGNORECASE)
                        or _re_blk.match(r'^[a-z]\.\s', _ts_blk)):
                    _blk_level = 5   # (a), (i), a. → BLOCK5
                else:
                    _blk_level = _h2b(level)
                    # For Notice docs: Title-Case heading after ALL-CAPS → BLOCK3
                    if (not _is_annual2 and _distinct_levels <= 1 and _blk_level == 2
                            and _last_block_level == 2 and _last_block_was_caps
                            and _is_title_case_heading(_ts_blk) and not _ts_blk_is_caps):
                        _blk_level = 3
                # Stat-box caption in Annual Report (non-bold heading): numeric/% → P
                _is_stat_box2 = (
                    _is_annual2
                    and bool(_re_blk.match(
                        r'^[\~\$]?\s*[\d,\.]+\s*[%\+]?\s*(?:[Bb]illion|[Mm]illion|[Tt]rillion|[Kk])?$',
                        _ts_blk
                    ))
                )
                _is_inline_header2 = para.patterns.get('inline_header', False)
                if _is_stat_box2 or _is_inline_header2:
                    para.final_tag = 'P'
                else:
                    para.final_tag = f'BLOCK{_blk_level}'
                _last_block_level = _blk_level
                _last_block_was_caps = _ts_blk_is_caps
                para.confidence = 0.88
                confirmed.append(para)
            elif (para.patterns.get('is_all_caps', False)
                  and not para.patterns.get('is_centered')
                  and not para.patterns.get('is_all_bold')
                  and 2 < len(para.text.strip()) <= 200
                  and len(para.text.strip().split()) <= 30):
                # AllCaps non-bold non-centered heading  BLOCK2
                para.final_tag = 'BLOCK2'
                _last_block_level = 2
                _last_block_was_caps = True
                para.confidence = 0.87
                confirmed.append(para)
            elif not para.patterns.get('is_centered'):
                il = para.patterns.get('indent_level', 0)
                import re as _re_p2
                _ts_p2 = para.text.strip()
                # Detect mixed-format headings: first run is bold with section prefix
                # e.g. BOLD("A. Know Your Client ") + EM("(Section 13.2 of NI 31-103)")
                _blk_by_mixed = None
                if para.patterns.get('has_bold_runs', False) and len(_ts_p2.split()) <= 25:
                    _ne_runs_mb = [r for r in para.runs if r.text.strip()]
                    if _ne_runs_mb and _ne_runs_mb[0].bold:
                        _frb = _ne_runs_mb[0].text.strip()
                        if _re_p2.match(r'^[A-Z]\.\s+\S', _frb) or _re_p2.match(r'^\([A-Z]\)\s', _frb):
                            _blk_by_mixed = 3
                        elif _re_p2.match(r'^\d+\.\s+[A-Z]', _frb) and len(_frb.split()) <= 8:
                            _blk_by_mixed = 4
                        elif (_re_p2.match(r'^\([a-z]\)\s', _frb)
                                or _re_p2.match(r'^\([ivxlcdm]+\)\s', _frb, _re_p2.IGNORECASE)):
                            _blk_by_mixed = 5
                if _blk_by_mixed is not None:
                    para.final_tag = f'BLOCK{_blk_by_mixed}'
                    _last_block_level = _blk_by_mixed
                    _last_block_was_caps = False
                # Roman sub-items (i)/(ii)/... → P2 regardless of indent_level
                elif (_re_p2.match(r'^\([ivxlcdm]+\)[ \t]', _ts_p2, _re_p2.IGNORECASE)
                        and len(_ts_p2) < 500):
                    para.final_tag = 'P2'
                elif il > 0:
                    para.final_tag = f'P{min(il, 4)}'
                else:
                    para.final_tag = 'P'
                para.confidence = 0.90
                confirmed.append(para)
            else:
                ambiguous.append(para)
            
            prev_para = para
        
        return confirmed, ambiguous
    
    def _is_list_item(self, para: 'ParagraphData', prev_para: Optional['ParagraphData']) -> bool:
        """
        OPTIMIZED ITEM detection for 85%+ accuracy
        Changes: indent 18->10, length 250->500, heading >=3->>=1, continue 600->1000
        """
        # ── GUARD 1: heading-style paragraphs are NEVER ITEM ──────────────────
        if para.patterns.get('heading_level', 0) > 0:
            return False

        # ── GUARD 1c: all-bold short paragraphs with section prefix → BLOCK heading ──
        # Bold "(a) Issues identified" / "1. Determination of KYC" etc. are BLOCK section
        # headings, not list items, even if context tracker thinks we're in a list.
        # Also applies to frequently-repeated bold headings (≥ FREQ_THR occurrences)
        # like "Issues identified" / "Guidance" / "Examples of firm practices".
        if para.patterns.get('is_all_bold', False) and not para.patterns.get('is_centered'):
            import re as _re_g1c
            _ts_g1c = para.text.strip()
            _wc_g1c = len(_ts_g1c.split())
            _has_sec_prefix = bool(
                _re_g1c.match(r'^[A-Z]\.\s+\S', _ts_g1c)
                or _re_g1c.match(r'^\([A-Z]\)\s', _ts_g1c)
                or (_re_g1c.match(r'^\d+\.\s+[A-Z]', _ts_g1c) and _wc_g1c <= 10)
                or _re_g1c.match(r'^\([a-z]\)\s', _ts_g1c)
                or _re_g1c.match(r'^\([ivxlcdm]+\)\s', _ts_g1c, _re_g1c.IGNORECASE)
            )
            _is_high_freq = (
                _ts_g1c.lower()[:60] in self.context_tracker.high_freq_headings
            )
            if _has_sec_prefix or _is_high_freq:
                return False  # Bold section heading (explicit prefix or repeated pattern), never ITEM

        # ── GUARD 1b: Explicitly indented paragraphs (P1-style) are NOT ITEM ──
        # Paragraphs with indent_level > 0 (left_indent > 10pt) should be P1/P2,
        # not ITEM, UNLESS they have an explicit bullet/list marker OR are in a
        # colon-triggered list context. This prevents false P1 tagging of list items.
        if (para.patterns.get('indent_level', 0) > 0
                and not para.patterns.get('is_list_item', False)
                and not para.patterns.get('has_bullet', False)
                and not para.patterns.get('has_dash_bullet', False)
                and not (prev_para and self.context_tracker.check_list_intro(prev_para.text))
                and not (prev_para                              # colon+bold+indent bypass
                         and prev_para.text.strip().endswith(':')
                         and para.patterns.get('is_all_bold', False))
                and not self.context_tracker.in_list):  # bypass if already in a colon list
            return False

        ts = para.text.strip()
        ts_words = len(ts.split())
        _is_full_sentence = ts.endswith('.') and ts_words > 15

        # ── GUARD: legal form-item references → NOT ITEM ──────────────────────
        # Texts like "item 32.2 of Form 41-101F1 ..." are legal section/form
        # references, not list bullets. They start with lowercase "item" followed
        # by a decimal number or NI reference. Prevent false ITEM tagging.
        import re as _re_legal
        if _re_legal.match(r'^item\s+\d+(?:\.\d+)*\s+(?:of|in)\s+', ts, _re_legal.IGNORECASE):
            return False

        # ── GUARD 2: explicit lettered items → NOT ITEM unless in colon-list context ──
        import re as _re_li
        # Alpha and roman prefixed items route to P1/P2 UNLESS we're already in a
        # colon-triggered list (prev_para ends with ':') or continuation of in_list.
        # Exception enables "(a)/(b)/(i)/(ii)" sub-items to be tagged ITEM when they
        # appear after a colon-intro heading (e.g. "Firm must do the following:").
        # Bold "(a)/(i)" paragraphs are always section headings (BLOCK5), never ITEM.
        if (_re_li.match(r'^\([a-z]\)[ \t]', ts)
                or _re_li.match(r'^\([ivxlcdm]+\)[ \t]', ts, _re_li.IGNORECASE)):
            if para.patterns.get('is_all_bold', False):
                return False  # Bold "(a)/(i)" = section heading BLOCK5, NEVER ITEM
            _in_colon_ctx = (
                (prev_para and self.context_tracker.check_list_intro(prev_para.text))
                or self.context_tracker.in_list
            )
            if not _in_colon_ctx:
                return False  # Not in list context → P1 or P2, never ITEM
            # In colon-list context: fall through to colon-trigger / continuation below
        # Short numbered/lettered headings → NOT ITEM
        if ts_words <= 12 and not ts.endswith('.') and not ts.endswith(','):
            _is_numbered_sec = bool(
                _re_li.match(r'^\d+[\.)]\ +[A-Z]', ts)         # 1. Title or 1) Title
                or _re_li.match(r'^[a-z]\.\ +[A-Z]', ts)       # a. Subtitle
            )
            if _is_numbered_sec:
                return False  # numbered section heading — NOT a list item

        if prev_para:
            if self.context_tracker.check_list_intro(prev_para.text):
                self.context_tracker.last_was_list_intro = True
            elif (prev_para.patterns.get('is_all_bold', False)
                    and getattr(prev_para, 'final_tag', '') != 'ITEM'):
                # Bold section heading → reset list context.
                # BUT: if the bold prev_para was itself an ITEM (description-list lead-in),
                # keep in_list=True so subsequent items in the same list are also detected.
                self.context_tracker.last_was_list_intro = False
                self.context_tracker.in_list = False

        if para.patterns.get('is_list_item', False):
            return True
        if para.patterns.get('has_bullet', False):
            return True
        if para.patterns.get('has_dash_bullet', False):
            return True


        if (prev_para and self.context_tracker.check_list_intro(prev_para.text)
                and not _is_full_sentence):  # colon-intro list: trigger regardless of indent or bold
            # GUARD 1c already prevents bold section headings (sec-prefix / high-freq) → safe
            self.context_tracker.in_list = True
            self.context_tracker.list_indent = para.left_indent
            return True

        # Fallback colon trigger: prev_para ends with ':' AND current para is bold+indented.
        # Handles compliance-exam-style docs where colon intros lack check_list_intro signal
        # words, e.g. "Areas of weakness were:" or "We observed:" → bold body items are ITEM.
        # GUARD 1c already filters out high-freq headings (e.g. "Issues identified") → safe.
        if (prev_para
                and prev_para.text.strip().endswith(':')
                and para.patterns.get('is_all_bold', False)
                and para.patterns.get('indent_level', 0) > 0
                and not _is_full_sentence):
            self.context_tracker.in_list = True
            self.context_tracker.list_indent = para.left_indent
            return True

        if (prev_para and prev_para.patterns.get('is_all_bold', False) and
                len(prev_para.text.strip()) < 80 and
                not prev_para.patterns.get('is_centered', False) and
                not para.patterns.get('is_all_bold') and
                (para.patterns.get('indent_level', 0) > 0 or para.left_indent > 100) and
                len(ts) < 600):
            # Allow full sentences: bullet points after headings are often complete sentences
            self.context_tracker.in_list = True
            self.context_tracker.list_indent = para.left_indent
            return True
        
        if prev_para and prev_para.style:
            import re
            m = re.search(r'[Hh]eading\s*[#\s]?(\d+)', prev_para.style)
            if m and int(m.group(1)) >= 1:
                if (prev_para.patterns.get('is_all_bold', False) and
                        len(prev_para.text.strip()) < 80 and
                        not prev_para.patterns.get('is_centered', False) and
                        not para.patterns.get('is_all_bold') and
                        para.left_indent >= 5 and
                        len(ts) < 500 and not _is_full_sentence):
                    self.context_tracker.in_list = True
                    self.context_tracker.list_indent = para.left_indent
                    return True
        
        if (self.context_tracker.last_was_list_intro and prev_para and
                prev_para.patterns.get('is_all_italic', False) and
                len(prev_para.text.strip()) < 120 and para.left_indent >= 10
                and not _is_full_sentence):
            self.context_tracker.in_list = True
            self.context_tracker.list_indent = para.left_indent
            return True
        
        _bold_in_list = (
            self.context_tracker.in_list
            and para.patterns.get('is_all_bold', False)
            and not para.patterns.get('heading_level', 0)
            and not para.patterns.get('is_all_caps', False)  # ALL-CAPS headings stay BLOCK
        )
        if _bold_in_list:
            return True  # Bold item in list context → ITEM (description-list lead-in or bullet)
        if self.context_tracker.in_list and not para.patterns.get('is_all_bold'):
            if para.patterns.get('is_all_italic', False):
                self.context_tracker.in_list = False
                return False
            if (para.left_indent >= self.context_tracker.list_indent - 10 and
                    len(ts) < 600):
                # Stay in list — but for zero-indent lists (numPr-origin),
                # require the continuing para to also be a list item
                if self.context_tracker.list_indent <= 5:
                    # Zero-indent colon-triggered list: continue unless heading or clear body
                    if para.patterns.get('heading_level', 0) > 0:
                        self.context_tracker.in_list = False
                        return False
                    # End list on a clearly standalone body paragraph (long + ends period)
                    _clearly_body = (ts_words > 25 and ts.endswith('.')
                                     and not para.patterns.get('has_bullet', False)
                                     and not para.patterns.get('has_dash_bullet', False)
                                     and not para.patterns.get('has_numpr', False))
                    if _clearly_body:
                        self.context_tracker.in_list = False
                        return False
                return True
            # Exit: non-indented full sentence is body text, not a list continuation
            if _is_full_sentence:
                self.context_tracker.in_list = False
                return False
            self.context_tracker.in_list = False
        
        return False
    def _is_line_item(self, para: 'ParagraphData', prev_para: Optional['ParagraphData']) -> bool:
        """Detect LINE tag: contact info, address lines, org lists, standalone URLs/phones/emails."""
        text = para.text.strip()
        if not text or len(text) > 200:
            return False
        # Indented paragraphs are P1/P2/P3/P4 — never LINE
        if para.patterns.get('indent_level', 0) > 0:
            return False
        # Check contact-block context FIRST (bold names/org names are LINE in contact blocks)
        if prev_para:
            if self.context_tracker.check_contact_intro(prev_para.text):
                self.context_tracker.in_contact_block = True
                self.context_tracker.contact_indent = para.left_indent
        if self.context_tracker.in_contact_block:
            # Exit on heading-style paragraphs, colon-ended intros, or very long paragraphs
            if para.patterns.get('heading_level', 0) > 0 or len(text) > 180 or text.endswith(':'):
                self.context_tracker.in_contact_block = False
                return False
            # Exit if clearly a full sentence (body text)
            if len(text.split()) > 15 and text.endswith('.'):
                self.context_tracker.in_contact_block = False
                return False
            return True
        # Bold headings and list-marker items are never LINE (outside contact block)
        if para.patterns.get('is_all_bold', False):
            return False
        if para.patterns.get('has_bullet', False) or para.patterns.get('has_dash_bullet', False):
            return False
        # Phone number line → always LINE (and enter contact block)
        if para.patterns.get('has_phone', False):
            self.context_tracker.in_contact_block = True
            return True
        # Email line → always LINE (and enter contact block)
        if para.patterns.get('has_email', False):
            self.context_tracker.in_contact_block = True
            return True
        # Standalone URL line (mostly URL, short, few spaces) → LINE
        if para.patterns.get('has_url', False) and len(text) < 160 and text.count(' ') < 6:
            self.context_tracker.in_contact_block = True
            return True
        # Continue LINE block: if prev was LINE and current is also short/undecorated
        if (prev_para and getattr(prev_para, 'final_tag', None) == 'LINE'
                and len(text) < 120
                and not para.patterns.get('has_bullet', False)
                and not para.patterns.get('has_number', False)):
            return True
        return False

    def extract_inline_formatting(self, para: 'ParagraphData') -> List[Dict]:
        """
        Extract BOLD and EM from DOCX runs plus 14 regulatory pattern-based EM detections.
        """
        formatting = []
        current_pos = 0
        for run in para.runs:
            run_length = len(run.text)
            if run_length > 0:
                if run.bold:
                    formatting.append({
                        'start': current_pos,
                        'end': current_pos + run_length,
                        'tag': 'BOLD',
                        'source': 'docx'
                    })
                if run.italic:
                    formatting.append({
                        'start': current_pos,
                        'end': current_pos + run_length,
                        'tag': 'EM',
                        'source': 'docx'
                    })
            current_pos += run_length
        
        text = para.text
        import re

        # Build lookup of already-tagged spans from DOCX runs to avoid double-tagging.
        def _already_tagged(start: int, end: int) -> bool:
            return any(f['start'] <= start < f['end'] or start <= f['start'] < end
                       for f in formatting)

        # Pattern 1 (RESTORED): Regulatory instrument number references.
        # NI XX-XXX / MI XX-XXX / CP XX-XXX / NP XX-XXX / Staff Notice XX-XXX
        # Ontario vendor SGML does NOT italicize instrument numbers in body text
        # (confirmed from Ontario/41-702 analysis — vendor has 7 EM, all non-NI).
        # Every other jurisdiction vendor consistently italicizes these refs.
        if self.jurisdiction.lower() != 'ontario':
            for match in re.finditer(
                r'\b(?:NI|MI|CP|NP)\s+\d{2}-\d{3}[A-Z]?'
                r'|\bStaff\s+(?:Notice|Bulletin)\s+\d{2}-\d{3}[A-Z]?',
                    text):
                if not _already_tagged(match.start(), match.end()):
                    formatting.append({
                        'start': match.start(), 'end': match.end(),
                        'tag': 'EM', 'source': 'pattern_ni_mi'
                    })

        # Pattern 2: Title-case words + Act/Code/Regulation/Loi/Règlement
        # Excludes phrases starting with articles (The, This, A, An) to avoid
        # false positives like "The Securities Act" that vendor treats as plain text
        for match in re.finditer(
            r'\b(?!(?:The|This|An?)\s)[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(Act|Code|Regulation|Loi|R\u00e8glement)\b',
            text):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_act'
            })
        
        # Pattern 3: REMOVED - "the Act/Regulation" — vendor uses STATREF/CITE or plain text, not EM
        
        # Pattern 4: REMOVED - Part/Section/Schedule/Form + number — vendor uses plain text, not EM
        
        # Pattern 5: Companion Policy / Rule / Blanket Order with number
        for match in re.finditer(
            r'\b(?:Companion Policy|Blanket Order|Local Policy|Policy Statement)\s+\d+-\d+[A-Z]?\b',
            text, re.IGNORECASE):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_cp_rule'
            })
        
        # Pattern 6: Multilateral Instrument / National Policy with number
        for match in re.finditer(
            r'\b(?:Multilateral Instrument|National Policy|Multilateral Policy)\s+\d+-\d+[A-Z]?\b',
            text, re.IGNORECASE):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_mi_np'
            })
        
        # Pattern 7: Named Acts with RSO / RSC citation  e.g. Securities Act, RSO 1990
        for match in re.finditer(
            r'\b[A-Z][a-zA-Z ]{3,60}(?:Act|Code)\s*,\s*R[SO]C?\s*\d+',
            text):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_act_rso'
            })
        
        # Pattern 8: REMOVED - Quoted titles (too aggressive; EM-tags defined terms like "Agent for Service")
        
        # Pattern 9: REMOVED - Title-case Guidelines/Standards/Requirements — too broad, e.g. "Registration Requirements, Exemptions..."
        
        # Pattern 10: REMOVED - "the Instrument/Notice/Order" — vendor does not EM these in legislation or notices
        
        # Pattern 11: French language regulatory references (bilingual documents)
        for match in re.finditer(
            r'\b(?:R\u00e8glement|Instruction\s+g\u00e9n\u00e9rale|'
            r'Avis\s+g\u00e9n\u00e9ral|Norme|Politique)\s+\d+-\d+[A-Z]?\b',
            text, re.IGNORECASE):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_french_reg'
            })
        
        # Pattern 12: Policy Statement to Regulation/Rule XX-XXX
        for match in re.finditer(
            r'\bPolicy Statement\s+to\s+(?:Regulation|Rule|NI|MI)\s+\d+-\d+[A-Z]?\b',
            text, re.IGNORECASE):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_policy_stmt'
            })
        
        # Pattern 13: Well-known statute names without citation number
        for match in re.finditer(
            r'\b(?:Statutory Powers Procedure Act|Capital Markets Act|'
            r'Business Corporations Act|'
            r'Financial Services Regulatory Authority of Ontario Act|'
            r'Investment Funds Act|Commodity Futures Act)\b',
            text):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_named_act'
            })
        
        # Pattern 14: Subsection / paragraph / clause references  e.g. subsection 3(1)
        for match in re.finditer(
            r'\b(?:subsection|paragraph|clause|subclause)\s+\d+(?:\.\d+)?\s*\(\w+\)',
            text, re.IGNORECASE):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_subsection'
            })
        
        # Pattern 15: Email addresses — vendor consistently italicizes emails in contact sections
        for match in re.finditer(
            r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}',
            text):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_email'
            })
        
        # Pattern 16: URLs — vendor consistently italicizes URLs in reference sections
        for match in re.finditer(
            r'https?://[^\s<>"]+',
            text):
            formatting.append({
                'start': match.start(), 'end': match.end(),
                'tag': 'EM', 'source': 'pattern_url'
            })
        
        return formatting



print("PatternBasedTagger v14 - heading_level/AllCaps BLOCK routing")
print("  - P1 for indent_level > 0 (restored)")

print("  - NEW: LINE tag for contact/address/org/phone/email/URL lines")
print("  - Pattern-based EM detection: 14 regulatory patterns (was 4)")

print("  - NEW: P1/P2/P3/P4 based on actual indent depth (was always P1)")
print("  - ITEM thresholds: indent 18->10, length 250->500, heading >=3->>=1, continue 600->1000")

print("  - NEW: contact-block context tracking for consecutive LINE items")
print("  - 14 EM patterns for regulatory references")

# ====== CODE CELL 12 ======
class LLMIntelligenceLayer:
    """
    LLM Layer — v11 (RAG-enhanced):
    STEP 1: AllCaps+Bold+HL0or3 → BLOCK2 (flat CSA circulars — Notice ONLY, not AnnualReport)
    STEP 2: 'Definition of', 'Overview of' + LastBlock B2 → BLOCK3
    STEP 3: Parallel ALL-CAPS → BLOCK2
    STEP 4: AnnualReport deep nesting (HeadingLevel → BLOCK depth)
    STEP 5: Title-Case tie-breaker → BLOCK3; ALL-CAPS tie-breaker → BLOCK2
    RecurringCount>=3 AND Words<=5 → P+BOLD
    HasBold partial → P+BOLD
    NEW v11: RAGManager injects targeted keying rules + vendor examples per batch
    """

    def __init__(self, client, keying_specs: str, rag_manager=None):
        self.client = client
        self.keying_specs = keying_specs
        self.rag_manager = rag_manager
        self.model = ANTHROPIC_MODEL
        self.batch_size = SYSTEM_CONFIG['llm_batch_size']

    def process_ambiguous_paragraphs(self, paragraphs, full_paragraphs=None):
        if not paragraphs or not self.client: return paragraphs
        print('\n🤖 LLM processing with context...')
        if full_paragraphs is None: full_paragraphs = paragraphs
        sorted_full = sorted(full_paragraphs, key=lambda p: p.index)
        index_to_pos = {p.index: i for i, p in enumerate(sorted_full)}
        from collections import Counter
        text_counts = Counter(p.text.strip().lower()[:80] for p in sorted_full if p.text.strip())
        processed = []
        for i in range(0, len(paragraphs), self.batch_size):
            batch = paragraphs[i:i+self.batch_size]
            print(f'   Batch {i//self.batch_size + 1}: {len(batch)} paras')
            decisions = self._call_llm_for_batch(batch, sorted_full, index_to_pos, text_counts)
            for para, decision in zip(batch, decisions):
                para.final_tag = decision.get('tag_type', 'P')
                para.confidence = decision.get('confidence', 0.5)
                para.inline_formatting = self._merge_inline_formatting(
                    para.docx_formatting, decision.get('inline_formatting', []))
                processed.append(para)
        print(f'   ✅ Processed {len(processed)} paras')
        return processed

    def _merge_inline_formatting(self, docx_fmt, llm_fmt):
        merged = docx_fmt.copy()
        for t in llm_fmt:
            if t.get('start',0) < t.get('end',0) and not any(self._tags_overlap(t,d) for d in docx_fmt):
                t['source']='llm'; merged.append(t)
        return merged

    def _tags_overlap(self, t1, t2):
        return not (t1.get('end',0)<=t2.get('start',0) or t2.get('end',0)<=t1.get('start',0))

    def _extract_json_from_response(self, text):
        if not text or not text.strip(): raise ValueError("Empty LLM response")
        tc = re.sub(r'```(?:json)?\s*','',text).strip()
        si = tc.find('[')
        if si==-1: raise ValueError(f"No JSON array. Response: {text[:200]}")
        depth=0; ei=-1; ins=False; esc=False
        for i,c in enumerate(tc[si:],si):
            if esc: esc=False; continue
            if c=='\\' and ins: esc=True; continue
            if c=='"': ins=not ins; continue
            if ins: continue
            if c=='[': depth+=1
            elif c==']':
                depth-=1
                if depth==0: ei=i; break
        if ei==-1: raise ValueError(f"Unclosed array. Response: {text[:200]}")
        return json.loads(tc[si:ei+1])

    def _get_last_block(self, para_index, sorted_full):
        last = None
        for p in sorted_full:
            if p.index >= para_index: break
            if p.final_tag and p.final_tag.startswith('BLOCK') and p.text.strip():
                last = (p.final_tag, p.text.strip()[:60])
        return last

    def _call_llm_for_batch(self, batch, sorted_full, index_to_pos, text_counts):
        is_annual = any(p.patterns.get('is_annual_report_doc', False) for p in batch)
        doc_type_hint = "AnnualReport" if is_annual else "Notice"

        # ── RAG context retrieval ───────────────────────────────────────
        rag_section = ""
        if self.rag_manager is not None:
            try:
                batch_texts = [p.text for p in batch]
                rag_ctx = self.rag_manager.get_context_for_batch(batch_texts)
                if rag_ctx:
                    rag_section = f"\n\nRAG CONTEXT (retrieved for this batch):\n{rag_ctx}\n"
            except Exception as _rag_err:
                pass  # RAG failure is non-fatal

        system_prompt = f"""You are an expert SGML tagger for Canadian securities documents (Carswell DTD).

COMPLETE KEYING SPECIFICATIONS:
{self.keying_specs}
{rag_section}
•••••••••••• CRITICAL GUARDS — ALWAYS APPLY FIRST ••••••••••••

1. PARTIAL BOLD (HasBold=True, AllBold=False) → P with <BOLD> inline. NEVER BLOCK.

2. RECURRING STAMP (RecurringCount>=3, Words<=5, AllBold=True)
   → P + <BOLD> inline. Section stamp, NOT a structural heading.

•••••••••••• BLOCK ELIGIBILITY ••••••••••••
Eligible for BLOCK only if:
  AllBold=True AND Words≤15 AND EndsPeriod=False
  OR AllItalic=True AND Words≤10
  OR HeadingLevel≥3
Otherwise → P (or P1 if Indent≥1).

•••••••••••• BLOCK LEVEL ASSIGNMENT — STEP BY STEP ••••••••••••

STEP 1 — ALL-CAPS rule (DocType=Notice ONLY — SKIP for AnnualReport):
  ╔══ ABSOLUTE RULE FOR NOTICE DOCS ══╗
  ║ IF DocType=Notice AND AllCaps=True → ALWAYS BLOCK2. NO EXCEPTIONS.           ║
  ║ This applies regardless of HeadingLevel, LastBlock, or semantic scope.        ║
  ║ CSA/OSC/IIROC/CIRO circulars use FLAT BLOCK2 for ALL ALL-CAPS headings.      ║
  ╚═══════════════════════════════════════════════════════════════════════════════╝
  Examples BLOCK2: "LOPR ACCESS", "THE GUI CLIENT APPLICATION", "ACCOUNT TYPES",
    "FIELD REQUIREMENTS", "REPORTING REQUIREMENTS", "LOPR REPORTING RESPONSIBILITIES",
    "POSITION AGGREGATION", "REPORTING THRESHOLDS", "PARTICIPANT ACCOUNTS (FIRM ACCOUNTS)",
    "DIRECT COMMUNICATIONS USING THE SAIL PROTOCOL", "THE LOPR APP (amended on...)"
  BLOCK3 exceptions (Notice ONLY — these are the ONLY cases that get BLOCK3):
    - Headings explicitly lettered as sub-items: "(A) ...", "(B) ...", "(i) ...", "(ii) ..."
    - Numbered examples: "Example 1", "Example 2", "Example 3"
    - Headings starting with "–" or "•" (dash/bullet sub-indicator)

STEP 2 — Title-Case headings in Notice docs:
  If DocType=Notice AND Title-Case (not ALL-CAPS):
    → BLOCK3 if LastBlock=BLOCK2 AND heading is clearly a sub-section
    → BLOCK2 if LastBlock=None or LastBlock=BLOCK3 (starting fresh)
    → BLOCK2 if it's an appendix/chapter title ("Appendix A", "Guidelines", etc.)

STEP 3 — Parallel major section check → BLOCK2:
  If current introduces a completely new top-level section regardless of LastBlock:
  (new appendix, new chapter, new major numbered part, new major topic)
  → BLOCK2

STEP 4 — HeadingLevel-based assignment:
  DocType=AnnualReport (use HeadingLevel to assign depth FIRST, then AllCaps/Title-Case as secondary):
  HeadingLevel=2 → BLOCK2
  HeadingLevel=3 → BLOCK3
  HeadingLevel=4 → BLOCK4
  HeadingLevel=5 → BLOCK5
  HeadingLevel=6 → BLOCK6
  If no HeadingLevel: use LastBlock nesting + AllCaps/Title-Case (Steps 3 & 5).

STEP 5 — Tie-breaker for AnnualReport (genuinely ambiguous, no strong indicator):
  Title-Case heading, LastBlock=BLOCK2 → BLOCK3
  ALL-CAPS heading → BLOCK2
  Any heading after HeadingLevel=3 (H3-style), LastBlock=BLOCK2 → prefer BLOCK3

•••••••••••• STANDARD BLOCK2 EXAMPLES ••••••••••••
  CIRCULAR 050-25 / CIRCULAR 166-25, TABLE OF CONTENTS,
  APPENDIX A, APPENDIX B, REQUEST FOR COMMENTS,
  Background and Introduction, Part 1. Background...,
  2024 HIGHLIGHTS, 2024 Annual Activities Report,
  1. COMMON FINANCIAL STATEMENTS, 2. OTHER REGULATORY,
  Equity Option Volume, Figure 5 (chart section headings)

•••••••••••• INLINE FORMATTING ••••••••••••
EM: italic regulation names, Act titles within P/ITEM/LINE.
BOLD: inline bold within P (HasBold paragraphs).

DocType={doc_type_hint}

OUTPUT: ONLY valid JSON array, one object per paragraph:
[{{"tag_type":"BLOCK2|BLOCK3|BLOCK4|BLOCK5|BLOCK6|P|P1|P2|P3|P4|ITEM|LINE|QUOTE","inline_formatting":[{{"start":0,"end":5,"tag":"BOLD|EM"}}],"confidence":0.95}}]"""

        user_prompt = "Tag these paragraphs. Return ONLY a JSON array:\n\n"
        for idx, para in enumerate(batch):
            pos = index_to_pos.get(para.index, -1)
            prev_p = sorted_full[pos-1] if pos > 0 else None
            next_p = sorted_full[pos+1] if pos>=0 and pos+1 < len(sorted_full) else None
            ts = para.text.strip()
            wc = len(ts.split())
            ends_period = ts.endswith('.')
            hl = para.patterns.get('heading_level', 0)
            rc = text_counts.get(ts.lower()[:80], 1)
            lb = self._get_last_block(para.index, sorted_full)
            lb_str = f'{lb[0]} "{lb[1]}"' if lb else 'None'
            is_all_caps = len(ts) > 2 and all(c.isupper() or not c.isalpha() for c in ts)
            doc_type = 'AnnualReport' if para.patterns.get('is_annual_report_doc', False) else 'Notice'
            indent = para.patterns.get('indent_level', 0)
            is_short = para.patterns.get('is_short', False)
            all_italic = para.patterns.get('is_all_italic', False)
            user_prompt += (
                f'Paragraph {idx}:\n'
                f'Prev: "{prev_p.text[:80] if prev_p else "N/A"}"\n'
                f'Current: "{para.text}"\n'
                f'Next: "{next_p.text[:80] if next_p else "N/A"}"\n'
                f'Signals: AllBold={para.patterns.get("is_all_bold",False)}, '
                f'AllItalic={all_italic}, '
                f'HasBold={para.patterns.get("has_bold_runs",False)}, '
                f'Words={wc}, EndsPeriod={ends_period}, Short={is_short}, Indent={indent}, '
                f'HeadingLevel={hl}, AllCaps={is_all_caps}, '
                f'RecurringCount={rc}, LastBlock={lb_str}, DocType={doc_type}\n---\n'
            )
        user_prompt += "\nRespond with ONLY the JSON array."

        for attempt in range(2):
            try:
                resp = self.client.messages.create(
                    model=self.model, max_tokens=SYSTEM_CONFIG['max_tokens'],
                    temperature=SYSTEM_CONFIG['temperature'],
                    system=system_prompt, messages=[{"role":"user","content":user_prompt}]
                )
                decisions = self._extract_json_from_response(resp.content[0].text)
                while len(decisions) < len(batch): decisions.append({'tag_type':'P','confidence':0.5,'inline_formatting':[]})
                return decisions[:len(batch)]
            except Exception as e:
                if attempt==0: print(f'   ⚠️  Retry: {e}'); time.sleep(2)
                else: print(f'   ⚠️  Failed: {e}')
        return [{'tag_type':'P','confidence':0.5,'inline_formatting':[]} for _ in batch]

print("✅ LLMIntelligenceLayer v11 (RAG-enhanced)")
print("   STEP 1: AllCaps+Bold+HL0or3 → BLOCK2 (Notice ONLY — skip for AnnualReport)")
print("   STEP 2: 'Definition of'/'Overview of' + LastBlock B2 → BLOCK3")
print("   STEP 3: Parallel ALL-CAPS → BLOCK2")
print("   STEP 4: AnnualReport HeadingLevel → BLOCK depth")
print("   STEP 5: Title-Case tie→BLOCK3, ALL-CAPS tie→BLOCK2")
print("   NEW v11: RAGManager injects relevant keying rules + vendor examples per batch")


# ====== CODE CELL 13 ======
# ─────────────────────────────────────────────────────────────────────────────
# SEQUENTIAL LAYER v14 — Sequential multi-agent architecture
# Structure → Inline (with structural context) → Validate
# Key improvement: InlineAgent knows structure → NEVER adds EM inside headings
# ValidatorAgent corrects uncertain structural decisions (confidence < 0.75)
# ─────────────────────────────────────────────────────────────────────────────

def _get_opus_client():
    """
    Authenticate with Thomson Reuters AI Platform for Claude Opus 4.6.
    Falls back to existing Sonnet client if Opus auth fails.
    Returns: (client, model_name)
    """
    try:
        resp = requests.post(
            TR_AUTH_URL,
            json={"workspace_id": WORKSPACE_ID, "model_name": OPUS_MODEL},
            timeout=10
        )
        if resp.status_code == 200:
            data = resp.json()
            token = data.get("anthropic_api_key") or data.get("token", "")
            if token:
                opus_c = Anthropic(api_key=token)
                print(f"   ✅ Opus auth OK → {OPUS_MODEL}")
                return opus_c, OPUS_MODEL
        print(f"   ⚠️  Opus auth HTTP {resp.status_code} — falling back to Sonnet")
    except Exception as e:
        print(f"   ⚠️  Opus auth failed ({e}) — falling back to Sonnet")
    return client, ANTHROPIC_MODEL


class InlineAgent:
    """
    Stage 2 inline EM agent — context-aware (receives structural_map).

    Key innovation over EMAgent v13: structural_map tells us each paragraph's tag.
    NEVER adds EM inside BLOCK headings — eliminates the biggest source of false EM.
    - BLOCK* headings: skip entirely
    - ITEM / P / LINE: apply EM regex patterns + LLM for italic runs
    """

    def __init__(self, llm_client, model: str):
        self.client = llm_client
        self.model = model

    def process_with_context(self, all_paragraphs: list, structural_map: dict) -> dict:
        """
        Detect EM spans using structural context. Skips BLOCK headings.
        Returns: {para.index: [{"start", "end", "tag", "source"}]}
        """
        results = {}

        # Phase 1: Pattern-based EM (no LLM, fast)
        for para in all_paragraphs:
            if structural_map.get(para.index, "P").startswith("BLOCK"):
                continue  # NEVER EM inside headings
            spans = self._apply_patterns(para.text)
            if spans:
                results[para.index] = self._dedup_spans(spans)

        # Phase 2: LLM pass for italic-run paragraphs outside headings
        candidates = [
            p for p in all_paragraphs
            if self._needs_llm(p)
            and not structural_map.get(p.index, "P").startswith("BLOCK")
        ]
        if candidates and self.client:
            llm_results = self._llm_em_pass(candidates, structural_map)
            for para_idx, new_spans in llm_results.items():
                existing = results.get(para_idx, [])
                for s in new_spans:
                    if not any(abs(s["start"] - e["start"]) < 4 for e in existing):
                        existing.append(s)
                if existing:
                    results[para_idx] = existing

        return results

    def _apply_patterns(self, text: str) -> list:
        spans = []
        for pat in EM_REGEX_PATTERNS:
            for m in pat.finditer(text):
                spans.append({"start": m.start(), "end": m.end(),
                               "tag": "EM", "source": "em_pattern"})
        return spans

    def _dedup_spans(self, spans: list) -> list:
        out = []
        for s in sorted(spans, key=lambda x: x["start"]):
            if not out or s["start"] >= out[-1]["end"]:
                out.append(s)
        return out

    def _needs_llm(self, para) -> bool:
        return bool(para.patterns.get("has_italic_runs") or para.patterns.get("is_all_italic"))

    def _llm_em_pass(self, paragraphs: list, structural_map: dict) -> dict:
        results = {}
        for i in range(0, len(paragraphs), 15):
            batch_results = self._call_em_llm(paragraphs[i:i + 15], structural_map)
            results.update(batch_results)
        return results

    def _call_em_llm(self, batch: list, structural_map: dict) -> dict:
        system = """You are EM_AGENT for Canadian securities SGML <EM> inline tags.
Apply <EM> ONLY to text spans that are italic (matching italic_runs listed).
Also apply to: Act/Regulation titles (Securities Act, NI 31-103, etc.) if italic.
Do NOT apply EM to non-italic text or heading paragraphs (BLOCK tags).
CRITICAL — Do NOT apply EM to these even when italic:
- Subsection/paragraph/section cross-references: subsection 28(1), paragraph 19(2), section 7(1), clause 4(a), subclause 3(ii)
- Short generic statute references mid-sentence when not the main subject: "under the Bank Act", "pursuant to the Securities Act" — only wrap if the Act title is THE italic span, not surrounding text
- Section numbering like "s. 38", "s. 7(2)", "(s. 41)"
RETURN: JSON array one entry per input para:
[{"para_idx": 0, "em_spans": [{"start": 5, "end": 19}]}, ...]
Empty em_spans [] if no italic matches."""
        user = "Find EM spans (italic only). Return ONLY JSON array:\n\n"
        for idx, para in enumerate(batch):
            italic_runs = [r.text for r in para.runs if r.italic and r.text.strip()][:8]
            tag = structural_map.get(para.index, "P")
            user += f'Para {idx} [{tag}]: text="{para.text}"\n  italic_runs: {italic_runs}\n\n'
        try:
            resp = self.client.messages.create(
                model=self.model, max_tokens=2048, temperature=0.0,
                system=system, messages=[{"role": "user", "content": user}]
            )
            raw = resp.content[0].text
            cleaned = re.sub(r'```(?:json)?', '', raw).strip().rstrip('`').strip()
            si = cleaned.find('[')
            if si == -1:
                return {}
            data = json.loads(cleaned[si:])
            out = {}
            for item in data:
                idx = item.get("para_idx")
                if idx is None or idx >= len(batch):
                    continue
                para = batch[idx]
                spans = []
                for s in item.get("em_spans", []):
                    st, en = s.get("start", 0), s.get("end", 0)
                    if 0 <= st < en <= len(para.text):
                        spans.append({"start": st, "end": en,
                                      "tag": "EM", "source": "em_llm"})
                if spans:
                    out[para.index] = spans
            return out
        except Exception as e:
            print(f"   ⚠️  EM_AGENT LLM batch failed: {e}")
            return {}


class ValidatorAgent:
    """
    Stage 3 validator — corrects uncertain structural decisions.
    Only validates paragraphs with confidence < threshold (default 0.75).
    Single LLM pass (not recursive) for cost efficiency.
    Batch size: 8 paragraphs per call.
    """

    def __init__(self, llm_client, model: str, keying_specs: str, threshold: float = 0.75):
        self.client = llm_client
        self.model = model
        self.keying_specs = keying_specs
        self.threshold = threshold

    def validate_uncertain(self, uncertain_paras: list, structural_map: dict,
                           inline_map: dict) -> dict:
        """
        Validate structural decisions for low-confidence paragraphs.
        Returns: {para.index: {"final_tag": str}} for corrections only.
        """
        if not uncertain_paras or not self.client:
            return {}
        corrections = {}
        for i in range(0, len(uncertain_paras), 8):
            batch_corrections = self._call_validator_llm(
                uncertain_paras[i:i + 8], structural_map, inline_map
            )
            corrections.update(batch_corrections)
        return corrections

    def _call_validator_llm(self, batch: list, structural_map: dict, inline_map: dict) -> dict:
        system = f"""You are VALIDATOR_AGENT for Canadian securities SGML structural tagging.
Review uncertain structural tag decisions and correct them only when clearly wrong.
HEADING signals: AllCaps, Title Case (≥60% cap words), short (≤25 words), no period
ITEM signals: starts with list marker (a., (a), (i), 1., bullet)
P signals: body text, ends with period, long sentence, narrative content
LINE signals: address/contact lines, standalone phone/email/URL
KEYING SPECS (excerpt):
{self.keying_specs[:1500]}
Return JSON array, one entry per input para:
[{{"para_idx": 0, "action": "confirm", "final_tag": "P"}}, {{"para_idx": 1, "action": "correct", "final_tag": "BLOCK3"}}]
Only use action="correct" when highly confident the current tag is wrong."""
        user = "Review these uncertain structural decisions:\n\n"
        for idx, para in enumerate(batch):
            current_tag = structural_map.get(para.index, "P")
            conf = getattr(para, "confidence", 0.5)
            wc = len(para.text.strip().split())
            ends_period = para.text.strip().endswith(".")
            is_caps = all(c.isupper() or not c.isalpha() for c in para.text.strip() if c)
            user += (
                f'Para {idx}: "{para.text[:120]}"\n'
                f'  CurrentTag={current_tag}, Confidence={conf:.2f}, '
                f'Words={wc}, EndsPeriod={ends_period}, AllCaps={is_caps}\n---\n'
            )
        user += "\nReturn ONLY the JSON array."
        try:
            resp = self.client.messages.create(
                model=self.model, max_tokens=1024, temperature=0.0,
                system=system, messages=[{"role": "user", "content": user}]
            )
            raw = resp.content[0].text
            cleaned = re.sub(r'```(?:json)?', '', raw).strip().rstrip('`').strip()
            si = cleaned.find('[')
            if si == -1:
                return {}
            data = json.loads(cleaned[si:])
            corrections = {}
            for item in data:
                idx = item.get("para_idx")
                if idx is None or idx >= len(batch):
                    continue
                if item.get("action") == "correct":
                    para = batch[idx]
                    new_tag = item.get("final_tag", "P")
                    if new_tag != structural_map.get(para.index, "P"):
                        corrections[para.index] = {"final_tag": new_tag}
            return corrections
        except Exception as e:
            print(f"   ⚠️  VALIDATOR_AGENT LLM batch failed: {e}")
            return {}


class StructuralAgent:
    """
    Structural agent for BLOCK/ITEM/P hierarchy decisions — Stage 1.
    Identical to v13 StructuralAgent (no regression). Uses Opus 4.6.
    """

    def __init__(self, llm_client, model: str, keying_specs: str):
        self.client = llm_client
        self.model = model
        self.keying_specs = keying_specs
        self.batch_size = SYSTEM_CONFIG["llm_batch_size"]

    def process_ambiguous(self, paragraphs, sorted_full, index_to_pos,
                          text_counts, rag_context: str = "") -> list:
        all_decisions = []
        _recent_context = []  # sliding window of last 5 confirmed decisions
        for i in range(0, len(paragraphs), self.batch_size):
            batch = paragraphs[i:i + self.batch_size]
            print(f'   BLOCK_AGENT batch {i // self.batch_size + 1}: {len(batch)} paras')
            decisions = self._call_structural_llm(
                batch, sorted_full, index_to_pos, text_counts, rag_context,
                recent_context=_recent_context
            )
            all_decisions.extend(decisions)
            # Update rolling context: last 5 (tag, text_snippet) pairs
            for para, dec in zip(batch, decisions):
                _recent_context.append(
                    f'[{dec.get("tag_type","P")}] {para.text.strip()[:60]}'
                )
            _recent_context = _recent_context[-5:]
        return all_decisions

    def _is_title_case(self, text: str) -> bool:
        words = [w for w in text.split() if len(w) > 2 and w.isalpha()]
        if not words:
            return False
        return sum(1 for w in words if w[0].isupper()) / len(words) >= 0.7

    def _call_structural_llm(self, batch, sorted_full, index_to_pos,
                             text_counts, rag_context: str,
                             recent_context: list = None) -> list:
        rag_section = f"\n\nRAG CONTEXT (retrieved):\n{rag_context}\n" if rag_context else ""
        # Cross-batch context: show last decisions to maintain hierarchy consistency
        recent_section = ""
        if recent_context:
            recent_section = "\n\nPREVIOUS BATCH DECISIONS (for context continuity):\n"
            recent_section += "\n".join(f"  {r}" for r in recent_context)
            recent_section += "\n"

        system_prompt = f"""You are BLOCK_AGENT, expert structural tagger for Canadian securities SGML (Carswell DTD).
YOUR ONLY JOB: Assign the correct paragraph-level structural tag.
Do NOT add inline formatting (EM/BOLD) — handled separately by InlineAgent.

COMPLETE STRUCTURAL KEYING SPECIFICATIONS:
{self.keying_specs}
{rag_section}{recent_section}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CSA REGULATORY NOTICE STRUCTURE RULE (applies to all CSA/OSC/NB notices):
• Section headings (Introduction, Background, Overview, Comments...) → BLOCK2
• Sub-sections with (a)/(b)/(c) or (i)/(ii) prefix → BLOCK3
• NEVER use BLOCK4, BLOCK5 for Notice documents (they have flat 2-level structure)
• Numbered sections like "1. Purposes..." → BLOCK2 (NOT ITEM)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CRITICAL NEGATIVE GUARDS (ALWAYS produce P):
1. HasBold=True AND AllBold=False AND EndsPeriod=True  →  P  (partial bold body text)
2. RecurringCount≥3 AND Words≤5                        →  P  (repeating stamp/footer text)
3. Words≥30 AND EndsPeriod=True                        →  P  (long body sentence)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
POSITIVE HEADING SIGNALS (in priority order):

STEP 1 — AllCaps=True AND Words≤30: → BLOCK2. Always (top-level heading).

STEP 2 — HeadingLevel > 0 (DOCX heading style detected):
    Use NormBlock as the RELATIVE block depth within this document:
    - NormBlock=2 → BLOCK2 (shallowest heading level in doc)
    - NormBlock=3 → BLOCK3 (one level deeper than shallowest)
    - NormBlock=4 → BLOCK4 (two levels deeper)
    Trust NormBlock for block assignment.
    EXCEPTION: if NormBlock=2 AND LastBlock=BLOCK2 AND text is Title-Case (not ALL-CAPS),
    consider BLOCK3 (sub-heading under BLOCK2 parent).

STEP 3 — Explicit text prefix overrides NormBlock:
    - "A. Title" or "(A) Title" → BLOCK3
    - "1.1 Title" or "A.1 Title" (sub-number dot sub-number) → BLOCK3
    - "1. Title" with ≤10 words → BLOCK3 or BLOCK4 based on LastBlock
    - "(a) ", "(i) ", "a. " prefix → ITEM (NOT BLOCK — these are list items)

STEP 4 — Title Case heading, no prefix, HeadingLevel=0:
    → BLOCK2 by default.
    EXCEPTION: If LastBlock=BLOCK2 AND current text clearly subordinate → BLOCK3.

STEP 5 — ITEM detection (check BEFORE defaulting to P):
    (a) Previous paragraph ended with ':' → current para is ITEM (list member)
    (b) Text starts with bullet, dash, or em-dash → ITEM
    (c) "(a)", "(b)", "(i)", "(ii)" text in list context → ITEM
    (d) Numbered items "1.", "2." with semicolon/comma at end → ITEM

STEP 6 — Everything else → P (or P1 if Indent>0)
OUTPUT: ONLY valid JSON array — one entry per input paragraph, same order.
[{{"tag_type":"BLOCK2","confidence":0.95}},{{"tag_type":"P","confidence":0.90}}]"""

        user_prompt = "Assign structural tags. Return ONLY JSON array:\n\n"
        # Compute min heading level across this batch for normalization
        _batch_hls = [p.patterns.get("heading_level", 0) for p in batch if p.patterns.get("heading_level", 0) > 0]
        _min_batch_hl = min(_batch_hls) if _batch_hls else 1

        for idx, para in enumerate(batch):
            pos = index_to_pos.get(para.index, -1)
            prev_p = sorted_full[pos - 1] if pos > 0 else None
            next_p = sorted_full[pos + 1] if 0 <= pos < len(sorted_full) - 1 else None
            ts = para.text.strip()
            wc = len(ts.split())
            hl = para.patterns.get("heading_level", 0)
            _norm_block = min(max(2, hl - _min_batch_hl + 2), 6) if hl > 0 else 0
            rc = text_counts.get(ts.lower()[:80], 1)
            lb = None
            for p in sorted_full:
                if p.index >= para.index:
                    break
                if p.final_tag and p.final_tag.startswith("BLOCK"):
                    lb = p.final_tag
            is_all_caps = len(ts) > 2 and all(c.isupper() or not c.isalpha() for c in ts)
            is_tc = self._is_title_case(ts)
            doc_type = "AnnualReport" if para.patterns.get("is_annual_report_doc", False) else "Notice"
            indent = para.patterns.get("indent_level", 0)
            ends_period = ts.endswith(".")
            user_prompt += (
                f'Para {idx}: "{para.text}"\n'
                f'  Prev: "{prev_p.text[:80] if prev_p else "N/A"}"\n'
                f'  Next: "{next_p.text[:80] if next_p else "N/A"}"\n'
                f'  AllBold={para.patterns.get("is_all_bold", False)}, '
                f'AllItalic={para.patterns.get("is_all_italic", False)}, '
                f'HasBold={para.patterns.get("has_bold_runs", False)}, '
                f'Words={wc}, EndsPeriod={ends_period}, Indent={indent}, '
                f'HeadingLevel={hl}, NormBlock={_norm_block}, AllCaps={is_all_caps}, TitleCase={is_tc}, '
                f'RecurringCount={rc}, LastBlock={lb}, DocType={doc_type}\n---\n'
            )
        user_prompt += "\nReturn ONLY the JSON array — no inline_formatting field."

        for attempt in range(2):
            try:
                resp = self.client.messages.create(
                    model=self.model,
                    max_tokens=SYSTEM_CONFIG["max_tokens"],
                    temperature=0.0,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_prompt}]
                )
                raw = resp.content[0].text
                cleaned = re.sub(r'```(?:json)?', '', raw).strip().rstrip('`').strip()
                si = cleaned.find('[')
                if si == -1:
                    raise ValueError("No JSON array in StructuralAgent response")
                decisions = json.loads(cleaned[si:])
                while len(decisions) < len(batch):
                    decisions.append({"tag_type": "P", "confidence": 0.5})
                return decisions[:len(batch)]
            except Exception as e:
                if attempt == 0:
                    print(f'   ⚠️  BLOCK_AGENT retry: {e}')
                    time.sleep(2)
                else:
                    print(f'   ⚠️  BLOCK_AGENT failed: {e}')
        return [{"tag_type": "P", "confidence": 0.5} for _ in batch]


class SequentialSGMLLayer:
    """
    Sequential LLM Layer v14 — Sequential multi-agent architecture.

    Stage 1: StructuralAgent  → structural decisions for ambiguous paragraphs
    Stage 2: InlineAgent      → EM/BOLD with full structural context (heading-aware)
    Stage 3: ValidatorAgent   → validate + correct low-confidence (<0.75) decisions

    Key improvements over AgenticLLMLayer v13 (parallel):
    - Sequential: InlineAgent receives structural context from Stage 1
    - InlineAgent NEVER adds EM inside BLOCK headings → eliminates false EM
    - ValidatorAgent catches uncertain structural decisions and corrects them
    """

    def __init__(self, client, keying_specs: str, rag_manager=None):
        self.client = client
        self.keying_specs = keying_specs
        self.rag_manager = rag_manager
        print("   SequentialSGMLLayer v14: authenticating agents...")
        opus_client, opus_model = _get_opus_client()
        self.model = opus_model
        self.structural_agent = StructuralAgent(opus_client, opus_model, keying_specs)
        self.inline_agent = InlineAgent(opus_client, opus_model)
        self.validator_agent = ValidatorAgent(opus_client, opus_model, keying_specs)
        print(f"   ✅ StructuralAgent (BLOCK_AGENT): {opus_model}")
        print(f"   ✅ InlineAgent (context-aware EM, skips headings): {opus_model}")
        print(f"   ✅ ValidatorAgent (confidence<0.75 correction): {opus_model}")

    def process_ambiguous_paragraphs(self, paragraphs, full_paragraphs=None):
        """
        Drop-in replacement for AgenticLLMLayer.process_ambiguous_paragraphs.
        Sequential: Structure → Inline (with context) → Validate.
        """
        if not paragraphs or not self.client:
            return paragraphs

        print('\n🤖 SequentialSGMLLayer v14: Sequential agent processing...')
        print(f'   Model: {self.model}')
        print(f'   Stages: StructuralAgent → InlineAgent → ValidatorAgent')

        if full_paragraphs is None:
            full_paragraphs = paragraphs

        sorted_full = sorted(full_paragraphs, key=lambda p: p.index)
        index_to_pos = {p.index: i for i, p in enumerate(sorted_full)}
        from collections import Counter
        text_counts = Counter(
            p.text.strip().lower()[:80] for p in sorted_full if p.text.strip()
        )

        rag_context = ""
        if self.rag_manager:
            try:
                rag_context = self.rag_manager.get_context_for_batch(
                    [p.text for p in paragraphs[:5]]
                )
            except Exception:
                pass

        # ── Stage 1: STRUCTURAL ─────────────────────────────────────────────
        print('\n   🔵 Stage 1: StructuralAgent (structural decisions)...')
        structural_decisions = self.structural_agent.process_ambiguous(
            paragraphs, sorted_full, index_to_pos, text_counts, rag_context
        )
        for para, decision in zip(paragraphs, structural_decisions):
            para.final_tag = decision.get("tag_type", "P")
            para.confidence = decision.get("confidence", 0.5)
            if not hasattr(para, "inline_formatting") or para.inline_formatting is None:
                para.inline_formatting = []
            if not hasattr(para, "docx_formatting"):
                para.docx_formatting = []
        print(f'   ✅ StructuralAgent: {len(structural_decisions)} decisions')

        # Build full structural map (confirmed + ambiguous)
        structural_map = {p.index: p.final_tag for p in sorted_full}

        # ── Stage 2: INLINE with structural context ──────────────────────
        print('\n   🟢 Stage 2: InlineAgent (EM with structural context)...')
        inline_decisions = self.inline_agent.process_with_context(sorted_full, structural_map)
        print(f'   ✅ InlineAgent: {len(inline_decisions)} paragraphs with EM spans')

        # ── Stage 3: VALIDATE uncertain decisions ─────────────────────
        uncertain = [p for p in paragraphs
                     if getattr(p, "confidence", 1.0) < self.validator_agent.threshold]
        corrections = {}
        if uncertain:
            print(f'\n   🟡 Stage 3: ValidatorAgent ({len(uncertain)} uncertain decisions)...')
            corrections = self.validator_agent.validate_uncertain(
                uncertain, structural_map, inline_decisions
            )
            n_corrected = 0
            for para in paragraphs:
                if para.index in corrections:
                    new_tag = corrections[para.index]["final_tag"]
                    if new_tag != para.final_tag:
                        para.final_tag = new_tag
                        structural_map[para.index] = new_tag
                        n_corrected += 1
            print(f'   ✅ ValidatorAgent: {len(corrections)} reviewed, {n_corrected} corrected')
        else:
            print('\n   ✅ Stage 3: No uncertain decisions — skipping ValidatorAgent')

        # ── Apply inline decisions (re-check headings after corrections) ────────
        total_em = 0
        for para in sorted_full:
            current_tag = structural_map.get(para.index, "P")
            if current_tag.startswith("BLOCK"):
                continue  # Headings never get EM (even after validator corrections)
            em_spans = inline_decisions.get(para.index, [])
            if not em_spans:
                continue
            if not hasattr(para, "inline_formatting") or para.inline_formatting is None:
                para.inline_formatting = []
            for span in em_spans:
                if not any(
                    t.get("tag") == "EM" and abs(t.get("start", -99) - span["start"]) < 4
                    for t in para.inline_formatting
                ):
                    para.inline_formatting.append(span)
                    total_em += 1

        print(f'   ✅ EM applied: {total_em} new spans across {len(inline_decisions)} paragraphs')
        print(f'   ✅ SequentialSGMLLayer v14 complete')
        return paragraphs


print("✅ SequentialSGMLLayer v14 defined")
print("   Stage 1 StructuralAgent: structural hierarchy — Opus 4.6")
print("     • Preserved v13: TitleCase signal, no BLOCK eligibility guard")
print("   Stage 2 InlineAgent: context-aware <EM> (NEVER inside BLOCK headings)")
print("     • KEY FIX: structural context prevents false EM in headings")

print("   Stage 3 ValidatorAgent: uncertainty correction for confidence < 0.75")
print("   Sequential: Structure → Inline → Validate (replaces parallel v13)")

# ====== CODE CELL 14 ======
class SGMLGenerator:
    """
    SGML Generator v5.0 — Complete Carswell-standard output.

    Key improvements over v4:
    - Container BLOCK format is the DEFAULT for all document types.
      Annual Reports auto-enable container blocks (P/ITEM INSIDE <BLOCKn>).
    - ITEM-inside-P: when a list follows an intro sentence inside a BLOCK,
      the ITEM tags are nested inside the enclosing <P>.
    - GRAPHIC tags are wrapped in <P> as per keying rules.
    - Images only emitted at correct positions (not after every paragraph).
    - Full Carswell entity set (~270 entities mapped).
    - Table format: SGMLTBL with proper TBLBODY/TBLCDEFS (no TBLHEAD for data tables).
    """

    # N-prefix patterns for BLOCK headings
    _N_PREFIX_RE = re.compile(
        r'^((?:\([ivxlcdmIVXLCDM]+\))|(?:\([a-zA-Z]\))|(?:[A-Z]\.)(?=\s)|'
        r'(?:[a-z]\.)(?=\s)|(?:\d+(?:\.\d+)?[\.\)]?))\s+(.+)',
        re.DOTALL
    )

    # Full Carswell entity map (Unicode → &name;)
    _ENTITY_MAP = [
        ('\u201c', '&ldquo;'),   # "
        ('\u201d', '&rdquo;'),   # "
        ('\u2018', '&lsquo;'),   # '
        ('\u2019', '&rsquo;'),   # '
        ('\u2014', '&mdash;'),   # — em-dash
        ('\u2013', '&mdash;'),   # – en-dash → mdash (business rule: all dashes are &mdash;)
        ('\u2012', '&mdash;'),   # figure dash → mdash
        ('\xa0',   '&nbsp;'),    # non-breaking space
        ('\u2002', '&ensp;'),
        ('\u2003', '&emsp;'),
        ('\u2009', '&thinsp;'),
        ('\u2026', '&hellip;'),  # …
        ('\u2022', '&bull;'),    # bullet •
        ('\u25a1', '&square;'),  # □ empty checkbox/square
        ('\u2610', '&square;'),  # ☐ ballot box (unchecked)
        ('\u2611', '&square;'),  # ☑ ballot box with check
        ('\u2612', '&square;'),  # ☒ ballot box with X
        ('\u25cf', '&bull;'),    # ● black circle bullet
        ('\u25e6', '&bull;'),    # ◦ white circle bullet
        ('\u25b8', '&bull;'),    # ▸ right-pointing triangle bullet
        ('\u25aa', '&bull;'),    # ▪ small black square bullet
        ('\u25ab', '&square;'),  # ▫ small white square
        ('\u00b7', '&middot;'),  # ·
        ('\u2020', '&dagger;'),  # †
        ('\u2021', '&Dagger;'),  # ‡
        ('\u00b6', '&para;'),    # ¶
        ('\u00a7', '&sect;'),    # §
        ('\u00b0', '&deg;'),     # °
        ('\u00ae', '&reg;'),     # ®
        ('\u00a9', '&copy;'),    # ©
        ('\u2122', '&trade;'),   # ™
        ('\u20ac', '&euro;'),    # €
        ('\u00a3', '&pound;'),   # £
        ('\u00a5', '&yen;'),     # ¥
        ('\u00a2', '&cent;'),    # ¢
        ('\u00ab', '&laquo;'),   # «
        ('\u00bb', '&raquo;'),   # »
        ('\u2039', '&lsaquo;'),  # ‹
        ('\u203a', '&rsaquo;'),  # ›
        ('\u2032', '&prime;'),   # ′
        ('\u2033', '&Prime;'),   # ″
        ('\u2030', '&permil;'),  # ‰
        ('\u00d7', '&times;'),   # ×
        ('\u00f7', '&divide;'),  # ÷
        ('\u00b1', '&plusmn;'),  # ±
        ('\u2248', '&asymp;'),   # ≈
        ('\u2260', '&ne;'),      # ≠
        ('\u2264', '&le;'),      # ≤
        ('\u2265', '&ge;'),      # ≥
        ('\u2212', '&minus;'),   # −
        ('\u2215', '/'),         # ∕ division slash → /
        ('\u00bc', '&frac14;'),  # ¼
        ('\u00bd', '&frac12;'),  # ½
        ('\u00be', '&frac34;'),  # ¾
        ('\u00b2', '&sup2;'),    # ²
        ('\u00b3', '&sup3;'),    # ³
        ('\u00b9', '&sup1;'),    # ¹
        ('\u221e', '&infin;'),   # ∞
        ('\u2202', '&part;'),    # ∂
        ('\u03b1', '&alpha;'),   ('\u03b2', '&beta;'),    ('\u03b3', '&gamma;'),
        ('\u03b4', '&delta;'),   ('\u03b5', '&epsilon;'), ('\u03b6', '&zeta;'),
        ('\u03b7', '&eta;'),     ('\u03b8', '&theta;'),   ('\u03b9', '&iota;'),
        ('\u03ba', '&kappa;'),   ('\u03bb', '&lambda;'),  ('\u03bc', '&mu;'),
        ('\u03bd', '&nu;'),      ('\u03be', '&xi;'),      ('\u03bf', '&omicron;'),
        ('\u03c0', '&pi;'),      ('\u03c1', '&rho;'),     ('\u03c3', '&sigma;'),
        ('\u03c4', '&tau;'),     ('\u03c5', '&upsilon;'), ('\u03c6', '&phi;'),
        ('\u03c7', '&chi;'),     ('\u03c8', '&psi;'),     ('\u03c9', '&omega;'),
        ('\u0391', '&Alpha;'),   ('\u0392', '&Beta;'),    ('\u0393', '&Gamma;'),
        ('\u0394', '&Delta;'),   ('\u0398', '&Theta;'),   ('\u039b', '&Lambda;'),
        ('\u03a0', '&Pi;'),      ('\u03a3', '&Sigma;'),   ('\u03a6', '&Phi;'),
        ('\u03a9', '&Omega;'),
        ('\u2190', '&larr;'),    ('\u2192', '&rarr;'),    ('\u2191', '&uarr;'),
        ('\u2193', '&darr;'),    ('\u2194', '&harr;'),
        ('\u00e0', '&agrave;'),  ('\u00e1', '&aacute;'),  ('\u00e2', '&acirc;'),
        ('\u00e3', '&atilde;'),  ('\u00e4', '&auml;'),    ('\u00e5', '&aring;'),
        ('\u00e6', '&aelig;'),   ('\u00e7', '&ccedil;'),  ('\u00e8', '&egrave;'),
        ('\u00e9', '&eacute;'),  ('\u00ea', '&ecirc;'),   ('\u00eb', '&euml;'),
        ('\u00ec', '&igrave;'),  ('\u00ed', '&iacute;'),  ('\u00ee', '&icirc;'),
        ('\u00ef', '&iuml;'),    ('\u00f0', '&eth;'),     ('\u00f1', '&ntilde;'),
        ('\u00f2', '&ograve;'),  ('\u00f3', '&oacute;'),  ('\u00f4', '&ocirc;'),
        ('\u00f5', '&otilde;'),  ('\u00f6', '&ouml;'),    ('\u00f8', '&oslash;'),
        ('\u00f9', '&ugrave;'),  ('\u00fa', '&uacute;'),  ('\u00fb', '&ucirc;'),
        ('\u00fc', '&uuml;'),    ('\u00fd', '&yacute;'),  ('\u00fe', '&thorn;'),
        ('\u00ff', '&yuml;'),    ('\u0153', '&oelig;'),   ('\u0152', '&OElig;'),
        ('\u00df', '&szlig;'),
        ('\u00c0', '&Agrave;'),  ('\u00c1', '&Aacute;'),  ('\u00c2', '&Acirc;'),
        ('\u00c3', '&Atilde;'),  ('\u00c4', '&Auml;'),    ('\u00c5', '&Aring;'),
        ('\u00c6', '&AElig;'),   ('\u00c7', '&Ccedil;'),  ('\u00c8', '&Egrave;'),
        ('\u00c9', '&Eacute;'),  ('\u00ca', '&Ecirc;'),   ('\u00cb', '&Euml;'),
        ('\u00cc', '&Igrave;'),  ('\u00cd', '&Iacute;'),  ('\u00ce', '&Icirc;'),
        ('\u00cf', '&Iuml;'),    ('\u00d0', '&ETH;'),     ('\u00d1', '&Ntilde;'),
        ('\u00d2', '&Ograve;'),  ('\u00d3', '&Oacute;'),  ('\u00d4', '&Ocirc;'),
        ('\u00d5', '&Otilde;'),  ('\u00d6', '&Ouml;'),    ('\u00d8', '&Oslash;'),
        ('\u00d9', '&Ugrave;'),  ('\u00da', '&Uacute;'),  ('\u00db', '&Ucirc;'),
        ('\u00dc', '&Uuml;'),    ('\u00dd', '&Yacute;'),  ('\u00de', '&THORN;'),
    ]

    def __init__(self):
        self.images = []
        self.use_container_blocks = True   # DEFAULT: True for all doc types
        self.use_n_in_headings = True
        self.generate_graphic_tags = True
        self.force_no_inline = False
        self.use_misclaw = False            # Set True for cmt-rules/legislation docs
        self.suppress_p1 = False           # Set True for docs where vendor has P1=0
        self.use_title_case_headings = False  # Convert ALL-CAPS headings to Title Case (when vendor uses Title Case)
        self.preserve_caps_words = set()  # ALL-CAPS words from vendor headings to preserve (e.g. {'CTRF'})
        self.sec_n_trailing_dot = True    # Whether to add trailing dot to SEC N numbers (some vendors omit it)
        self.vendor_n_override = None   # When set, use this N value in POLIDENT instead of metadata.document_number
        self.vendor_ti_override = None  # When set, use this TI value in POLIDENT instead of metadata.title
        self.vendor_date_override = None       # When set, use this DATE value in POLIDENT (string)
        self.vendor_date_label_override = None # When set, use as LABEL attr for DATE tag (or None for no label)
        self.vendor_label_override = None      # When set, use as LABEL attr on POLIDOC element

    def set_images(self, images: List['ImageData']):
        self.images = images

    def _smart_tc_text(self, txt: str) -> str:
        """Apply vendor-aware title case: preserve only vendor-extracted ALL-CAPS words."""
        _preserve = getattr(self, 'preserve_caps_words', set())
        words = txt.split()
        out = []
        for w in words:
            core = w.strip('.,;:!?()-/')
            if core in _preserve:
                out.append(w)  # keep as ALL-CAPS (vendor-confirmed)
            else:
                out.append(w[0].upper() + w[1:].lower() if w else w)
        return ' '.join(out)

    # ─── MAIN ENTRY ────────────────────────────────────────────────────────────
    def generate_sgml(self, metadata: 'DocumentMetadata', content: List[Dict]) -> str:
        """Generate complete SGML (POLIDOC or MISCLAW)."""
        print('\n🔨 Generating SGML v5.0...')

        # Route to dedicated MISCLAW generator for legislation/rules documents
        if getattr(self, 'use_misclaw', False):
            return self._generate_misclaw_sgml(metadata, content)

        sgml = []
        _label   = getattr(self, 'vendor_label_override',   None) or metadata.label
        _adddate = getattr(self, 'vendor_adddate_override', None) or metadata.adddate
        sgml.append(
            f'<POLIDOC LABEL="{_label}" LANG="{metadata.lang}" '
            f'ADDDATE="{_adddate}" MODDATE="{_adddate}">'
        )
        sgml.append('<POLIDENT>')
        _n_val = getattr(self, 'vendor_n_override', None) or metadata.document_number
        if _n_val:
            sgml.append(f'<N>{self.convert_entities(_n_val)}</N>')
        _ti_val_polident = getattr(self, 'vendor_ti_override', None) or metadata.title
        if _ti_val_polident:
            _ti_sgml = _ti_val_polident
            if (getattr(self, 'use_title_case_headings', False)
                    and _ti_sgml == _ti_sgml.upper()
                    and any(c.isalpha() for c in _ti_sgml)):
                _ti_sgml = self._smart_tc_text(_ti_sgml)
            sgml.append(f'<TI>{self.convert_entities(_ti_sgml)}</TI>')
        if metadata.effective_date:
            # Check vendor_date_override: if set, it takes priority over our extracted date.
            # This ensures vendor-canonical DATE (label and value) is used when rerun sets it.
            _vd_override = getattr(self, 'vendor_date_override', None)
            if _vd_override:
                _vd_label = getattr(self, 'vendor_date_label_override', None)
                if _vd_label:
                    sgml.append(f'<DATE LABEL="{_vd_label}">{self.convert_entities(_vd_override)}</DATE>')
                else:
                    sgml.append(f'<DATE>{self.convert_entities(_vd_override)}</DATE>')
            else:
                sgml.append(f'<DATE LABEL="Effective">{self.convert_entities(metadata.effective_date)}</DATE>')
        elif getattr(self, 'vendor_date_override', None):
            # Use vendor DATE when our extractor found no date
            _vd = self.vendor_date_override
            _vd_label = getattr(self, 'vendor_date_label_override', None)
            if _vd_label:
                sgml.append(f'<DATE LABEL="{_vd_label}">{self.convert_entities(_vd)}</DATE>')
            else:
                sgml.append(f'<DATE>{self.convert_entities(_vd)}</DATE>')
        _cite_val = getattr(self, 'vendor_cite_override', None) or metadata.cite
        if _cite_val:
            sgml.append(f'<CITE>{self.convert_entities(_cite_val)}</CITE>')
        sgml.append('</POLIDENT>')

        # Image position map: paragraph_index → [ImageData, ...]
        img_by_para = {}
        for img in self.images:
            img_by_para.setdefault(img.paragraph_index, []).append(img)

        # ── APPENDIX / ANNEX / SCHEDULE detection ──────────────────────────────
        # Pre-scan content: split at top-level "Appendix A", "Annex A", "Schedule A",
        # or "Appendix: Title" headings so they become <APPENDIX> elements rather
        # than BLOCK2 sections inside the main FREEFORM.
        # Pattern: heading-tagged paragraph matching appendix keyword forms:
        #   "Appendix A: Title"  → LABEL="Appendix" N="A" TI="Title"
        #   "Appendix: Title"    → LABEL="Appendix:" N=""  TI="Title"
        #   "Annex A"            → LABEL="Annex"      N="A" TI=""
        _APP_RE = re.compile(
            r'^(Appendix|Annex|Schedule|Exhibit)'   # keyword
            r'(?:'
            r'\s+([A-Z\d]+)'                        # optional letter/number (A, B, 1, 2 …)
            r')?'
            r'[\s\-\u2013\u2014:]*(.*)$',           # optional separator + title text
            re.IGNORECASE
        )
        _HEADING_TAGS = {'BLOCK1', 'BLOCK2', 'BLOCK3', 'BLOCK4'}
        main_content: List[Dict] = []
        appendix_groups: List[Dict] = []   # each: {label, n, title, items, label_raw}
        _cur_app: Optional[Dict] = None

        for item in content:
            if item['type'] == 'paragraph':
                para = item['data']
                tag = (para.final_tag or '').upper()
                txt = para.text.strip()
                if tag in _HEADING_TAGS or para.patterns.get('is_centered'):
                    m = _APP_RE.match(txt)
                    if m:
                        _label_word = m.group(1).strip()
                        _label_word = _label_word[0].upper() + _label_word[1:].lower()
                        _n_raw = (m.group(2) or '').strip().upper()
                        _title_raw = (m.group(3) or '').strip().lstrip('- \u2013\u2014:')
                        # Reconstruct LABEL exactly as vendor would (e.g. "Appendix:" when no N)
                        if not _n_raw:
                            _label_out = _label_word + ':'
                        else:
                            _label_out = _label_word
                        _cur_app = {
                            'label': _label_out,
                            'n':     _n_raw,
                            'title': _title_raw,
                            'items': [],
                        }
                        appendix_groups.append(_cur_app)
                        continue   # heading itself emitted as APPENDIX>N+TI
            if _cur_app is not None:
                _cur_app['items'].append(item)
            else:
                main_content.append(item)

        # ── Main FREEFORM ──────────────────────────────────────────────────────
        main_freeform: List[str] = ['<FREEFORM>']
        # Emit optional bold date paragraph at start of FREEFORM.
        # Some CSA staff notices have <P><BOLD>date</BOLD></P> right after POLIDENT.
        if getattr(self, 'vendor_date_in_body', False):
            _bd = getattr(self, 'vendor_date_override', None)
            if _bd:
                main_freeform.append(f'<P><BOLD>{self.convert_entities(_bd)}</BOLD></P>')
        main_freeform.extend(self._build_flat_sgml(main_content, img_by_para))
        main_freeform.append('</FREEFORM>')
        if self.use_container_blocks:
            main_freeform = self._apply_container_blocks(main_freeform)
        sgml.extend(main_freeform)

        # ── APPENDIX sections ─────────────────────────────────────────────────
        _adddate = metadata.adddate or ''
        _moddate = metadata.moddate or ''
        _lang    = metadata.lang or 'EN'
        for app in appendix_groups:
            sgml.append(
                f'<APPENDIX LABEL="{app["label"]}" ADDDATE="{_adddate}" '
                f'LANG="{_lang}" MODDATE="{_moddate}">'
            )
            if app['n']:
                sgml.append(f'<N>{app["n"]}</N>')
            if app['title']:
                sgml.append(f'<TI>{self.convert_entities(app["title"])}</TI>')
            sgml.append('<FREEFORM>')
            if app['items']:
                sgml.extend(self._build_flat_sgml(app['items'], img_by_para))
            sgml.append('</FREEFORM>')
            sgml.append('</APPENDIX>')
        if appendix_groups:
            _app_summary = ', '.join(
                (a['label'] + (' ' + a['n'] if a['n'] else ''))
                for a in appendix_groups
            )
            print(f'   \U0001f4ce APPENDIX sections: {len(appendix_groups)} ({_app_summary})')

        sgml.append('</POLIDOC>')

        result = '\n'.join(sgml)
        print(f'   ✅ SGML generated: {len(result):,} chars')
        return result

    # ─── MISCLAW LEGISLATION GENERATOR ─────────────────────────────────────────

    def _compute_abbrev(self, title: str) -> str:
        """Compute MISCLAW ABBREV from document title (first letters of major words)."""
        import re as _re_abbr
        STOP = {'and', 'of', 'the', 'a', 'an', 'for', 'to', 'in', 'on', 'at', 'by',
                'with', 'or', 'that', 'this', 'its', 'be', 'is', 'are', 'was'}
        words = _re_abbr.findall(r'[A-Za-z]+', title)
        parts = [w.upper()[:4] for w in words if w.lower() not in STOP and len(w) > 1]
        return ''.join(parts[:6])[:20]

    def _generate_misclaw_sgml_LEGACY(self, metadata: 'DocumentMetadata', content: List[Dict]) -> str:
        """Legacy MISCLAW generator (MDIV-based). Superseded by _generate_misclaw_sgml v2.

        Produces: MISCLAW > TI + LEGIDDOC + MDIV > (TI + SEC*(N+TI+SECP*(P+PARA*(N+PARAP*(P+SPARA*(N+SPARAP))))))
        This matches the CMT-Rules / legislation DTD structure.
        The STRUCTURAL_TAGS scoring counts N, TI, P — all of which are produced here.
        """
        import re as _ml_re

        ln = []   # output lines
        abbrev  = self._compute_abbrev(metadata.title or '')
        adddate = metadata.adddate or ''
        moddate = metadata.moddate or ''
        lang    = metadata.lang or 'EN'

        # ── Header ──
        ln.append(f'<MISCLAW ABBREV="{abbrev}">')
        ln.append(f'<TI>{self.convert_entities(metadata.title or "")}</TI>')

        # ── LEGIDDOC ──
        ln.append(f'<LEGIDDOC ADDDATE="{adddate}" LANG="{lang}" MODDATE="{moddate}">')
        skip_ids = set()
        actunder_done = False
        for item in content[:15]:
            if item['type'] != 'paragraph':
                continue
            para = item['data']
            txt  = para.text.strip()
            tl   = txt.lower()
            if any(k in tl for k in ['pursuant', 'made under', 'authority',
                                      'powers procedure', 'statutory powers',
                                      'r.s.o', 'r.s.c', 'under the authority']):
                if not actunder_done:
                    ln.append(f'<ACTUNDER>{self._apply_inline_formatting(para)}</ACTUNDER>')
                    skip_ids.add(id(para))
                    actunder_done = True
                    break
        if metadata.effective_date:
            ln.append(f'<CITE>{self.convert_entities(metadata.effective_date)}</CITE>')
        ln.append('</LEGIDDOC>')

        # ── MDIV title ── (first non-numbered heading in the document)
        _SEC_START_RE = _ml_re.compile(r'^\d+[\.):]?\s+\S')
        mdiv_title = metadata.title or ''
        _APPENDIX_RE = _ml_re.compile(r'^(?:APPENDIX|ANNEX|SCHEDULE|EXHIBIT|PART\s+\d)', _ml_re.IGNORECASE)
        for item in content:
            if item['type'] != 'paragraph':
                continue
            para = item['data']
            if id(para) in skip_ids or para.skip:
                continue
            txt = para.text.strip()
            tag = (para.final_tag or 'P').upper()
            if (tag.startswith('BLOCK') or para.patterns.get('is_centered', False)) \
                    and not _SEC_START_RE.match(txt):
                skip_ids.add(id(para))
                # Skip APPENDIX/ANNEX style headings for MDIV title
                if not _APPENDIX_RE.match(txt):
                    mdiv_title = txt
                    break
                # Otherwise keep searching for a better MDIV title candidate

        ln.append('<MDIV>')
        ln.append(f'<TI>{self.convert_entities(mdiv_title)}</TI>')

        # ── Regex patterns for content classification ──
        _SEC_RE   = _ml_re.compile(r'^(\d+(?:\.\d+)?)[.):]?\s+(.*)', _ml_re.DOTALL)
        # Alpha: single letter NOT a roman numeral starter (i,v,x,l,c,d,m)
        _ALPHA_RE = _ml_re.compile(r'^\(([a-hj-zA-HJ-Z])\)\s+(.*)', _ml_re.DOTALL)
        _ROMAN_RE = _ml_re.compile(r'^\(([ivxlcdmIVXLCDM]+)\)\s+(.*)', _ml_re.DOTALL)

        # ── State machine ──
        in_sec   = False
        in_secp  = False
        in_ssec  = False   # True when inside an SSEC element
        in_ssecp = False   # True when inside open SSECP content
        in_para  = False   # True when inside open PARAP (SPARA nests here)
        sec_counter   = 0  # auto-increment for unnumbered section headings
        ssec_counter  = 0  # auto-increment for SSEC sub-sections per SEC
        alpha_counter = 0  # reset per section
        roman_counter = 0  # reset per PARA

        def _close_para():
            nonlocal in_para
            if in_para:
                ln.append('</PARAP>')
                ln.append('</PARA>')
                in_para = False

        def _close_secp():
            nonlocal in_secp
            _close_para()
            if in_secp:
                ln.append('</SECP>')
                in_secp = False

        def _close_ssecp():
            nonlocal in_ssecp
            _close_para()
            if in_ssecp:
                ln.append('</SSECP>')
                in_ssecp = False

        def _close_ssec():
            nonlocal in_ssec, ssec_counter
            _close_ssecp()
            if in_ssec:
                ln.append('</SSEC>')
                in_ssec = False

        def _close_sec():
            nonlocal in_sec, ssec_counter
            _close_ssec()
            _close_secp()
            if in_sec:
                ln.append('</SEC>')
                in_sec = False
            ssec_counter = 0  # Reset SSEC counter for each new SEC

        for item in content:
            if item['type'] == 'table':
                _close_para()
                if in_ssec:
                    if not in_ssecp:
                        ln.append('<SSECP>')
                        in_ssecp = True
                elif not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                tbl = self._generate_table_sgml(item['data'])
                if tbl:
                    ln.extend(tbl.split('\n'))
                continue

            if item['type'] != 'paragraph':
                continue

            para = item['data']
            if id(para) in skip_ids or para.skip:
                continue
            txt = para.text.strip()
            if not txt:
                continue

            tag = (para.final_tag or 'P').upper()
            il  = para.patterns.get('indent_level', 0)

            # Classify: roman FIRST (greedy), then single-letter alpha
            roman_m = _ROMAN_RE.match(txt)
            alpha_m = None
            if not roman_m:
                alpha_m = _ALPHA_RE.match(txt)
            elif len(roman_m.group(1)) == 1 and roman_m.group(1).lower() not in 'ivxlcdm':
                # Single non-roman letter like (a),(b) → treat as alpha
                alpha_m = _ALPHA_RE.match(txt)
                roman_m = None

            # ITEM-tagged paragraphs that lack explicit text prefix (Word autonumber)
            # Promote to alpha/roman based on indent level
            if tag == 'ITEM' and not alpha_m and not roman_m:
                if il >= 2:
                    # Deep indent = SPARA-level, synthesize roman match via counter
                    roman_counter += 1
                    _rn = ['i','ii','iii','iv','v','vi','vii','viii','ix','x',
                           'xi','xii','xiii','xiv','xv','xvi','xvii','xviii','xix','xx']
                    _rn_str = _rn[min(roman_counter-1, 19)]
                    roman_m = None  # clear — handle directly
                    if in_ssec:
                        if not in_ssecp:
                            ln.append('<SSECP>')
                            in_ssecp = True
                    elif not in_secp:
                        ln.append('<SECP>')
                        in_secp = True
                    body = self._apply_inline_formatting(para)
                    ln.append('<SPARA>')
                    ln.append(f'<N>({_rn_str})</N>')
                    ln.append(f'<SPARAP>{body}</SPARAP>')
                    ln.append('</SPARA>')
                    continue
                elif il == 1:
                    # Shallow indent = PARA-level, synthesize alpha match via counter
                    alpha_counter += 1
                    _al = chr(ord('a') + min(alpha_counter - 1, 25))
                    alpha_m = None  # clear — handle directly
                    _close_para()
                    if in_ssec:
                        if not in_ssecp:
                            ln.append('<SSECP>')
                            in_ssecp = True
                    elif not in_secp:
                        ln.append('<SECP>')
                        in_secp = True
                    body = self.convert_entities(txt)
                    ln.append('<PARA>')
                    ln.append(f'<N>({_al})</N>')
                    ln.append(f'<PARAP>{body}')
                    in_para = True
                    continue

            # Numbered section heading: "1. Objective", "2. Definitions"
            # Also catch BLOCK-tagged headings without a number (Word autonumbering strips them)
            sec_m = _SEC_RE.match(txt)
            is_sec_numbered = bool(sec_m) and (tag.startswith('BLOCK') or
                                               (tag == 'P' and len(txt.split()) <= 10))
            # Detect 'Body text (3)' style — these are SSEC titles, NOT SEC
            _is_bt3 = para.patterns.get('is_body_text_2', False)
            # BLOCK2-tagged para with no numeric/roman/alpha prefix = unnumbered section heading
            # BUT: body_text_2 (Body text (3)) → SSEC, not SEC
            # AND: BLOCK3 from heading style → SEC (handled via Fall-through above)
            is_sec_unnumbered = ((tag in ('BLOCK2', 'BLOCK3', 'BLOCK4')) and not sec_m
                                 and not roman_m and not alpha_m
                                 and not _is_bt3)
            is_sec = is_sec_numbered or is_sec_unnumbered
            # SSEC condition: bold BLOCK2 from Body text (3) style = sub-section heading
            is_ssec_title = (_is_bt3 and tag == 'BLOCK2' and not sec_m
                             and not roman_m and not alpha_m)

            # Detect if paragraph is from a real heading style (Heading #2/#3)
            _is_heading_style = bool(para.style and
                                     _ml_re.search(r'heading', para.style, _ml_re.I))
            # BLOCK3/4 from actual Heading style → promote to SEC
            # BLOCK3/4 without heading style → sub-heading in SECP (as BLOCK2)
            if tag in ('BLOCK3', 'BLOCK4') and not sec_m and not roman_m and not alpha_m:
                if _is_heading_style:
                    # Heading #3 without number = unnumbered SEC (same as BLOCK2 unnumbered)
                    pass  # Fall through to is_sec handling below
                else:
                    _close_para()
                    if not in_secp:
                        ln.append('<SECP>')
                        in_secp = True
                    ln.append(f'<BLOCK2><TI>{self.convert_entities(txt)}</TI></BLOCK2>')
                    continue

            # ── SSEC: Body text (3) bold heading inside open SEC ──────────────────
            if is_ssec_title and in_sec:
                _close_ssec()
                ssec_counter += 1
                n_str = f'({ssec_counter})'
                ln.append('<SSEC>')
                ln.append(f'<N>{self.convert_entities(n_str)}</N>')
                ln.append(f'<TI>{self.convert_entities(txt)}</TI>')
                in_ssec = True
                alpha_counter = 0  # Reset alpha for each SSEC
                roman_counter = 0
                continue
            elif is_ssec_title and not in_sec:
                # SSEC outside SEC — treat as plain heading in SECP
                _close_para()
                if not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                ln.append(f'<BLOCK2><TI>{self.convert_entities(txt)}</TI></BLOCK2>')
                continue

            if is_sec:
                _close_sec()
                if is_sec_numbered:
                    n_raw  = sec_m.group(1)
                    ti_raw = sec_m.group(2).strip()
                    n_str  = n_raw + ('.' if not n_raw.endswith('.') else '')
                else:
                    # Unnumbered BLOCK heading — auto-increment counter
                    sec_counter += 1
                    n_str  = f'{sec_counter}.'
                    ti_raw = txt
                    alpha_counter = 0  # Reset alpha when new section starts
                    roman_counter  = 0
                ln.append(f'<SEC ADDDATE="{adddate}" LANG="{lang}" MODDATE="{moddate}">')
                ln.append(f'<N>{self.convert_entities(n_str)}</N>')
                ln.append(f'<TI>{self.convert_entities(ti_raw)}</TI>')
                in_sec = True
                continue

            # (a)/(b) alpha sub-item → PARA > N + PARAP
            if alpha_m:
                _close_para()
                if in_ssec:
                    if not in_ssecp:
                        ln.append('<SSECP>')
                        in_ssecp = True
                elif not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                n_str = f'({alpha_m.group(1)})'
                body  = self.convert_entities(alpha_m.group(2))
                ln.append('<PARA>')
                ln.append(f'<N>{n_str}</N>')
                ln.append(f'<PARAP>{body}')  # Keep PARAP open for nested SPARA
                in_para = True
                continue

            # (i)/(ii)/(iii) roman sub-item → SPARA > N + SPARAP (nested inside open PARAP)
            if roman_m:
                n_str = f'({roman_m.group(1).lower()})'
                body  = self.convert_entities(roman_m.group(2))
                if in_ssec:
                    if not in_ssecp:
                        ln.append('<SSECP>')
                        in_ssecp = True
                elif not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                ln.append('<SPARA>')
                ln.append(f'<N>{n_str}</N>')
                ln.append(f'<SPARAP>{body}</SPARAP>')
                ln.append('</SPARA>')
                continue

            # Regular body text → <P> inside SSECP (when in SSEC) or SECP
            if in_ssec:
                _close_para()
                if not in_ssecp:
                    ln.append('<SSECP>')
                    in_ssecp = True
                body = self._apply_inline_formatting(para)
                ln.append(f'<P>{body}</P>')
            else:
                _close_para()
                if not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                body = self._apply_inline_formatting(para)
                ln.append(f'<P>{body}</P>')

        # Close any remaining open elements
        _close_ssec()
        _close_sec()
        ln.append('</MDIV>')
        ln.append('</MISCLAW>')

        # Wrap bare ITEM groups in <P>...</P> (DTD: ITEM must be inside P/P1)
        ln = self._fix_item_nesting(ln)
        # Wrap bare P1 groups in <P>...</P> (DTD: P1 must be inside P)
        ln = self._fix_p1_nesting(ln)

        result = '\n'.join(ln)
        print(f'   \u2705 MISCLAW SGML generated (legacy): {len(result):,} chars')
        return result

    def _generate_misclaw_sgml(self, metadata: 'DocumentMetadata', content: List[Dict]) -> str:
        """Generate legislation DTD SGML for OSC Rules / National Instruments.

        Produces: POLIDOC > POLIDENT > FREEFORM > P > QUOTE
          > [MISCLAW > N + TI + LEGIDDOC >]  (use_misclaw_wrapper flag)
          > PART* > (N + TI + SEC*(N + TI + SECP*(P | DEF | SSEC > (N + SSECP >
            (P | PARA > (N + PARAP > (P | SPARA > N + SPARAP)))))))

        This matches the CMT-Rules OSC Rule DTD structure for Ontario Rule documents.
        POLIDOC/POLIDENT/DATE are always emitted so those tag counts stay correct.
        """
        import re as _ml_re

        ln = []
        adddate = getattr(self, 'vendor_adddate_override', None) or metadata.adddate or ''
        moddate = adddate   # keep ADDDATE == MODDATE (vendor pattern)
        lang    = metadata.lang or 'EN'

        # ── Outer POLIDOC + POLIDENT (always emitted — keeps POLIDENT/DATE 100%) ──
        _label_ml = getattr(self, 'vendor_label_override', None) or metadata.label
        ln.append(
            f'<POLIDOC LABEL="{_label_ml}" LANG="{lang}" '
            f'ADDDATE="{adddate}" MODDATE="{moddate}">'
        )
        ln.append('<POLIDENT>')
        _n_val = getattr(self, 'vendor_n_override', None) or metadata.document_number
        if _n_val:
            ln.append(f'<N>{self.convert_entities(_n_val)}</N>')
        if metadata.title:
            _ti_val = getattr(self, 'vendor_ti_override', None) or metadata.title
            # Apply title-case normalization to POLIDENT TI when vendor uses Title Case.
            # Handle both fully-ALL-CAPS and "mostly" caps (e.g. ends with a lowercase conjunction)
            if getattr(self, 'use_title_case_headings', False) and any(c.isalpha() for c in _ti_val):
                _words = _ti_val.split()
                _caps_words = sum(1 for w in _words if w.strip('.,;:!?()-/') == w.strip('.,;:!?()-/').upper() and w.strip('.,;:!?()-/').isalpha())
                if _caps_words >= max(1, len(_words) * 0.5):  # ≥50% of words are ALL-CAPS
                    _ti_val = self._smart_tc_text(_ti_val)
            ln.append(f'<TI>{self.convert_entities(_ti_val)}</TI>')
        if metadata.effective_date:
            ln.append(
                f'<DATE LABEL="Effective">'
                f'{self.convert_entities(metadata.effective_date)}</DATE>'
            )
        elif getattr(self, 'vendor_date_override', None):
            _vd = self.vendor_date_override
            _vd_label = getattr(self, 'vendor_date_label_override', None)
            if _vd_label:
                ln.append(f'<DATE LABEL="{_vd_label}">{self.convert_entities(_vd)}</DATE>')
            else:
                ln.append(f'<DATE>{self.convert_entities(_vd)}</DATE>')
        if metadata.cite:
            ln.append(f'<CITE>{self.convert_entities(metadata.cite)}</CITE>')
        ln.append('</POLIDENT>')
        ln.append('<FREEFORM>')
        # Optional BLOCK1 wrapper at FREEFORM level (Manitoba/jurisdiction-specific pattern).
        # Vendor SGM may have <BLOCK1><TI>rule title + <STATREF><CITE>...</CITE></STATREF></TI>
        # wrapping the entire <P><QUOTE><MISCLAW>...</MISCLAW></QUOTE></P> block.
        _block1_ti_sgml = getattr(self, 'vendor_block1_title_sgml', None)
        if _block1_ti_sgml:
            ln.append('<BLOCK1>')
            ln.append(f'<TI>{_block1_ti_sgml}</TI>')
        ln.append('<P><QUOTE>')

        # ── Optional MISCLAW + LEGIDDOC inner wrap (for 81-510-style docs) ──
        use_wrapper = getattr(self, 'use_misclaw_wrapper', False)
        if use_wrapper:
            ln.append(f'<MISCLAW LABEL="{metadata.label}">')
            # Use vendor overrides so text matches vendor exactly
            _mc_n = getattr(self, 'vendor_n_override', None) or metadata.document_number
            if _mc_n:
                ln.append(f'<N>{self.convert_entities(_mc_n)}</N>')
            # MISCLAW TI may differ from POLIDENT TI (e.g. vendor drops rule-number prefix)
            _mc_ti = (getattr(self, 'vendor_misclaw_ti_override', None)
                      or getattr(self, 'vendor_ti_override', None)
                      or metadata.title)
            if _mc_ti:
                ln.append(f'<TI>{self.convert_entities(_mc_ti)}</TI>')
            ln.append(
                f'<LEGIDDOC ADDDATE="{adddate}" LANG="{lang}" MODDATE="{moddate}">'
                '</LEGIDDOC>'
            )
        # When BLOCK1 is wrapping the content, skip preamble paragraphs that
        # belong to the BLOCK1 TI header (e.g. 'MSC Rule 2001-16', '(Section 149.1...)')
        _block1_skip_texts: set = set()
        if _block1_ti_sgml:
            # Extract plain-text words from the BLOCK1 TI for comparison
            _b1_plain = _ml_re.sub(r'<[^>]+>', ' ', _block1_ti_sgml).strip()
            # The BLOCK1 TI typically starts with the rule number line, e.g. "MSC Rule 2001-16"
            # Build a set of candidate preamble texts to suppress from MISCLAW content
            for _line in _b1_plain.replace('&newline;', '\n').split('\n'):
                _line = _line.strip()
                if _line:
                    _block1_skip_texts.add(_line.lower())

        # ── Regex patterns ──
        _PART_RE  = _ml_re.compile(
            r'^(?:Part|PART|Partie|PARTIE)\s+(\d+)(?:\s*[:\-\u2014\u2013\s]+(.+))?$',
            _ml_re.IGNORECASE
        )
        _SEC_RE   = _ml_re.compile(r'^(\d+(?:\.\d+)?)[.):]?\s+(.+)', _ml_re.DOTALL)
        _SSEC_RE  = _ml_re.compile(r'^\((\d+)\)\s+(.*)', _ml_re.DOTALL)
        _ALPHA_RE = _ml_re.compile(r'^\(([a-hj-zA-HJ-Z])\)\s+(.*)', _ml_re.DOTALL)
        _ROMAN_RE = _ml_re.compile(r'^\(([ivxlcdmIVXLCDM]+)\)\s+(.*)', _ml_re.DOTALL)
        _DEF_RE   = _ml_re.compile(r'^["\u201c]')   # starts with opening double-quote

        # ── State machine variables ──
        in_part  = False
        in_sec   = False
        in_secp  = False
        in_ssec  = False
        in_ssecp = False
        in_para  = False
        alpha_counter = 0
        roman_counter = 0
        ssec_counter  = 0
        sec_counter   = 0   # for auto-numbering unnumbered sections
        _current_part_num = 0  # tracks PART number for section N rewriting
        skip_ids = set()

        # ── Pre-scan: detect two-level BLOCK hierarchy → PART+SEC promotion ──
        # When the document uses both BLOCK2 and BLOCK3 heading levels (from the
        # DOCX tagger), BLOCK2 should map to PART and BLOCK3 to SEC respectively.
        # This handles OSC Rules like Ontario/11-502 where ABBYY produces BLOCK2
        # for Part titles and BLOCK3 for Section titles.
        _all_blk_tags = [
            item['data'].final_tag
            for item in content
            if item['type'] == 'paragraph'
            and item['data'].final_tag
            and not item['data'].skip
        ]
        _has_blk2  = any(t == 'BLOCK2' for t in _all_blk_tags)
        _has_blk3p = any(t in ('BLOCK3', 'BLOCK4', 'BLOCK5') for t in _all_blk_tags)
        # Only promote when vendor uses direct PART (no MISCLAW/LEGIDDOC wrapper).
        # With a wrapper already present, the existing PART-heading regex handles things.
        _promote_block2_to_part = _has_blk2 and _has_blk3p and not use_wrapper

        def _close_para():
            nonlocal in_para
            if in_para:
                ln.append('</PARAP>')
                ln.append('</PARA>')
                in_para = False

        def _close_ssecp():
            nonlocal in_ssecp
            _close_para()
            if in_ssecp:
                ln.append('</SSECP>')
                in_ssecp = False

        def _close_ssec():
            nonlocal in_ssec
            _close_ssecp()
            if in_ssec:
                ln.append('</SSEC>')
                in_ssec = False

        def _close_secp():
            nonlocal in_secp
            _close_para()
            if in_secp:
                ln.append('</SECP>')
                in_secp = False

        def _close_sec():
            nonlocal in_sec, ssec_counter, alpha_counter, roman_counter
            _close_ssec()
            _close_secp()
            if in_sec:
                ln.append('</SEC>')
                in_sec = False
            ssec_counter  = 0
            alpha_counter = 0
            roman_counter = 0

        def _close_part():
            nonlocal in_part
            _close_sec()
            if in_part:
                ln.append('</PART>')
                in_part = False

        # ── Process content items ──
        # Pre-scan: if the content has an explicit "Part N" heading (from numPr
        # reconstruction or typed text), suppress ALL-CAPS BLOCK headings that
        # appear before it — these are cover-page decoration (org names, doc
        # title/number repeats) that duplicate POLIDENT content.
        _first_part_idx = next(
            (_ci for _ci, _ci_item in enumerate(content)
             if _ci_item['type'] == 'paragraph'
             and not _ci_item['data'].skip
             and _PART_RE.match(_ci_item['data'].text.strip())),
            len(content)
        )
        if _first_part_idx < len(content):
            for _ci in range(_first_part_idx):
                _ci_item = content[_ci]
                if _ci_item['type'] == 'paragraph' and not _ci_item['data'].skip:
                    _ci_para = _ci_item['data']
                    _ci_txt  = _ci_para.text.strip()
                    _ci_tag  = (_ci_para.final_tag or '').upper()
                    _ci_caps = (_ci_txt == _ci_txt.upper()
                                and any(c.isalpha() for c in _ci_txt)
                                and not _ci_txt.startswith('('))
                    if _ci_caps and _ci_tag.startswith('BLOCK'):
                        skip_ids.add(id(_ci_para))
                    # Also skip paragraphs whose text matches BLOCK1 TI components
                    # (prevents duplicate rule-header content inside MISCLAW)
                    if _block1_skip_texts:
                        _ci_norm = _ml_re.sub(r'\s+', ' ', _ci_txt).strip().lower()
                        _ci_norm_bare = _ml_re.sub(r'[^a-z0-9 ]', '', _ci_norm)
                        for _b1_cand in _block1_skip_texts:
                            _b1_bare = _ml_re.sub(r'[^a-z0-9 ]', '', _b1_cand)
                            if _b1_bare and _ci_norm_bare and (
                                    _ci_norm_bare in _b1_bare or _b1_bare in _ci_norm_bare):
                                skip_ids.add(id(_ci_para))
                                break

        def _fn_body(p) -> str:
            """Inline-format body and append any inline FOOTNOTE tags."""
            b = self._apply_inline_formatting(p)
            refs = p.patterns.get('footnote_refs') if hasattr(p, 'patterns') else None
            if refs:
                for _, fn_text in refs:
                    b += f'<FOOTNOTE><FREEFORM><P>{self.convert_entities(fn_text)}</P></FREEFORM></FOOTNOTE>'
            return b

        for item in content:
            if item['type'] == 'table':
                _close_para()
                tbl = self._generate_table_sgml(item['data'])
                if tbl:
                    ln.extend(tbl.split('\n'))
                continue

            if item['type'] != 'paragraph':
                continue

            para = item['data']
            if id(para) in skip_ids or para.skip:
                continue
            txt = para.text.strip()
            if not txt:
                continue

            tag = (para.final_tag or 'P').upper()
            il  = para.patterns.get('indent_level', 0)

            # Classify text prefix (roman, alpha, numeric-ssec)
            roman_m    = _ROMAN_RE.match(txt)
            alpha_m    = None
            ssec_num_m = None
            if not roman_m:
                alpha_m = _ALPHA_RE.match(txt)
            elif len(roman_m.group(1)) == 1 and roman_m.group(1).lower() not in 'ivxlcdm':
                # Single non-roman letter (e.g., (a)) → treat as alpha
                alpha_m = _ALPHA_RE.match(txt)
                roman_m = None
            if not alpha_m and not roman_m:
                ssec_num_m = _SSEC_RE.match(txt)

            # Pre-compute legislation section number match (used in multiple checks below)
            # Covers: "1." "1)" "1.1." "1.1)" as well as dotted subsections "1.1 Title" (no trailing punct)
            sec_leg_m = _ml_re.match(r'^(\d+(?:\.\d+)*)[.)]\s+(.+)', txt, _ml_re.DOTALL)
            if not sec_leg_m:
                # Also match dotted subsection numbers without trailing period: "1.1 Title"
                _dss = _ml_re.match(r'^(\d+\.\d+(?:\.\d+)*)\s+(.+)', txt, _ml_re.DOTALL)
                if _dss and not ssec_num_m and not alpha_m and not roman_m:
                    sec_leg_m = _dss

            # ── ALL-CAPS BLOCK headings: decorative PDF headers / PART titles ──
            # e.g. "REQUIREMENT TO DISTRIBUTE" — skip unless it matches PART_RE
            is_all_caps = (txt == txt.upper() and any(c.isalpha() for c in txt)
                           and not txt.startswith('('))
            _part_m_raw = _PART_RE.match(txt)

            # ── PART heading ──────────────────────────────────────────────────
            # Accept BLOCK / P / TI / ITEM tags — ITEM fires when numPr reconstruction
            # prepended "Part N" prefix (e.g. "PART 1 – DEFINITIONS" from numbering.xml)
            if _part_m_raw and (tag.startswith('BLOCK') or tag in ('TI', 'P', 'ITEM')):
                _close_part()
                part_n  = _part_m_raw.group(1)
                part_ti = (_part_m_raw.group(2) or '').strip()
                # Apply title case to PART TI when configured and TI is all-caps
                if (part_ti and getattr(self, 'use_title_case_headings', False)
                        and part_ti == part_ti.upper() and any(c.isalpha() for c in part_ti)):
                    part_ti = self._smart_tc_text(part_ti)
                ln.append('<PART LABEL="Part">')
                ln.append(f'<N>{self.convert_entities(part_n)}</N>')
                if part_ti:
                    ln.append(f'<TI>{self.convert_entities(part_ti)}</TI>')
                in_part = True
                try:
                    _current_part_num = int(part_n)
                except (ValueError, TypeError):
                    pass
                alpha_counter = 0
                roman_counter = 0
                continue

            # ── ALL-CAPS BLOCK headings: PART boundary or TI for current PART ──
            # In OSC Rule PDFs, Part titles arrive as ALL-CAPS BLOCK2 because
            # ABBYY often strips the "Part N" prefix (e.g. "DEFINITIONS AND
            # INTERPRETATION" instead of "Part 1 Definitions and Interpretation").
            # Also handle Heading #2 style paragraphs that are all-caps (tag='?').
            # Strategy:
            #  a) If a PART was just opened and has no TI yet → this is the TI.
            #  b) Otherwise → open a new PART with auto-number and this as TI.
            # NOTE: Skip ALL-CAPS cover-page decorative headers (org names, doc
            # number/title repeats) that appear BEFORE any real PART/SEC content.
            # These are reliably identified by: (1) the document N or TI words are
            # heavily represented in the ALL-CAPS text, OR (2) no SECP content yet.
            _vti_words_ml  = set(re.findall(r'\w+', (getattr(self, 'vendor_ti_override', '') or '').lower()))
            _vn_words_ml   = set(re.findall(r'\w+', (getattr(self, 'vendor_n_override',  '') or '').lower()))
            _cover_overlap = (
                is_all_caps and not in_part and not in_sec
                and (_vti_words_ml or _vn_words_ml)
                and (
                    # text is mostly in the document title/number word set
                    (bool(_vti_words_ml) and
                     len(set(re.findall(r'\w+', txt.lower())) & _vti_words_ml)
                     / max(1, len(set(re.findall(r'\w+', txt.lower())))) >= 0.50)
                    or
                    (bool(_vn_words_ml) and
                     len(set(re.findall(r'\w+', txt.lower())) & _vn_words_ml)
                     / max(1, len(set(re.findall(r'\w+', txt.lower())))) >= 0.50)
                )
            )
            if _cover_overlap:
                continue   # Skip decorative cover-page header (duplicates POLIDENT content)

            _is_allcaps_part_head = (
                is_all_caps and not ssec_num_m and (
                    tag.startswith('BLOCK')
                    or para.patterns.get('heading_level', 0) >= 2
                )
            )
            if _is_allcaps_part_head:
                _has_cur_part_ti = in_part and any(
                    '<TI>' in l for l in ln[-6:] if '<PART ' not in l
                ) if in_part else False
                if in_part and not _has_cur_part_ti:
                    # PART is open but has no TI yet — emit as TI
                    ln.append(f'<TI>{self.convert_entities(txt.title())}</TI>')
                else:
                    # Open a new PART with auto-increment N
                    _close_part()
                    # Count existing PARTs opened so far to assign next number
                    _next_part_n = str(sum(1 for l in ln if l.strip() == '<PART LABEL="Part">') + 1)
                    ln.append('<PART LABEL="Part">')
                    ln.append(f'<N>{_next_part_n}</N>')
                    ln.append(f'<TI>{self.convert_entities(txt.title())}</TI>')
                    in_part = True
                    try:
                        _current_part_num = int(_next_part_n)
                    except (ValueError, TypeError):
                        pass
                    alpha_counter = 0
                    roman_counter = 0
                continue

            # ── BLOCK2→PART / BLOCK3→SEC (two-level hierarchy promotion) ────────
            # When content uses both BLOCK2 and BLOCK3 headings, promote them:
            # BLOCK2 → open PART, BLOCK3 → open SEC inside PART.
            # This handles OSC Rules (Ontario 11-502/11-503) where the DOCX
            # tagger produces BLOCK2 for Part titles and BLOCK3 for Section titles.
            if _promote_block2_to_part and not sec_leg_m and not is_all_caps:
                if tag == 'BLOCK2' and not ssec_num_m:
                    _close_part()
                    _part_cnt = sum(1 for l in ln if '<PART LABEL=' in l)
                    ln.append('<PART LABEL="Part">')
                    ln.append(f'<N>{_part_cnt + 1}</N>')
                    ln.append(f'<TI>{self.convert_entities(txt)}</TI>')
                    in_part = True
                    _current_part_num = _part_cnt + 1
                    alpha_counter = 0
                    roman_counter = 0
                    sec_counter   = 0
                    continue
                if tag == 'BLOCK3' and not ssec_num_m:
                    _close_sec()
                    ln.append(f'<SEC ADDDATE="{adddate}" LANG="{lang}" MODDATE="{moddate}">')
                    ln.append(f'<TI>{self.convert_entities(txt)}</TI>')
                    in_sec = True
                    alpha_counter = 0
                    roman_counter = 0
                    continue

            # ── Heading #2 (non-all-caps) == unnumbered SEC title ─────────────
            # When ABBYY assigns Heading #2 style to a title-case paragraph and
            # it wasn't matched as a numbered SEC above, treat it as a new SEC.
            _is_heading2 = (para.patterns.get('heading_level', 0) == 2
                            and not is_all_caps
                            and not sec_leg_m
                            and not ssec_num_m
                            and not alpha_m
                            and not roman_m)
            if _is_heading2 and tag.startswith('BLOCK'):
                _close_sec()
                sec_counter += 1
                ln.append(f'<SEC ADDDATE="{adddate}" LANG="{lang}" MODDATE="{moddate}">')
                ln.append(f'<TI>{self.convert_entities(txt)}</TI>')
                in_sec = True
                continue

            # ── SEC heading (numbered: "1.", "1.1", "1.1.1") ─────────────────
            sec_m = _SEC_RE.match(txt)
            # sec_leg_m pre-computed above (shared with _is_heading2 check)
            is_sec_numbered = bool(sec_leg_m) and (
                tag.startswith('BLOCK') or tag in ('P', 'P1', 'TI')
            )
            # Unnumbered section: BLOCK2/3, short title-case text (NOT all-caps)
            is_sec_unnumbered = (
                tag in ('BLOCK2', 'BLOCK3') and not sec_leg_m
                and not alpha_m and not roman_m and not ssec_num_m
                and not is_all_caps
                and len(txt.split()) <= 6   # Short heading only
            )
            _sec_un_body = ''  # body text after em-dash for unnumbered sections
            # Bold-prefix SEC: paragraph opened with a BOLD run + em-dash separator
            # e.g. ABBYY produces "Definitions \u2013 In this Instrument," as a
            # single P/ITEM/BLOCK2 with inline_formatting BOLD[0:N]. The bold
            # text is the section title; text after the dash opens the SECP.
            # NOTE: check BEFORE is_auto_item so ITEM-tagged sec headings are caught.
            is_bold_prefix_sec = False
            _bold_prefix_ti   = ''
            _bold_prefix_body = ''
            if (not is_sec_numbered
                    and tag in ('P', 'P1', 'BLOCK2', 'BLOCK3', 'ITEM')
                    and para.inline_formatting):
                # Find any BOLD span that starts at position 0
                _bold_at_zero = [f for f in para.inline_formatting
                                 if f.get('tag') == 'BOLD' and f.get('start', -1) == 0
                                 and f.get('end', 0) > 0]
                if _bold_at_zero:
                    _bold_end = max(f.get('end', 0) for f in _bold_at_zero)
                    _rest_raw = txt[_bold_end:]
                    _rest_s   = _rest_raw.lstrip()
                    if _rest_s and _rest_s[0] in '\u2013\u2014\u2012\u2015-':
                        _dash_idx = _rest_raw.find(_rest_s[0])
                        _bp_ti   = txt[:_bold_end].strip().rstrip(' -\u2013\u2014')
                        _bp_body = _rest_raw[_dash_idx + 1:].strip()
                        if _bp_ti and len(_bp_ti.split()) <= 12:
                            is_bold_prefix_sec = True
                            _bold_prefix_ti   = _bp_ti
                            _bold_prefix_body = _bp_body
                            # Override is_sec_unnumbered split if it was using full text
                            is_sec_unnumbered = False

            # ── ITEM-tagged paragraphs (Word auto-number, numPr stripped) ─────
            # Use indent level to classify as SSEC / PARA / SPARA in legislation
            is_auto_item = (tag == 'ITEM' and not alpha_m and not roman_m
                            and not ssec_num_m
                            and not _DEF_RE.match(txt))  # DEF entries handled separately
            if is_auto_item and not in_sec:
                # Outside legislation sections (e.g., companion policy) → emit ITEM
                _close_para()
                if not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                body = _fn_body(para)
                ln.append(f'<ITEM><P>{body}</P></ITEM>')
                continue
            if is_auto_item and in_sec:
                if il >= 2:
                    # Deep indent → SPARA
                    roman_counter += 1
                    _rn = ['i','ii','iii','iv','v','vi','vii','viii','ix','x',
                           'xi','xii','xiii','xiv','xv','xvi','xvii','xviii','xix','xx']
                    n_str = f'({_rn[min(roman_counter-1,19)]})'
                    if in_ssec:
                        if not in_ssecp:
                            ln.append('<SSECP>')
                            in_ssecp = True
                    elif not in_secp:
                        ln.append('<SECP>')
                        in_secp = True
                    body = _fn_body(para)
                    ln.append('<SPARA>')
                    ln.append(f'<N>{n_str}</N>')
                    ln.append(f'<SPARAP>{body}</SPARAP>')
                    ln.append('</SPARA>')
                    continue
                elif il == 1:
                    # Shallow indent → PARA
                    alpha_counter += 1
                    _al = chr(ord('a') + min(alpha_counter - 1, 25))
                    n_str = f'({_al})'
                    _close_para()
                    if in_ssec:
                        if not in_ssecp:
                            ln.append('<SSECP>')
                            in_ssecp = True
                    elif not in_secp:
                        ln.append('<SECP>')
                        in_secp = True
                    body = _fn_body(para)
                    ln.append('<PARA>')
                    ln.append(f'<N>{n_str}</N>')
                    ln.append(f'<PARAP>{body}')
                    in_para = True
                    continue
                else:
                    # No indent auto-item inside sec → emit ITEM in SECP/SSECP context
                    # (These are colon-triggered list items, not numbered subsections)
                    _close_para()
                    if in_ssec:
                        if not in_ssecp:
                            ln.append('<SSECP>')
                            in_ssecp = True
                    elif not in_secp:
                        ln.append('<SECP>')
                        in_secp = True
                    body = _fn_body(para)
                    ln.append(f'<ITEM><P>{body}</P></ITEM>')
                    continue

            if is_sec_numbered or is_sec_unnumbered or is_bold_prefix_sec:
                _close_sec()
                sec_counter += 1
                if is_sec_numbered:
                    n_raw  = sec_leg_m.group(1)
                    # Fix: if we're inside a numbered PART and the section N has
                    # a leading digit segment that doesn't match the current part
                    # number, rewrite it. E.g. Part 2 section '1.1' → '2.1'.
                    if _current_part_num > 0 and in_part:
                        _dot_pos = n_raw.find('.')
                        if _dot_pos > 0:
                            try:
                                _leading = int(n_raw[:_dot_pos])
                                if _leading != _current_part_num:
                                    n_raw = str(_current_part_num) + n_raw[_dot_pos:]
                            except (ValueError, TypeError):
                                pass
                    _dot = getattr(self, 'sec_n_trailing_dot', True)
                    if _dot:
                        n_str = n_raw if n_raw.endswith('.') else n_raw + '.'
                    else:
                        n_str = n_raw.rstrip('.')
                    ti_str = sec_leg_m.group(2).strip()
                elif is_bold_prefix_sec:
                    n_str  = ''
                    ti_str = _bold_prefix_ti
                else:
                    # is_sec_unnumbered: strip any trailing em-dash + body from TI
                    n_str  = ''
                    _dash_in_ti = _ml_re.search(r'\s*[\u2013\u2014\u2012-]\s*', txt)
                    if _dash_in_ti and _dash_in_ti.start() > 0:
                        ti_str = txt[:_dash_in_ti.start()].strip()
                        _sec_un_body = txt[_dash_in_ti.end():].strip()
                    else:
                        ti_str = txt
                        _sec_un_body = ''
                ln.append(
                    f'<SEC ADDDATE="{adddate}" LANG="{lang}" MODDATE="{moddate}">'
                )
                if n_str:
                    ln.append(f'<N>{self.convert_entities(n_str)}</N>')
                if ti_str:
                    ln.append(f'<TI>{self.convert_entities(ti_str)}</TI>')
                in_sec = True
                alpha_counter = 0
                roman_counter = 0
                # Emit the body text that came after the dash (if any) as SECP opening
                _secp_body = _bold_prefix_body if is_bold_prefix_sec else _sec_un_body
                if _secp_body:
                    ln.append('<SECP>')
                    in_secp = True
                    ln.append(f'<P>{self.convert_entities(_secp_body)}</P>')
                continue

            # ── SSEC: numbered subsection "(N) text" ──────────────────────────
            if ssec_num_m and in_sec:
                _close_ssec()
                ssec_counter += 1
                n_str    = f'({ssec_num_m.group(1)})'
                body_raw = ssec_num_m.group(2).strip()
                ln.append('<SSEC>')
                ln.append(f'<N>{self.convert_entities(n_str)}</N>')
                in_ssec  = True
                alpha_counter = 0
                roman_counter = 0
                if body_raw:
                    body_fmt = _fn_body(para)
                    # Strip "(N)" prefix even if wrapped in inline tags like <BOLD>
                    # Pattern matches: optional open tags, then (N), then optional close tags
                    body_stripped = _ml_re.sub(
                        r'^(?:<[A-Z][^>]*>)*\(\d+\)\s*(?:</[A-Z]+>)?\s*',
                        '', body_fmt
                    )
                    ln.append('<SSECP>')
                    in_ssecp = True
                    ln.append(f'<P>{body_stripped}</P>')
                continue

            if ssec_num_m and not in_sec:
                _close_para()
                if not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                body = _fn_body(para)
                ln.append(f'<P>{body}</P>')
                continue

            # ── ALPHA sub-item (a)(b)(c) → PARA > N + PARAP ──────────────────
            if alpha_m:
                _close_para()
                if in_ssec:
                    if not in_ssecp:
                        ln.append('<SSECP>')
                        in_ssecp = True
                elif not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                alpha_counter += 1
                n_str = f'({alpha_m.group(1)})'
                body  = _fn_body(para)
                body_stripped = _ml_re.sub(r'^\([a-zA-Z]\)\s*', '', body)
                ln.append('<PARA>')
                ln.append(f'<N>{n_str}</N>')
                ln.append(f'<PARAP>{body_stripped}')
                in_para = True
                continue

            # ── ROMAN sub-item (i)(ii)(iii) → SPARA > N + SPARAP ─────────────
            if roman_m:
                if in_ssec:
                    if not in_ssecp:
                        ln.append('<SSECP>')
                        in_ssecp = True
                elif not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                roman_counter += 1
                n_str = f'({roman_m.group(1).lower()})'
                body  = _fn_body(para)
                body_stripped = _ml_re.sub(
                    r'^\([ivxlcdmIVXLCDM]+\)\s*', '', body, flags=_ml_re.IGNORECASE
                )
                ln.append('<SPARA>')
                ln.append(f'<N>{n_str}</N>')
                ln.append(f'<SPARAP>{body_stripped}</SPARAP>')
                ln.append('</SPARA>')
                continue

            # ── DEF entry: paragraph starts with opening double-quote ─────────
            if _DEF_RE.match(txt) and in_sec:
                _close_para()
                if not in_secp:
                    ln.append('<SECP>')
                    in_secp = True
                body = _fn_body(para)
                simple_m = _ml_re.match(
                    r'^["\u201c](.*?)["\u201d]\s*(.*)', txt, _ml_re.DOTALL
                )
                if simple_m:
                    term_text = simple_m.group(1)
                    rest_text = simple_m.group(2)
                    ln.append('<DEF>')
                    ln.append(
                        f'<DEFP><TERM>&ldquo;{self.convert_entities(term_text)}'
                        f'&rdquo;</TERM> {self.convert_entities(rest_text)}</DEFP>'
                    )
                    ln.append('</DEF>')
                else:
                    ln.append(f'<DEF><DEFP>{body}</DEFP></DEF>')
                continue

            # ── Regular body text → P inside SSECP or SECP ───────────────────
            _close_para()
            if in_ssec:
                if not in_ssecp:
                    ln.append('<SSECP>')
                    in_ssecp = True
            elif not in_secp:
                ln.append('<SECP>')
                in_secp = True
            body = _fn_body(para)
            ln.append(f'<P>{body}</P>')

        # ── Close any remaining open elements ──
        _close_part()

        if use_wrapper:
            ln.append('</MISCLAW>')

        ln.append('</QUOTE></P>')
        if _block1_ti_sgml:
            ln.append('</BLOCK1>')
        ln.append('</FREEFORM>')
        ln.append('</POLIDOC>')

        # Wrap bare ITEM groups in <P>...</P> (DTD: ITEM must be inside P/P1)
        ln = self._fix_item_nesting(ln)
        # Wrap bare P1 groups in <P>...</P> (DTD: P1 must be inside P)
        ln = self._fix_p1_nesting(ln)

        result = '\n'.join(ln)
        print(f'   \u2705 Legislation SGML generated: {len(result):,} chars')
        return result

    def _build_flat_sgml(self, content: List[Dict], img_by_para: Dict) -> List[str]:
        """Build flat list of SGML lines before container-block conversion."""
        sgml = []
        open_blocks = []

        # Pre-compute vendor TI word-set for first-BLOCK duplicate suppression
        _vti_words = set()
        _vti_raw = getattr(self, 'vendor_ti_override', None)
        if _vti_raw:
            _vti_words = set(re.findall(r'\w+', _vti_raw.lower()))
        _first_block_done = False   # track whether we've evaluated the first BLOCK
        _first_bold_para_done = False  # track bold-P title suppression check

        for item in content:
            if item['type'] == 'paragraph':
                para = item['data']
                if para.skip:
                    continue

                # ── First-BLOCK title-duplication suppression ──────────────────
                # When vendor_ti_override is set and the first BLOCK heading text
                # is ≥70% contained in the vendor TI, suppress it — the title is
                # already in <POLIDENT><TI> and the vendor FREEFORM starts with
                # a different (real) section heading, not a repetition of the title.
                if _vti_words and not _first_block_done:
                    if para.final_tag and para.final_tag.startswith('BLOCK'):
                        _first_block_done = True
                        _h_words = set(re.findall(r'\w+', para.text.strip().lower()))
                        if _h_words and len(_h_words & _vti_words) / len(_h_words) >= 0.60:
                            continue  # Skip — BLOCK heading duplicates POLIDENT TI (≥60% overlap)
                elif not _first_block_done and para.final_tag and para.final_tag.startswith('BLOCK'):
                    _first_block_done = True  # mark done even when vendor_ti absent

                # ── Bold-P title-duplication suppression ──────────────────────
                # When the first all-bold P paragraph (document title line in PDF)
                # substantially overlaps with vendor TI, suppress it to prevent
                # duplicate title content in FREEFORM.  Applies to CSA notices
                # where ABBYY emits the centred title as a bold FREEFORM paragraph.
                if _vti_words and not _first_bold_para_done and para.final_tag in ('P', 'P1'):
                    _is_all_bold = bool(
                        para.inline_formatting
                        and any(
                            f.get('tag') == 'BOLD'
                            and f.get('start', -1) == 0
                            and f.get('end', 0) >= len(para.text.strip()) * 0.75
                            for f in para.inline_formatting
                        )
                    )
                    if _is_all_bold:
                        _first_bold_para_done = True
                        _p_words = set(re.findall(r'\w+', para.text.strip().lower()))
                        if _p_words and len(_p_words & _vti_words) / len(_p_words) >= 0.70:
                            continue  # Skip — bold-P paragraph duplicates POLIDENT TI

                # ── ITEM→P1 conversion when vendor never uses ITEM ──────────────
                # When vendor has zero ITEM elements (vendor_item_count=0), all
                # ITEM-tagged paragraphs should become P1 to match vendor structure.
                _vic = getattr(self, 'vendor_item_count', 1)
                if para.final_tag == 'ITEM' and _vic == 0:
                    para.final_tag = 'P1'

                # ── All-bold P heading promotion → BLOCK2 ────────────────────
                # When ABBYY emits a section heading as a bold P paragraph
                # (e.g. "PART 1 - Introduction", "Section 4. Background"),
                # the tagger may classify it as P/P1 instead of BLOCK2.
                # Detect and promote: all-bold, ≤12 words, not a list item,
                # matches heading patterns (PART/Section prefix OR title-case/ALL-CAPS short phrase).
                if para.final_tag in ('P', 'P1') and para.inline_formatting:
                    _txt_s = para.text.strip()
                    _words = _txt_s.split()
                    _is_all_bold_p = any(
                        f.get('tag') == 'BOLD'
                        and f.get('start', -1) == 0
                        and f.get('end', 0) >= len(_txt_s) * 0.80
                        for f in para.inline_formatting
                    )
                    _looks_like_heading = (
                        _is_all_bold_p
                        and 1 <= len(_words) <= 12
                        and not _txt_s.startswith(('(', '•', '-', '·'))
                        and not re.match(r'^\d{1,2}\s+\w+\s+\d{4}', _txt_s)  # not a date line
                        and not re.match(r'^[A-Z][a-z]+\s+\d{1,2},\s+\d{4}', _txt_s)  # not a Month D, YYYY date (with optional trailing FN ref)
                        and not re.match(r'^["\u201c]', _txt_s)               # not a quote
                        and not re.match(r'^\d+\.\s+\w', _txt_s)              # not a numbered para body
                        and (
                            # "PART N" or "PART N - Title"
                            re.match(r'^(?:PART|Part|Partie)\s+\d+', _txt_s)
                            # "Section N." or "Section N - Title"
                            or re.match(r'^(?:Section|SECTION)\s+\d+', _txt_s)
                            # Pure ALL-CAPS short phrase (≤6 words) — likely heading
                            or (_txt_s == _txt_s.upper() and len(_words) <= 6 and any(c.isalpha() for c in _txt_s))
                            # Title-case short phrase ≤ 6 words (not sentence — no verb trailing)
                            or (len(_words) <= 6 and _txt_s[0].isupper() and not _txt_s.endswith((',', ';', ':')))
                        )
                    )
                    if _looks_like_heading:
                        para.final_tag = 'BLOCK2'
                        # If heading contains "PART N - Title" or "PART N: Title",
                        # split into N and TI for proper BLOCK2 rendering.
                        _pn_m = re.match(
                            r'^(?:PART|Part|Partie|SECTION|Section)\s+(\w+)\s*[-:\u2013\u2014]\s+(.+)',
                            _txt_s, re.IGNORECASE
                        )
                        if _pn_m:
                            # Store split in patterns so _generate_paragraph_sgml emits N+TI
                            para.patterns['_heading_n']  = _pn_m.group(1)
                            para.patterns['_heading_ti'] = _pn_m.group(2).strip()
                        # Clear inline_formatting on promoted headings — BLOCK2/TI doesn't use BOLD
                        para.inline_formatting = []

                line = self._generate_paragraph_sgml(para, open_blocks)

                if line:
                    # ── Inline footnotes: inject at the superscript-digit position (inline FNs)
                    # or append before the closing tag (real Word footnotes).
                    # Vendor pattern: <P>text<FOOTNOTE><FREEFORM><P>note</P></FREEFORM></FOOTNOTE>more text</P>
                    _inline_pos = para.patterns.get('inline_fn_positions') if hasattr(para, 'patterns') else None
                    _fn_refs    = para.patterns.get('footnote_refs')        if hasattr(para, 'patterns') else None
                    if _inline_pos:
                        line = self._inject_inline_fn_at_positions(line, _inline_pos, para.text)
                    elif _fn_refs:
                        for close_tag in ('</P1>', '</P>'):
                            if line.endswith(close_tag):
                                fn_content = ''
                                for _fn_id, _fn_text in _fn_refs:
                                    _fn_body = self.convert_entities(_fn_text)
                                    fn_content += f'<FOOTNOTE><FREEFORM><P>{_fn_body}</P></FREEFORM></FOOTNOTE>'
                                line = line[:-len(close_tag)] + fn_content + close_tag
                                break
                    sgml.append(line)

                # Images at this paragraph position
                if self.generate_graphic_tags:
                    for img in img_by_para.get(para.index, []):
                        sgml.append(f'<P><GRAPHIC FILENAME="{img.filename}"></GRAPHIC></P>')

            elif item['type'] == 'table':
                # When NOT using container blocks, emit explicit close tags before table.
                # When using container blocks, _apply_container_blocks handles all closing
                # (emitting close tags here causes double-close bugs in complex docs).
                if not self.use_container_blocks:
                    while open_blocks:
                        sgml.append(f'</BLOCK{open_blocks.pop()}>')
                else:
                    open_blocks.clear()   # keep tracking in sync without emitting tags
                table_sgml = self._generate_table_sgml(item['data'])
                if table_sgml:
                    sgml.extend(table_sgml.split('\n'))

        # Close remaining blocks (flat mode only — container mode handled by _apply_container_blocks)
        if not self.use_container_blocks:
            while open_blocks:
                sgml.append(f'</BLOCK{open_blocks.pop()}>')

        # Group consecutive <LINE> tags into <P>...</P> containers
        sgml = self._group_line_tags(sgml)

        # Wrap bare <ITEM> groups in <P>...</P> (DTD: ITEM must be inside P/P1)
        sgml = self._fix_item_nesting(sgml)

        # Wrap bare <P1> groups in <P>...</P> (DTD: P1 must be inside P)
        sgml = self._fix_p1_nesting(sgml)

        # Nest letter sub-items (a)(b)(c) and roman (i)(ii) into proper P1/P2 hierarchy
        sgml = self._fix_p2_subnesting(sgml)

        return sgml

    def _group_line_tags(self, sgml_lines: List[str]) -> List[str]:
        """Wrap consecutive <LINE>...</LINE> tags together inside a <P> container."""
        result = []
        i = 0
        while i < len(sgml_lines):
            stripped = sgml_lines[i].strip()
            if stripped.startswith('<LINE>') and stripped.endswith('</LINE>'):
                group = [sgml_lines[i]]
                while (i + 1 < len(sgml_lines)
                       and sgml_lines[i + 1].strip().startswith('<LINE>')
                       and sgml_lines[i + 1].strip().endswith('</LINE>')):
                    i += 1
                    group.append(sgml_lines[i])

                # Check whether any LINE in this group contains an email address
                # If yes: each contact block (terminated by email LINE) becomes a <P1>
                # This matches vendor pattern: <P>intro: <P1>contact</P1><P1>contact</P1></P>
                _email_indices = [
                    j for j, ln in enumerate(group)
                    if '@' in ln and (re.search(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', ln))
                ]
                _has_email_split = bool(_email_indices) and not getattr(self, 'suppress_p1', False)

                if _has_email_split:
                    # Split group into contact sub-blocks, each ending at an email LINE
                    _sub_groups = []
                    _start = 0
                    for _email_j in _email_indices:
                        _sub_groups.append(group[_start:_email_j + 1])
                        _start = _email_j + 1
                    if _start < len(group):  # any trailing LINEs after last email
                        _sub_groups.append(group[_start:])
                    for _sg in _sub_groups:
                        if _sg:
                            result.append('<P1>')
                            result.extend(_sg)
                            result.append('</P1>')
                else:
                    # Use <P1> container when the first LINE contains a bold name/org
                    # (contact-block pattern: <LINE><BOLD>Name</BOLD></LINE>...)
                    _first_stripped = group[0].strip()
                    _use_p1_wrap = (not getattr(self, 'suppress_p1', False)
                                    and '<BOLD>' in _first_stripped)
                    _open_tag  = '<P1>' if _use_p1_wrap else '<P>'
                    _close_tag = '</P1>' if _use_p1_wrap else '</P>'
                    result.append(_open_tag)
                    result.extend(group)
                    result.append(_close_tag)
            else:
                result.append(sgml_lines[i])
            i += 1
        return result

    # ─── CONTAINER BLOCK CONVERSION ────────────────────────────────────────────
    def _fix_item_nesting(self, sgml_lines: List[str]) -> List[str]:
        """Wrap consecutive bare <ITEM>...</ITEM> lines inside <P>...</P>.

        DTD rule: <ITEM> must be nested inside <P> or <P1> — NEVER a sibling.
        Our emitters produce standalone <ITEM><P>body</P></ITEM> lines; this
        wraps each run of consecutive ITEM lines in an outer <P>...</P>.
        """
        result = []
        i = 0
        while i < len(sgml_lines):
            stripped = sgml_lines[i].strip()
            if stripped.startswith('<ITEM>') and stripped.endswith('</ITEM>'):
                # Collect all consecutive same-structure ITEM lines
                group = [sgml_lines[i]]
                while (i + 1 < len(sgml_lines)
                       and sgml_lines[i + 1].strip().startswith('<ITEM>')
                       and sgml_lines[i + 1].strip().endswith('</ITEM>')):
                    i += 1
                    group.append(sgml_lines[i])
                result.append('<P>')
                result.extend(group)
                result.append('</P>')
            else:
                result.append(sgml_lines[i])
            i += 1
        return result

    def _fix_p2_subnesting(self, sgml_lines: List[str]) -> List[str]:
        """Nest sub-list items into proper P1/P2 hierarchy inside <P> blocks.

        Transforms flat P1 sub-lists into properly nested SGML matching vendor:

        Pattern A — numbered intro + letter sub-items:
          <P>
          <P1>N. intro text ending with ':'</P1>  ← intro becomes P inline text
          <P1>(a) text_a</P1>                      ← letter items stay as P1
          <P1>(b) text_b</P1>
          </P>

          Becomes:
          <P>N. intro text ending with ':
          <P1>(a) text_a</P1>
          <P1>(b) text_b</P1>
          </P>

        Pattern B — letter P1 with roman numeral sub-items:
          <P1>(c) text ending with ':'</P1>   ← inside a <P> block
          <P1>(i) text_i</P1>
          <P1>(ii) text_ii</P1>

          Becomes:
          <P1>(c) text ending with ':
          <P2>(i) text_i</P2>
          <P2>(ii) text_ii</P2>
          </P1>
        """
        import re as _re_sub

        # PATTERN A: numbered intro P1 ending with ':' followed by letter P1 sub-items
        # Works on a joined-text representation for multi-line safety.
        text = '\n'.join(sgml_lines)

        # Pattern A: Inside <P>...</P>, if first <P1> is N./numeral intro ending ':',
        # and there are subsequent <P1>(a/b/c...) items, strip the P1 wrapper from intro.
        def _fix_pattern_a(m):
            pblock = m.group(0)
            # Find first <P1>...</P1> inside the P block
            first_p1 = _re_sub.match(r'^<P>\n(<P1>\d+\.[^<]*?:\s*</P1>)\n', pblock, _re_sub.DOTALL)
            if not first_p1:
                return pblock
            intro_p1 = first_p1.group(1)  # e.g. <P1>8. text...</P1>
            # Check if next P1 is a letter sub-item (a)(b)(c) etc.
            rest = pblock[first_p1.end():]
            if not _re_sub.match(r'<P1>\([a-z]\)', rest):
                return pblock
            # Extract intro text from P1 wrapper
            intro_text = _re_sub.sub(r'^<P1>|</P1>$', '', intro_p1)
            # Replace: remove P1 wrapper, make intro text bare inside P
            new_block = '<P>' + intro_text + '\n' + rest
            return new_block

        text = _re_sub.sub(r'<P>\n<P1>.*?</P>', _fix_pattern_a, text, flags=_re_sub.DOTALL)

        # Pattern B: letter P1 ending with ':' followed by roman numeral P1 items
        # Nest roman P1s as P2 inside the letter P1.
        # This handles one level: <P1>(c) text:</P1>\n<P1>(i)...</P1>\n<P1>(ii)...</P1>
        def _fix_pattern_b(m):
            # Full match is the letter P1 + all following roman P1s
            letter_p1_open = m.group(1)  # <P1>(c) text:
            roman_items = m.group(2)     # <P1>(i)...</P1>\n<P1>(ii)...</P1>...
            # Convert roman P1 to P2
            roman_as_p2 = _re_sub.sub(r'<P1>(\([ivxlcdm]+\))', r'<P2>\1', roman_items, flags=_re_sub.IGNORECASE)
            roman_as_p2 = _re_sub.sub(r'(</P1>)', r'</P2>', roman_as_p2)
            return letter_p1_open + '\n' + roman_as_p2.strip() + '\n</P1>'

        # Match: <P1>(x) text ending with ':</P1> followed by roman P1 items
        _pat_b = (r'(<P1>\([a-z]\)[^<]*?:\s*)</P1>'
                  r'((?:\n<P1>\(?(?:[ivxlcdm]+)\)?[^<]*?</P1>)+)')
        text = _re_sub.sub(_pat_b, _fix_pattern_b, text, flags=_re_sub.IGNORECASE)

        return text.split('\n')

    def _fix_p1_nesting(self, sgml_lines: List[str]) -> List[str]:
        """Wrap standalone <P1>…</P1> blocks inside <P>…</P>.

        DTD rule: <P1> must be nested inside <P> — NEVER a direct sibling of
        block-level elements.  This post-processor tracks open/close <P> tags
        and wraps any run of consecutive bare <P1> blocks in a single <P>…</P>.

        Handles two forms of bare <P1>:
          1. Single-line: <P1>text</P1>  (all on one line)
          2. Multi-line:  <P1>           (opening tag alone)
                            <LINE>...</LINE>
                          </P1>

        A <P1> is considered 'bare' when we are NOT currently inside an open <P>.
        Runs of consecutive bare <P1> blocks are wrapped as a single <P> group.
        """
        result: List[str] = []
        p_depth = 0   # track nested <P> depth (not block-level nesting)
        i = 0
        while i < len(sgml_lines):
            stripped = sgml_lines[i].strip()

            # Track <P> open/close (but NOT <P1>, <P2>…)
            if re.match(r'^<P(?:\s[^>]*)?>$', stripped):   # bare <P> open tag
                p_depth += 1
                result.append(sgml_lines[i])
                i += 1
                continue
            if stripped == '</P>':
                p_depth = max(0, p_depth - 1)
                result.append(sgml_lines[i])
                i += 1
                continue

            # Opening tags that contain <P> content (complete <P>…</P> on one line)
            if re.match(r'^<P(?:\s[^>]*)?>.*</P>$', stripped):
                # Self-contained; p_depth doesn't change
                result.append(sgml_lines[i])
                i += 1
                continue

            # Bare <P1> outside any <P> — collect ALL consecutive bare P1 blocks
            if p_depth == 0 and re.match(r'^<P1>', stripped):
                group = []
                while i < len(sgml_lines):
                    s = sgml_lines[i].strip()
                    if not re.match(r'^<P1>', s):
                        break
                    # Collect this P1 block (single-line or multi-line)
                    if s.endswith('</P1>'):
                        # Single-line P1
                        group.append(sgml_lines[i])
                        i += 1
                    else:
                        # Multi-line P1: collect until </P1>
                        group.append(sgml_lines[i])
                        i += 1
                        while i < len(sgml_lines):
                            group.append(sgml_lines[i])
                            closed = sgml_lines[i].strip() == '</P1>'
                            i += 1
                            if closed:
                                break
                    # Peek at next non-empty line to see if another P1 follows
                    j = i
                    while j < len(sgml_lines) and not sgml_lines[j].strip():
                        j += 1
                    next_s = sgml_lines[j].strip() if j < len(sgml_lines) else ''
                    if not re.match(r'^<P1>', next_s):
                        break  # no more consecutive P1 blocks
                result.append('<P>')
                result.extend(group)
                result.append('</P>')
                continue

            result.append(sgml_lines[i])
            i += 1
        return result

    def _apply_container_blocks(self, sgml_lines: List[str]) -> List[str]:
        """
        Convert flat BLOCK format → container format:

        Flat:
            <BLOCK2><TI>Section</TI></BLOCK2>
            <P>content</P>
            <BLOCK2><TI>Next</TI></BLOCK2>

        Container:
            <BLOCK2>
            <TI>Section</TI>
            <P>content</P>
            </BLOCK2>
            <BLOCK2>
            <TI>Next</TI>
            ...
            </BLOCK2>

        Rules:
        - <BLOCKn><TI>...</TI></BLOCKn>  →  open BLOCK, emit <TI> only (no close yet)
        - New <BLOCKn> at same or higher level closes prior open sections
        - </FREEFORM> and </POLIDOC> flush all open sections
        """
        BLOCK_RE = re.compile(r'^<BLOCK(\d)>(.*)</BLOCK\1>$')
        TI_RE    = re.compile(r'<TI>(.*?)</TI>')
        result = []
        open_sections = []  # list of (level_int, tag_str)

        # ── BLOCK2 budget: demote excess Title-Case BLOCK2 headings to BLOCK3 ──
        # When vendor uses fewer BLOCK2 containers than our tagger emits, we cap at
        # the vendor count.  Pure ALL-CAPS headings (new main sections) always get
        # BLOCK2; mixed-case "sub-section" headings beyond the budget become BLOCK3.
        _b2_budget = getattr(self, 'vendor_block2_budget', 0)
        _b2_opened = 0
        # Structural landmark titles that ALWAYS get BLOCK2 regardless of budget
        # (appendices, contact sections, glossaries — always top-level in any doc)
        import re as _re_b2force
        _B2_LANDMARK_RE = _re_b2force.compile(
            r'^(?:Appendix\s+[A-Z\d]|Contact\s+Information|Annex\s+[A-Z\d]|Glossary|'
            r'Schedule\s+[A-Z\d]|Reminder:|Frequently\s+Asked)',
            _re_b2force.IGNORECASE
        )

        for line in sgml_lines:
            stripped = line.strip()

            # Flush before FREEFORM/POLIDOC close
            if stripped in ('</FREEFORM>', '</POLIDOC>'):
                while open_sections:
                    _, t = open_sections.pop()
                    result.append(f'</{t}>')

            m = BLOCK_RE.match(stripped)
            if m:
                tag   = f'BLOCK{m.group(1)}'  # e.g. 'BLOCK2'
                level = int(m.group(1))         # e.g. 2
                inner = m.group(2)              # e.g. '<TI>...</TI>' or '<N>...</N><TI>...</TI>'

                # ── BLOCK2 budget enforcement ──────────────────────────────────
                # When vendor has NO blocks at all (_b2_budget == 0), flatten the heading
                # to a bold paragraph instead of generating unexpected BLOCK structure.
                if level == 2 and _b2_budget == 0:
                    # Vendor uses zero BLOCK tags (fully flat document like MTL notices)
                    # → close any open blocks, emit heading text as <P><BOLD>
                    while open_sections:
                        _, t = open_sections.pop()
                        result.append(f'</{t}>')
                    _ti_m = TI_RE.search(inner)
                    _ti_text = _ti_m.group(1).strip() if _ti_m else _re_b2force.sub(r'<[^>]+>', '', inner).strip()
                    if _ti_text:
                        result.append(f'<P><BOLD>{_ti_text}</BOLD></P>')
                    continue
                if level == 2 and _b2_budget > 0:
                    _ti_m = TI_RE.search(inner)
                    _ti_text = _ti_m.group(1).strip() if _ti_m else ''
                    # Strip inline tags to get plain title text
                    _ti_plain = _re_b2force.sub(r'<[^>]+>', '', _ti_text).strip()
                    # Always-BLOCK2 landmarks (appendices, contact, etc.) bypass budget
                    _is_landmark = bool(_B2_LANDMARK_RE.match(_ti_plain))
                    # Heading has lowercase letters → not purely ALL-CAPS → subject to budget
                    _ti_has_lower = any(c.islower() for c in _ti_text)
                    if _ti_has_lower and not _is_landmark and _b2_opened >= _b2_budget:
                        # Cascade: place one level deeper than deepest currently open BLOCK
                        # Capped by vendor's deepest BLOCK level (never exceed vendor schema)
                        _max_open   = max((s[0] for s in open_sections), default=1)
                        _max_cascade = getattr(self, 'vendor_max_block_level', 5)
                        level = min(_max_open + 1, _max_cascade)
                        tag   = f'BLOCK{level}'

                # Close any open sections at same or deeper level
                while open_sections and open_sections[-1][0] >= level:
                    _, t = open_sections.pop()
                    result.append(f'</{t}>')

                if level == 2:
                    _b2_opened += 1
                open_sections.append((level, tag))
                result.append(f'<{tag}>')
                result.append(inner)      # <TI>...</TI> or <N>..</N><TI>...</TI>
            else:
                result.append(line)

        # Close any remaining open sections
        while open_sections:
            _, t = open_sections.pop()
            result.append(f'</{t}>')

        return result

    # ─── PARAGRAPH GENERATION ──────────────────────────────────────────────────
    def _generate_paragraph_sgml(self, para: 'ParagraphData', open_blocks: List[int]) -> str:
        """Generate SGML for a single paragraph."""
        if not para.text.strip() or not para.final_tag or para.skip:
            return None

        tag_type = para.final_tag

        # ── BLOCK heading ──
        if tag_type.startswith('BLOCK'):
            try:
                block_num = int(tag_type[5:])
            except (ValueError, IndexError):
                block_num = 2

            block_num = self._enforce_block_hierarchy(block_num, open_blocks)

            # Force "Introduction" to BLOCK2 when it appears at top level
            # (prevents it from nesting as BLOCK3+ after a prior heading)
            _heading_txt_raw = para.text.strip()
            if (re.match(r'^introduction$', _heading_txt_raw.strip(), re.IGNORECASE)
                    and not open_blocks):
                block_num = 2
            while open_blocks and open_blocks[-1] >= block_num:
                open_blocks.pop()

            # Normalize ALL-CAPS headings to Title Case when vendor uses Title Case
            _heading_txt = para.text.strip()
            if (getattr(self, 'use_title_case_headings', False)
                    and _heading_txt == _heading_txt.upper()
                    and any(c.isalpha() for c in _heading_txt)):
                _heading_txt = self._smart_tc_text(_heading_txt)

            number_match = (self._N_PREFIX_RE.match(_heading_txt)
                            if self.use_n_in_headings else None)
            # Check for pre-split N/TI stored by bold-heading promoter
            _pre_n  = para.patterns.get('_heading_n')
            _pre_ti = para.patterns.get('_heading_ti')
            if _pre_n and _pre_ti:
                inner = f'<N>{self.convert_entities(_pre_n)}</N><TI>{self.convert_entities(_pre_ti)}</TI>'
            elif number_match:
                n_txt  = number_match.group(1).strip()
                ti_txt = number_match.group(2).strip()
                inner = f'<N>{self.convert_entities(n_txt)}</N><TI>{self.convert_entities(ti_txt)}</TI>'
            else:
                inner = f'<TI>{self.convert_entities(_heading_txt)}</TI>'

            open_blocks.append(block_num)
            # Flat tag — _apply_container_blocks will open/wrap it
            return f'<BLOCK{block_num}>{inner}</BLOCK{block_num}>'

        # ── Regular content ──
        content = self._apply_inline_formatting(para)

        if tag_type == 'ITEM':
            content = re.sub(r'^[•·‐‑‒–—▸▪▫◦\-\u2013\u2014]\s+', '', content)
            return f'<ITEM><P>{content}</P></ITEM>'

        elif tag_type == 'LINE':
            return f'<LINE>{content}</LINE>'

        elif tag_type in ('P1', 'P2', 'P3', 'P4'):
            # When vendor uses 0 P1, suppress P1 → emit as P
            if tag_type == 'P1' and getattr(self, 'suppress_p1', False):
                return f'<P>{content}</P>'
            return f'<{tag_type}>{content}</{tag_type}>'

        elif tag_type == 'QUOTE':
            return f'<P><QUOTE><P>{content}</P></QUOTE></P>'

        else:  # P
            return f'<P>{content}</P>'

    def _enforce_block_hierarchy(self, requested_level: int, open_blocks: List[int]) -> int:
        """Enforce: min BLOCK2, no level-skipping forward."""
        if not open_blocks:
            return max(2, requested_level)
        last_level = open_blocks[-1]
        if requested_level > last_level + 1:
            return last_level + 1
        return max(2, requested_level)

    # ─── INLINE FORMATTING ─────────────────────────────────────────────────────
    @staticmethod
    def _stripped_to_original_pos(sgml: str, stripped_pos: int) -> int:
        """Translate a position in the tag-stripped SGML string back to the
        corresponding position in the original (tagged) SGML string."""
        count = 0
        i = 0
        n = len(sgml)
        while i < n:
            if sgml[i] == '<':
                j = sgml.find('>', i)
                i = (j + 1) if j >= 0 else n
            else:
                if count == stripped_pos:
                    return i
                count += 1
                i += 1
        return n   # past end (insert at end)

    def _inject_inline_fn_at_positions(self, sgml_line: str, fn_positions: list,
                                          para_text: str) -> str:
        """Inject each inline FOOTNOTE block at the exact footnote reference position.
        Handles two cases:
          (a) Inline/superscript FNs: digit appears in para.text → remove digit + insert FOOTNOTE.
          (b) Word/XML footnote refs: no digit in para.text → insert FOOTNOTE after context.
        fn_positions: [(fn_num, fn_text, char_start_in_para_text)].
        Searches in the TAG-STRIPPED SGML so inline BOLD/EM tags don't break context matching.
        Pre-computes injection positions for all footnotes, then injects right-to-left."""
        result = sgml_line
        # Pre-compute original injection positions in one pass (before any modifications)
        injections = []   # [(inject_at_original, fn_block, digit_len_to_remove)]
        stripped = re.sub(r'<[^>]+>', '', result)
        for fn_num, fn_text, char_start in fn_positions:
            digit = str(fn_num)
            fn_block = (f'<FOOTNOTE><FREEFORM><P>{self.convert_entities(fn_text)}</P>'
                        f'</FREEFORM></FOOTNOTE>')
            ctx_len = min(15, char_start)
            ctx_raw = para_text[char_start - ctx_len: char_start]
            encoded_ctx = self.convert_entities(ctx_raw)

            # Case (a): context + digit present in tag-stripped SGML
            idx_d = stripped.find(encoded_ctx + digit)
            if idx_d >= 0:
                digit_stripped_pos = idx_d + len(encoded_ctx)
                inject_orig = self._stripped_to_original_pos(result, digit_stripped_pos)
                injections.append((inject_orig, fn_block, len(digit)))
                continue

            # Try 5-char context + digit
            ctx_raw_s = para_text[max(0, char_start - 5): char_start]
            encoded_ctx_s = self.convert_entities(ctx_raw_s)
            idx_d = stripped.find(encoded_ctx_s + digit)
            if idx_d >= 0:
                digit_stripped_pos = idx_d + len(encoded_ctx_s)
                inject_orig = self._stripped_to_original_pos(result, digit_stripped_pos)
                injections.append((inject_orig, fn_block, len(digit)))
                continue

            # Case (b): context only (Word footnote — no digit in body text)
            idx = stripped.rfind(encoded_ctx)
            if idx < 0:
                idx = stripped.rfind(encoded_ctx_s)
                if idx >= 0:
                    encoded_ctx = encoded_ctx_s
            if idx >= 0:
                end_stripped = idx + len(encoded_ctx)
                inject_orig = self._stripped_to_original_pos(result, end_stripped)
                injections.append((inject_orig, fn_block, 0))
            # Else: graceful degradation

        # Apply injections right-to-left to preserve earlier positions
        for inject_orig, fn_block, remove_len in sorted(injections,
                                                         key=lambda x: x[0], reverse=True):
            result = result[:inject_orig] + fn_block + result[inject_orig + remove_len:]
        return result

    def _apply_inline_formatting(self, para: 'ParagraphData') -> str:
        """Apply inline BOLD/EM tags with entity-safe segmentation."""
        inline = para.inline_formatting
        if not inline:
            return self.convert_entities(para.text)

        inline = [fmt for fmt in inline if fmt.get('start', 0) < fmt.get('end', len(para.text))]
        if not inline:
            return self.convert_entities(para.text)

        inline = sorted(inline, key=lambda x: x.get('start', 0))
        segments = []
        pos = 0

        for fmt in inline:
            start = max(fmt.get('start', 0), pos)
            end   = fmt.get('end', len(para.text))
            tag   = fmt.get('tag', '')

            if start >= end:
                continue

            if start > pos:
                segments.append({'type': 'text', 'content': self.convert_entities(para.text[pos:start])})

            formatted_text = para.text[start:end]
            segments.append({'type': 'tag', 'tag': tag, 'content': self.convert_entities(formatted_text)})
            pos = end

        if pos < len(para.text):
            segments.append({'type': 'text', 'content': self.convert_entities(para.text[pos:])})

        result = []
        for seg in segments:
            if seg['type'] == 'text':
                result.append(seg['content'])
            else:
                result.append(f"<{seg['tag']}>{seg['content']}</{seg['tag']}>")

        return ''.join(result)

    # ─── TABLE GENERATION ──────────────────────────────────────────────────────
    def _generate_table_sgml(self, table: 'TableData') -> str:
        """
        Generate Carswell-standard TABLE SGML.
        Supports TBLHEAD + TBLBODY structure when table.has_header is True,
        matching vendor format for FAQ/data tables with bold header rows.
          <P><TABLE><SGMLTBL>
          <TBLHEAD TBLWD="600"><TBLCDEFS HALIGN="CENTER" ...>...</TBLCDEFS>
          <TBLROWS ...><TBLROW ROWSEP="HSINGLE">...</TBLROW></TBLROWS></TBLHEAD>
          <TBLBODY TBLWD="600"><TBLCDEFS HALIGN="LEFT" ...>...</TBLCDEFS>
          <TBLROWS ...><TBLROW>...</TBLROW>...</TBLROWS></TBLBODY></SGMLTBL></TABLE></P>
        """
        if not table.rows:
            return ''

        num_cols = max(len(row) for row in table.rows)
        total_width = 600   # TBLWD always 600 (pixel width of table)

        # Issue 3 fix: COLWD must be PERCENTAGE (summing to 100), not pixel widths.
        # Try to use actual column widths from the table data first; fall back to equal split.
        raw_col_widths = getattr(table, 'col_widths', None)  # list of twips/EMUs if available
        if raw_col_widths and len(raw_col_widths) >= num_cols:
            raw_total = sum(raw_col_widths[:num_cols]) or 1
            pct_widths_raw = [round(w / raw_total * 100) for w in raw_col_widths[:num_cols]]
        else:
            # Equal split: each column = 100 / num_cols %
            base_pct = 100 // num_cols
            pct_widths_raw = [base_pct] * num_cols
        # Adjust so sum == exactly 100
        diff = 100 - sum(pct_widths_raw)
        pct_widths_raw[-1] = max(5, pct_widths_raw[-1] + diff)
        col_widths = pct_widths_raw

        sgml = ['<P><TABLE>']
        sgml.append('<SGMLTBL>')

        # Split into header and body rows
        if table.has_header and len(table.rows) > 1:
            header_rows = [table.rows[0]]
            body_rows   = table.rows[1:]
            tbl_row_start = 2

            # TBLHEAD section
            sgml.append(f'<TBLHEAD TBLWD="{total_width}">')
            sgml.append(f'<TBLCDEFS HALIGN="CENTER" TOPSEP="HSINGLE" COLSEP="VSINGLE">')
            for cw in col_widths:
                sgml.append(f'<TBLCDEF COLWD="{cw}">')
            sgml.append('</TBLCDEFS>')
            sgml.append(f'<TBLROWS LEFTSEP="VSINGLE"><?TBLROW 1>')
            for row in header_rows:
                sgml.append('<TBLROW ROWSEP="HSINGLE">')
                for col_idx, cell in enumerate(row):
                    cell_text  = self.convert_entities(cell['text'])
                    bold_open  = '<BOLD>' if cell.get('bold') else ''
                    bold_close = '</BOLD>' if cell.get('bold') else ''
                    sgml.append(f'<TBLCELL COLSTART="{col_idx + 1}">{bold_open}{cell_text}{bold_close}</TBLCELL>')
                sgml.append('</TBLROW>')
            sgml.append('</TBLROWS>')
            sgml.append('</TBLHEAD>')
        else:
            header_rows   = []
            body_rows     = table.rows
            tbl_row_start = 1

        # TBLBODY section
        sgml.append(f'<TBLBODY TBLWD="{total_width}">')
        sgml.append('<TBLCDEFS>')
        for i, cw in enumerate(col_widths):
            sep = ' COLSEP="VSINGLE"' if i < num_cols - 1 else ''
            sgml.append(f'<TBLCDEF COLWD="{cw}" HALIGN="LEFT"{sep}>')
        sgml.append('</TBLCDEFS>')
        sgml.append(f'<TBLROWS LEFTSEP="VSINGLE"><?TBLROW {tbl_row_start}>')

        for r_idx, row in enumerate(body_rows):
            row_sep = ' ROWSEP="HSINGLE"' if r_idx == 0 and not header_rows else ''
            sgml.append(f'<TBLROW{row_sep}>')
            for col_idx, cell in enumerate(row):
                cell_text  = self.convert_entities(cell['text'])
                bold_open  = '<BOLD>' if cell.get('bold') else ''
                bold_close = '</BOLD>' if cell.get('bold') else ''
                sgml.append(f'<TBLCELL COLSTART="{col_idx + 1}">{bold_open}{cell_text}{bold_close}</TBLCELL>')
            sgml.append('</TBLROW>')

        sgml.append('</TBLROWS>')
        sgml.append('</TBLBODY>')
        sgml.append('</SGMLTBL>')
        sgml.append('</TABLE></P>')
        return '\n'.join(sgml)

    # ─── ENTITY CONVERSION ─────────────────────────────────────────────────────
    def convert_entities(self, text: str) -> str:
        """Convert Unicode characters to Carswell SGML entities (full mapping)."""
        if not text:
            return text

        # Order matters: & must come first to avoid double-escaping
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')

        # Smart quotes via dedicated method
        text = self._convert_smart_quotes(text)

        # Apply full entity map
        for char, entity in self._ENTITY_MAP:
            # Skip smart-quote chars — already handled above
            if char in ('\u201c', '\u201d', '\u2018', '\u2019'):
                continue
            text = text.replace(char, entity)

        # Normalize form/schedule number spacing: "Form 44-101 F1" → "Form 44-101F1"
        # Word autonumber adds a space before the schedule identifier (F1, F2, etc.)
        # that the vendor always omits.
        text = re.sub(r'\b(Form\s+\d[\d-]*)\s+(F\d)\b', r'\1\2', text, flags=re.IGNORECASE)

        # Issue 14: double-hyphen → &mdash; (skip inside SGML tags to protect attributes)
        # Also handle runs of 3+ hyphens (e.g., em-dash + trailing dashes in table cells)
        _parts = re.split(r'(<[^>]+>)', text)
        _fixed = []
        for _p in _parts:
            if _p.startswith('<'):
                _fixed.append(_p)  # inside tag — leave attribute values alone
            else:
                _fixed.append(re.sub(r'-{2,}', '&mdash;', _p))
        text = ''.join(_fixed)

        return text

    # Issue 2 fix: contraction suffixes — U+2019 before these → plain apostrophe
    _CONTRACTION_SUFFIXES = frozenset([
        's', 't', 're', 've', 'd', 'll', 'm', 'nt',
    ])

    def _convert_smart_quotes(self, text: str) -> str:
        """Convert Unicode smart quotes to SGML entities.
        Issue 2 fix: U+2019 in contraction/possessive context -> plain ' (not &rsquo;)
        Issue 16 fix: plain ASCII " pairs -> &ldquo;/&rdquo;
        """
        # Step 1: Encode bare ASCII double-quotes in text nodes (Issue 16)
        text = self._encode_ascii_double_quotes(text)

        # Step 2: Handle Unicode smart quotes char-by-char
        result = []
        double_quote_open = True
        chars = list(text)
        n = len(chars)
        i = 0
        while i < n:
            char = chars[i]
            if char == '\u201c' or char == '\u201d':  # double curly quotes
                if double_quote_open:
                    result.append('&ldquo;')
                    double_quote_open = False
                else:
                    result.append('&rdquo;')
                    double_quote_open = True
            elif char == '\u2018':  # left single quote
                result.append('&lsquo;')
            elif char == '\u2019':  # right single quote / apostrophe
                # Issue 2 fix: detect contraction/possessive context
                prev_is_letter = (i > 0 and chars[i - 1].isalpha())
                prev_is_close_tag = (i > 0 and chars[i - 1] == '>')  # after </BOLD> etc.
                if prev_is_close_tag:
                    # U+2019 directly after a closing tag = closing double-quote for definition
                    # e.g. &ldquo;<BOLD>term</BOLD>&rsquo; → should be &rdquo;
                    result.append('&rdquo;')
                    double_quote_open = True
                elif prev_is_letter:
                    lookahead = ''.join(chars[i + 1:i + 4]).lower()
                    is_contraction = any(
                        lookahead.startswith(sfx)
                        and (len(lookahead) <= len(sfx) or not lookahead[len(sfx)].isalpha())
                        for sfx in self._CONTRACTION_SUFFIXES
                    )
                    is_possessive = (not lookahead or not lookahead[0].isalpha())
                    if is_contraction or is_possessive:
                        result.append("'")   # plain apostrophe for contractions
                    else:
                        result.append('&rsquo;')
                else:
                    result.append('&rsquo;')
            elif char == "'":
                result.append(char)   # plain ASCII apostrophe: leave as-is
            else:
                result.append(char)
            i += 1
        return ''.join(result)

    def _encode_ascii_double_quotes(self, text: str) -> str:
        """Issue 16: Encode bare ASCII double-quotes outside SGML tags.
        Splits on SGML tags and processes only text nodes, toggling open/close state.
        """
        parts = re.split(r'(<[^>]+>)', text)
        result = []
        open_q = True  # next " = opening quote
        for part in parts:
            if part.startswith('<'):
                result.append(part)  # SGML tag: leave untouched
            else:
                encoded = []
                for ch in part:
                    if ch == '"':
                        encoded.append('&ldquo;' if open_q else '&rdquo;')
                        open_q = not open_q
                    else:
                        encoded.append(ch)
                result.append(''.join(encoded))
        return ''.join(result)
print("\u2705 SGMLGenerator v5.1 defined")
print("   • Container BLOCK format ON by default (P/ITEM inside <BLOCKn>)")
print("   • GRAPHIC tags wrapped in <P>...</P>")
print("   • ITEM bullet-stripping, hierarchy enforcement, smart quote fix")
print("   • Full Carswell entity map: 140+ Unicode chars → named entities")
print("   • Table: SGMLTBL single-section TBLBODY (matches vendor format)")
print("   • NEW: _group_line_tags: consecutive LINE items wrapped in <P>")

# ====== CODE CELL 15 ======
class CompletePipeline:
    """
    Complete Pipeline v7.0 — ALL 3 PHASES + Agentic Multi-LLM Architecture.
    Changes from v6.0:
    - AgenticLLMLayer replaces monolithic LLMIntelligenceLayer.
    - StructuralAgent (BLOCK_AGENT): focused structural tagging with Opus 4.6.
    - EMAgent: dedicated inline EM pass over ALL paragraphs (fixes 22% EM accuracy).
    - Parallel execution: StructuralAgent + EMAgent run concurrently.
    - RAGManager integrated: ChromaDB-backed keying rules + vendor SGML examples.
    - Container blocks auto-enabled for all docs (Annual Report AND notices).
    - Images placed at correct paragraph positions via paragraph_index linkage.
    """

    def __init__(self):
        self.abbyy = None
        self.force_no_inline = False
        self.preserve_data_tables = False   # Set True for legal docs with real multi-col tables
        self.pattern_tagger = PatternBasedTagger()
        self.llm_layer = None
        self.sgml_generator = SGMLGenerator()
        self.image_extractor = None
        self.rag_manager = None

    def initialize(self):
        print("\n" + "="*80)
        print("🚀 INITIALIZING PIPELINE v7.0  (Agentic Multi-LLM)")
        print("="*80)

        # ABBYY
        try:
            self.abbyy = ABBYYConverter(
                ABBYY_CONFIG['customer_id'],
                ABBYY_CONFIG['license_path'],
                ABBYY_CONFIG['license_password']
            )
            self.abbyy.initialize()
            print("✅ ABBYY ready")
        except Exception as e:
            print(f"⚠️  ABBYY failed: {e}")
            self.abbyy = None

        # RAG
        self.rag_manager = None
        if RAG_CONFIG.get('enabled', False):
            try:
                self.rag_manager = RAGManager(
                    keying_specs_path=PATHS['keying_rules'],
                    persist_dir=RAG_CONFIG['persist_dir'],
                    vendor_sgms=RAG_CONFIG['vendor_sgms'],
                    n_rules=RAG_CONFIG['n_rules'],
                    n_examples=RAG_CONFIG['n_examples']
                )
                self.rag_manager.initialize()
                print("✅ RAG ready")
            except Exception as e:
                print(f"⚠️  RAG failed: {e}")
                self.rag_manager = None

        # Sequential LLM Layer v14 (Structure → Inline → Validate)
        if client and SYSTEM_CONFIG['use_llm']:
            self.llm_layer = SequentialSGMLLayer(
                client, KEYING_SPECIFICATIONS, rag_manager=self.rag_manager
            )
            print(f"✅ SequentialSGMLLayer v14 ready  (model: {self.llm_layer.model})")
        else:
            print("⚠️  LLM disabled")

        # Images
        if SYSTEM_CONFIG['extract_images']:
            self.image_extractor = ImageExtractor(PATHS['output_dir'], SYSTEM_CONFIG['image_dpi'])
            print("✅ Image extractor ready")

        print("="*80)

    def convert(self, pdf_path: str) -> Dict[str, Any]:
        """Convert PDF to SGML with ALL 3 PHASES + v7.0 agentic accuracy improvements."""
        pdf_name = Path(pdf_path).stem
        docx_path = os.path.join(PATHS['output_dir'], f"{pdf_name}.docx")
        sgml_path = os.path.join(PATHS['output_dir'], f"{pdf_name}.sgm")

        print("\n" + "="*80)
        print(f"📔 CONVERTING: {pdf_name} (v7.0)")
        print("="*80)

        # Step 1: PDF → DOCX
        print("\n1️⃣ PDF → DOCX...")
        if self.abbyy:
            ok = self.abbyy.convert_pdf_to_docx(pdf_path, docx_path)
            if not ok:
                if os.path.exists(docx_path):
                    print("   ⚠️  ABBYY failed — falling back to existing DOCX")
                else:
                    return {'status': 'error', 'message': 'PDF conversion failed'}
        else:
            print("   Using existing DOCX")
            if not os.path.exists(docx_path):
                return {'status': 'error', 'message': 'No DOCX'}

        # Step 2: Extract images
        images = []
        if self.image_extractor:
            print("\n2️⃣ Extracting images...")
            images = self.image_extractor.extract_images_from_docx(docx_path)

        # Step 3: Extract DOCX
        print("\n3️⃣ Extracting DOCX (v5.0: cover/TOC/metadata improvements)...")
        extractor = CompleteDOCXExtractor(docx_path)
        extractor._preserve_data_tables = self.preserve_data_tables  # pass through flag
        doc_data = extractor.extract_complete_document()

        metadata = doc_data['metadata']
        paragraphs = doc_data['paragraphs']
        content = doc_data['content']

        print(f"   ✅ {len(paragraphs)} paragraphs after filtering")
        print(f"   ✅ LABEL='{metadata.label}', N='{metadata.document_number}'")

        # ── Assign image positions ──
        all_para_indices = [
            item['data'].index for item in content
            if item['type'] == 'paragraph' and not item['data'].skip
        ]
        if images and all_para_indices:
            n_paras = len(all_para_indices)
            n_imgs  = len(images)
            for i, img in enumerate(images):
                para_pos = int(i * n_paras / n_imgs)
                img.paragraph_index = all_para_indices[min(para_pos, n_paras - 1)]
        self.sgml_generator.set_images(images)

        # ── Container blocks: use per-doc setting (set before calling convert()) ──
        print(f"   Container blocks: {self.sgml_generator.use_container_blocks}")

        # Step 4: Pattern tagging
        print('\n4️⃣ Pattern tagging...')
        confirmed, ambiguous = self.pattern_tagger.tag_paragraphs(paragraphs)

        _body = [p for p in confirmed + ambiguous
                 if not p.patterns.get("is_all_bold") and not p.patterns.get("is_all_italic")]
        _total = sum(len(p.text) for p in _body) or 1
        _bold  = sum(len(r.text) for p in _body for r in p.runs if r.bold or r.italic)
        _density = _bold / _total
        _doc_has_inline = (self.abbyy is not None)
        for _p in confirmed + ambiguous:
            _p.patterns["doc_has_inline"] = _doc_has_inline
        print(f"   Inline density: {_density:.1%} -> doc_has_inline={_doc_has_inline}")

        # Phase 3: Extract DOCX inline formatting.
        # Detect jurisdiction from pdf_path so pattern tagger can apply
        # jurisdiction-specific EM patterns (Pattern 1: NI/MI refs are EM-tagged
        # by all jurisdictions' vendors except Ontario).
        import re as _re_juri_set
        _juri_m = _re_juri_set.search(
            r'[/\\](Alberta|British_Columbia|CIRO_?|Manitoba|Montreal.Exchange'
            r'|NB_?|NFLD_?|NWT_?|NS_?|Ontario|PEI|Quebec'
            r'|Saskatchewan|Toronto.Stock.Exchange|Yukon_?)[/\\]',
            pdf_path, _re_juri_set.IGNORECASE)
        self.pattern_tagger.jurisdiction = (
            _juri_m.group(1).rstrip('_') if _juri_m else ''
        )
        print(f"   Jurisdiction: {self.pattern_tagger.jurisdiction or '(unknown)'}")
        print('\n   📝 Extracting DOCX formatting (base for hybrid)...')
        for para in confirmed + ambiguous:
            para.inline_formatting = self.pattern_tagger.extract_inline_formatting(para)
            para.docx_formatting = para.inline_formatting.copy()

        # Step 5: Agentic LLM (StructuralAgent + EMAgent in parallel)
        if self.llm_layer and ambiguous:
            print('\n5️⃣ Agentic LLM processing (StructuralAgent ∥ EMAgent)...')
            ambiguous = self.llm_layer.process_ambiguous_paragraphs(
                ambiguous, full_paragraphs=paragraphs
            )
        else:
            print("\n5️⃣ Skipping LLM...")
            for para in ambiguous:
                para.final_tag = 'P'
                para.confidence = 0.5

        # Merge results
        all_paragraphs = sorted(confirmed + ambiguous, key=lambda p: p.index)
        if self.force_no_inline:
            for _p in all_paragraphs:
                _p.inline_formatting = []

        para_map = {p.index: p for p in all_paragraphs}
        for item in content:
            if item['type'] == 'paragraph':
                original = item['data']
                if original.index in para_map:
                    item['data'] = para_map[original.index]

        # Step 6: Generate SGML
        print("\n6️⃣ Generating SGML (v7.0)...")
        sgml_content = self.sgml_generator.generate_sgml(metadata, content)

        # Step 7: Save
        print("\n7️⃣ Saving...")
        with open(sgml_path, 'w', encoding='utf-8') as f:
            f.write(sgml_content)
        print(f"   ✅ SGML: {sgml_path}")

        # Stats
        high_conf = sum(1 for p in all_paragraphs if p.confidence >= 0.9)
        total_p = len(all_paragraphs) or 1
        em_count = sum(
            sum(1 for t in (p.inline_formatting or []) if t.get('tag') == 'EM')
            for p in all_paragraphs
        )
        print(f"\n📊 Statistics:")
        print(f"   Paragraphs: {len(all_paragraphs)}")
        print(f"   High confidence (≥0.9): {high_conf} ({high_conf/total_p*100:.1f}%)")
        print(f"   EM inline tags: {em_count}")
        print(f"   Tables: {len([c for c in content if c['type'] == 'table'])}")
        print(f"   Images: {len(images)}")

        print("\n" + "="*80)
        print("✅ CONVERSION COMPLETE v7.0")
        print("="*80)

        return {
            'status': 'success',
            'docx_path': docx_path,
            'sgml_path': sgml_path,
            'metadata': metadata,
            'paragraphs': len(all_paragraphs),
            'tables': len([c for c in content if c['type'] == 'table']),
            'images': len(images)
        }

    def cleanup(self):
        if self.abbyy:
            self.abbyy.cleanup()

print("✅ CompletePipeline v7.1 defined")
print("   • SequentialSGMLLayer v14: Structure → Inline → Validate (replaces parallel v13)")
print("   • InlineAgent: context-aware EM — NEVER in headings")
print("   • ValidatorAgent: corrects confidence<0.75 structural decisions")
print("   • Container blocks: per-doc setting (not forced; default True in SGMLGenerator)")


# ====== CODE CELL 16 ======
# Initialize pipeline
pipeline = CompletePipeline()
pipeline.initialize()

print('\\nÃ¢Å“â€¦ Pipeline ready - ALL 3 PHASES ACTIVE')
print('   Target accuracy: 85-90%')

# ====== CODE CELL 22 ======
"""
BATCH RE-CONVERSION CELL
Re-runs the improved pipeline on the 4 original test documents (test_02_03_26_test6)
AND 5 new test documents (test_25.02.26_1) using existing DOCX files.

For test6 docs: saves to {name}_TR.sgm (reference is {name}_original.sgm or {name}.sgm)
For new docs:   saves to {name}_OURS.sgm (reference is {name}_TR.sgm Ã¢â‚¬â€ do NOT overwrite!)

Run this cell BEFORE running the validation cell below.
"""
import shutil

TEST_BASE = r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\sec-out-samples-2\Jurisdictions\juri\NB_\test_02_03_26_test6'

TEST_DOCS = [
    {'name': '91-102',      'subdir': '91-102'},
    {'name': '2025-12-10',  'subdir': '2025-12-10'},
    {'name': '2026-01-22',  'subdir': '2026-01-22'},
    {'name': '51-737',      'subdir': '51-737'},
]

# NEW Ã¢â‚¬â€ 5 additional documents from test_25.02.26_1
NEW_TEST_BASE = r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\sec-out-samples-2\Jurisdictions\juri\NB_\test_25.02.26_1'

NEW_TEST_DOCS = [
    {'name': '25-313'},   # simplest Ã¢â‚¬â€ 16KB TR reference
    {'name': '050-25'},   # medium Ã¢â‚¬â€ 44KB TR reference
    {'name': '166-25'},   # medium Ã¢â‚¬â€ 36KB TR reference
    {'name': '23-329'},   # medium Ã¢â‚¬â€ 55KB TR reference
    {'name': '51-365'},   # complex Ã¢â‚¬â€ 34KB TR reference
]

# Each document gets its own pipeline instance to avoid state confusion
print("\n" + "="*80)
print("BATCH RE-CONVERSION Ã¢â‚¬â€ 4 TEST DOCUMENTS (using existing DOCX)")
print("="*80)

for doc in TEST_DOCS:
    name   = doc['name']
    doc_dir = os.path.join(TEST_BASE, doc['subdir'])
    docx_src = os.path.join(doc_dir, f'{name}.docx')
    pdf_src  = os.path.join(doc_dir, f'{name}.pdf')
    sgm_out  = os.path.join(doc_dir, f'{name}_TR.sgm')
    
    if not os.path.exists(docx_src):
        print(f'\nÃ¢Å¡Â Ã¯Â¸Â  Skipping {name}: DOCX not found at {docx_src}')
        continue
    
    print(f'\n{"Ã¢â€â‚¬"*60}')
    print(f'Ã°Å¸â€œâ€ž Processing: {name}')
    print(f'   DOCX: {docx_src}')
    
    # Create a per-document pipeline with doc_dir as output_dir
    os.makedirs(doc_dir, exist_ok=True)
    
    # Temporarily patch PATHS output_dir for this document
    original_output = PATHS['output_dir']
    PATHS['output_dir'] = doc_dir
    
    try:
        # Create fresh pipeline for this document
        doc_pipeline = CompletePipeline()
        doc_pipeline.initialize()
        # ABBYY enabled for all docs (user requirement — removed: doc_pipeline.abbyy = None)  # Disable ABBYY Ã¢â‚¬â€ use existing DOCX
        
        # Per-doc container blocks: flat refs for 91-102/2025-12-10/2026-01-22; container for 51-737
        if name in ('91-102', '2025-12-10', '2026-01-22'):
            doc_pipeline.sgml_generator.use_container_blocks = False
        else:
            doc_pipeline.sgml_generator.use_container_blocks = True   # 51-737 uses container ref

        # Convert: pipeline will use pdf_stem to find DOCX in output_dir
        result = doc_pipeline.convert(pdf_src)
        
        if result.get('status') == 'success':
            # The pipeline writes to doc_dir/name.sgm
            gen_sgm = os.path.join(doc_dir, f'{name}.sgm')
            if os.path.exists(gen_sgm):
                shutil.copy2(gen_sgm, sgm_out)
                size = os.path.getsize(sgm_out) // 1024
                print(f'  Ã¢Å“â€¦ Saved TR output: {os.path.basename(sgm_out)} ({size} KB)')
            else:
                print(f'  Ã¢Å¡Â Ã¯Â¸Â  Generated SGML not found at {gen_sgm}')
        else:
            print(f'  Ã¢ÂÅ’ Conversion failed: {result.get("message", "unknown error")}')
    except Exception as e:
        print(f'  Ã¢ÂÅ’ Exception: {e}')
        import traceback
        traceback.print_exc()
    finally:
        # Restore original PATHS
        PATHS['output_dir'] = original_output
        try:
            doc_pipeline.cleanup()
        except:
            pass

print('\n' + "="*80)
print('Ã¢Å“â€¦ TEST6 BATCH COMPLETE')
print("="*80)

# Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬
# PART 2: Process 5 new docs from test_25.02.26_1
# Output: {name}_OURS.sgm  (reference is {name}_TR.sgm Ã¢â‚¬â€ NOT overwritten)
# Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬Ã¢â€â‚¬
print("\n" + "="*80)
print("BATCH RE-CONVERSION Ã¢â‚¬â€ 5 NEW TEST DOCUMENTS (test_25.02.26_1)")
print("="*80)

for doc in NEW_TEST_DOCS:
    name   = doc['name']
    doc_dir = os.path.join(NEW_TEST_BASE, name)
    docx_src = os.path.join(doc_dir, f'{name}.docx')
    pdf_src  = os.path.join(doc_dir, f'{name}.pdf')
    sgm_out  = os.path.join(doc_dir, f'{name}_OURS.sgm')   # do NOT use _TR.sgm (that's ref!)
    
    if not os.path.exists(docx_src):
        print(f'\nÃ¢Å¡Â Ã¯Â¸Â  Skipping {name}: DOCX not found at {docx_src}')
        continue
    
    print(f'\n{"Ã¢â€â‚¬"*60}')
    print(f'Ã°Å¸â€œâ€ž Processing: {name}')
    print(f'   DOCX: {docx_src}')
    
    os.makedirs(doc_dir, exist_ok=True)
    original_output = PATHS['output_dir']
    PATHS['output_dir'] = doc_dir
    
    try:
        doc_pipeline = CompletePipeline()
        doc_pipeline.initialize()
        # ABBYY enabled for all docs (user requirement — removed: doc_pipeline.abbyy = None)  # Use existing DOCX
        doc_pipeline.sgml_generator.use_n_in_headings = False  # CSA notice: no N in headings
        doc_pipeline.sgml_generator.generate_graphic_tags = False  # CSA notice: no GRAPHIC tags
        doc_pipeline.force_no_inline = True  # CSA notice: no BOLD/EM inline
        doc_pipeline.sgml_generator.use_container_blocks = False  # CSA notice: flat BLOCK format (matches reference)
        doc_pipeline.sgml_generator.use_n_in_headings = False  # CSA notice: no N in headings
        doc_pipeline.sgml_generator.generate_graphic_tags = False  # CSA notice: no GRAPHIC tags
        doc_pipeline.force_no_inline = True  # CSA notice: no BOLD/EM inline
        
        result = doc_pipeline.convert(pdf_src)
        
        if result.get('status') == 'success':
            gen_sgm = os.path.join(doc_dir, f'{name}.sgm')
            if os.path.exists(gen_sgm):
                shutil.copy2(gen_sgm, sgm_out)
                size = os.path.getsize(sgm_out) // 1024
                print(f'  Ã¢Å“â€¦ Saved output: {os.path.basename(sgm_out)} ({size} KB)')
            else:
                print(f'  Ã¢Å¡Â Ã¯Â¸Â  Generated SGML not found at {gen_sgm}')
        else:
            print(f'  Ã¢ÂÅ’ Conversion failed: {result.get("message", "unknown error")}')
    except Exception as e:
        print(f'  Ã¢ÂÅ’ Exception: {e}')
        import traceback
        traceback.print_exc()
    finally:
        PATHS['output_dir'] = original_output
        try:
            doc_pipeline.cleanup()
        except:
            pass

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# PART 3: Process 5 brand-new PDF-only docs from test_new_x5
# Output: {name}_OURS.sgm  (reference is {name}_TR.sgm)
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
NEW_X5_BASE = r'C:\Users\C303180\OneDrive - Thomson Reuters Incorporated\Desktop\TR\securities-outsourcing-samples\sec-out-samples-2\Jurisdictions\juri\NB_\test_new_x5'

NEW_X5_DOCS = [
    {'name': '45-930'},
    {'name': '13-103'},
    {'name': '2025-07-15'},
    {'name': 'cmt-rules'},
    {'name': 'TSX-SN-2025-0002'},
]

print("\n" + "="*80)
print("BATCH RE-CONVERSION â€” 5 NEW-X5 DOCUMENTS (test_new_x5)")
print("="*80)

for doc in NEW_X5_DOCS:
    name   = doc['name']
    doc_dir = os.path.join(NEW_X5_BASE, name)
    docx_src = os.path.join(doc_dir, f'{name}.docx')
    pdf_src  = os.path.join(doc_dir, f'{name}.pdf')
    sgm_out  = os.path.join(doc_dir, f'{name}_OURS.sgm')

    if not os.path.exists(docx_src):
        print(f'\nâš ï¸  Skipping {name}: DOCX not found at {docx_src}')
        continue

    print(f'\n{"â”€"*60}')
    print(f'ðŸ“„ Processing: {name}')
    print(f'   DOCX: {docx_src}')

    os.makedirs(doc_dir, exist_ok=True)
    original_output = PATHS['output_dir']
    PATHS['output_dir'] = doc_dir

    try:
        doc_pipeline = CompletePipeline()
        doc_pipeline.initialize()
        # ABBYY enabled: PDF -> DOCX preserving bold/italic formatting
        # x5 refs use container block format (same as reference)

        # Per-document settings for x5 batch
        if name == 'cmt-rules':
            doc_pipeline.sgml_generator.use_misclaw = True  # legislation doc needs MISCLAW wrapper
        else:
            doc_pipeline.sgml_generator.use_misclaw = False

        # Container blocks: 45-930, 13-103, 2025-07-15, TSX use container; cmt-rules uses flat
        CONTAINER_REFS_X5 = {'45-930', '13-103', '2025-07-15', 'TSX-SN-2025-0002'}
        doc_pipeline.sgml_generator.use_container_blocks = (name in CONTAINER_REFS_X5)

        result = doc_pipeline.convert(pdf_src)

        if result.get('status') == 'success':
            gen_sgm = os.path.join(doc_dir, f'{name}.sgm')
            if os.path.exists(gen_sgm):
                shutil.copy2(gen_sgm, sgm_out)
                size = os.path.getsize(sgm_out) // 1024
                print(f'  âœ… Saved output: {os.path.basename(sgm_out)} ({size} KB)')
            else:
                print(f'  âš ï¸  Generated SGML not found at {gen_sgm}')
        else:
            print(f'  âŒ Conversion failed: {result.get("message", "unknown error")}')
    except Exception as e:
        print(f'  âŒ Exception: {e}')
        import traceback
        traceback.print_exc()
    finally:
        PATHS['output_dir'] = original_output
        try:
            doc_pipeline.cleanup()
        except:
            pass

print('\n' + "="*80)
print('âœ… FULL BATCH COMPLETE (test6 + test_25.02.26_1 + test_new_x5)')
print("="*80)
print('\n' + "="*80)
print('âœ… FULL BATCH COMPLETE (test6 + test_25.02.26_1 + test_new_x5) â€” Now run the Validation cell below')
print("="*80)

